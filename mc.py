import latex2mathml.converter
import numpy as np
import math

from docx.oxml import ns
from lxml import etree
import matplotlib
from matplotlib import pyplot as plt
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Mm, Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION
from course import Table, Value, PerPercentTable, CalculateTable, ActivePassive, DirectCosts, WorkAndOther, ExtendedActivePassive

# formatting
document = Document()

title_text = None
subtitle_text = None
subtitle_2_text = None
main_text = None
table_name_text = None
formula_style = None
formula_style_12 = None
table_style = None
table_style_dense = None
table_style_12 = None
table_style_12_dense = None
table_style_10 = None


def init_styles():
    global main_text, title_text, table_name_text, subtitle_text, subtitle_2_text, \
        formula_style, formula_style_12, table_style, table_style_dense, \
        table_style_12, table_style_12_dense, table_style_10
    FONT_NAME = 'Times New Roman'

    def gen_paragraph_style(name, font_size, first_line_indent=0, space_before=0, space_after=0, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, base=None):
        ps = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base is not None:
            ps.base_style = base
        ps.paragraph_format.first_line_indent = first_line_indent
        ps.paragraph_format.alignment = alignment
        ps.font.name = FONT_NAME
        ps.font.size = font_size
        ps.paragraph_format.space_before = space_before
        ps.paragraph_format.space_after = space_after
        ps.font.bold = False
        return ps

    def gen_table_style(name, font_size, left_indent=None, right_indent=None, space_before=0, space_after=0):
        ts = document.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        ts.base_style = document.styles['Table Grid']
        if left_indent is not None:
            ts.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            ts.paragraph_format.right_indent = right_indent
        ts.paragraph_format.space_before = space_before
        ts.paragraph_format.space_after = space_after
        ts.font.name = FONT_NAME
        ts.font.size = font_size
        ts.next_paragraph_style = main_text
        return ts

    title_text = gen_paragraph_style('Title text', Pt(16), space_after=Pt(12), base=document.styles['Heading 1'])
    subtitle_text = gen_paragraph_style('Subtitle text', Pt(14), space_after=Pt(8), base=document.styles['Heading 2'])
    subtitle_2_text = gen_paragraph_style('Subtitle 2 text', Pt(14), space_after=Pt(8), base=document.styles['Heading 3'])
    main_text = gen_paragraph_style('Main text', Pt(14), Cm(1.25), alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    table_name_text = gen_paragraph_style('Table name text', Pt(12), space_before=Pt(12), space_after=Pt(4), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    formula_style = gen_paragraph_style('Formula style', Pt(14))
    formula_style_12 = gen_paragraph_style('Formula style 12', Pt(12))

    table_style = gen_table_style('Table style', Pt(14), space_before=Mm(1.5), space_after=Mm(1.5))
    table_style = gen_table_style('Table style dense', Pt(14), space_before=Mm(0.5), space_after=Mm(0.5))
    table_style_12 = gen_table_style('Table style 12', Pt(12), space_before=Mm(1.5), space_after=Mm(1.5))
    table_style_12_dense = gen_table_style('Table style 12 dense', Pt(12), space_before=Mm(0.15), space_after=Mm(0.15), left_indent=Mm(0.25), right_indent=Mm(0.25))
    table_style_10 = gen_table_style('Table style 10', Pt(10), space_before=Mm(1.0), space_after=Mm(1.0))


def dp(text='', style=None, no_indent=False):
    if style is None:
        p = document.add_paragraph(text, main_text)
    else:
        p = document.add_paragraph(text, style)
    if no_indent:
        p.paragraph_format.first_line_indent = 0
    return p


def latex_to_word(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def add_formula(latex, p=None, style=None):
    if p is None:
        if style is None:
            p = dp(style=formula_style)
        else:
            p = dp(style=style)
        p.add_run()._element.append(latex_to_word(latex))
    else:
        p.add_run()._element.append(latex_to_word(latex))
    return p


def add_formula_with_description(latex, description, style=None):
    p = add_formula(latex, style=style)
    fs = Pt(14)
    p.add_run('\nГде:\n').font.size = fs
    for i, e in enumerate(description):
        add_formula(e[0], p, formula_style)
        p.add_run(' - ' + e[1] + ('\n' if i + 1 < len(description) else '')).font.size = fs
    return p


def add_table(data, widths=None, first_bold=False, style=None):
    table = document.add_table(rows=len(data), cols=len(data[0]))
    if widths is not None:
        table.autofit = False
        table.allow_autofit = True
    if style is None:
        table.style = table_style
    else:
        table.style = style
    for i in range(len(table.rows)):
        row = table.rows[i]
        for j in range(len(row.cells)):
            cell = row.cells[j]
            if data[i][j] is not None:
                r = cell.paragraphs[0].add_run(data[i][j])
                if first_bold and i == 0:
                    r.font.bold = True
            if widths and widths[j] is not None:
                cell.width = widths[j]
    return table


def add_active_passive_table(active_passive):
    table = document.add_table(rows=21, cols=4)
    table.style = table_style_12_dense
    table.autofit = False
    table.allow_autofit = True
    for i, e in enumerate(['АКТИВ', 'руб', 'ПАССИВ', 'руб']):
        table.cell(0, i).paragraphs[0].add_run(e)
        table.cell(0, i).width = Cm(5.2 if i % 2 == 0 else 3.25)

    tbl = active_passive.to_table()

    for i in range(len(tbl)):
        for j in range(4):
            table.cell(i + 1, j).width = Cm(5.5 if j % 2 == 0 else 3.0)
            if tbl[i][j] is not None:
                if type(tbl[i][j]) != str:
                    table.cell(i + 1, j).paragraphs[0].add_run(fn(tbl[i][j]))
                else:
                    table.cell(i + 1, j).paragraphs[0].add_run(tbl[i][j].strip())
                    if tbl[i][j].startswith(' '):
                        table.cell(i + 1, j).paragraphs[0].paragraph_format.left_indent = Cm(0.5)

    for e in [
        [0, 0], [0, 2],
        [1, 0], [1, 2],
        [6, 0], [6, 2],
        [8, 0], [8, 2],
        [10, 2],
        [12, 2],
        [18, 0], [18, 2],
        [20, 0], [20, 2]
    ]:
        table.cell(e[0], e[1]).paragraphs[0].runs[0].bold = True
        table.cell(e[0], e[1]).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return table


def add_employee_structure_table(data):
    table = add_table(
        [['Должность',
          'Число сотрудников',
          'Тарифная ставка, руб./мес.',
          'Надбавки, руб./мес.',
          'Стимулирующие выплаты, руб./мес.',
          'Категория исполнителей',
          'Система оплаты',
          ]] +
        [[e.name, str(e.amount), fn(e.data, 0), 0, fn(round(e.data / 12), 0), '', ''] for e in data.sl.rows] +
        [[e.name, str(e.amount), fn(e.data, 0), 0, fn(round(e.data / 12), 0), '', ''] for e in data.vpr.rows] +
        [['Рабочий', str(data.R_opr), fn(data.opr_salary), fn(data.opr_extra, 0),
          fn(round(data.opr_salary / 12), 0), 'ОПР', 'Сдельная']],
        [Cm(3.7), Cm(1.75), Cm(2.5), Cm(1.75), Cm(2.75), Cm(2.0), Cm(2.1)],
        style=table_style_12
    )

    table.cell(1, 5).merge(table.cell(len(data.sl), 5))
    table.cell(1, 5).paragraphs[0].add_run('Служащие')

    table.cell(1 + len(data.sl), 5).merge(table.cell(1 + len(data.sl) + len(data.vpr) - 1, 5))
    table.cell(1 + len(data.sl), 5).paragraphs[0].add_run('ВПР')

    table.cell(1, 6).merge(table.cell(1 + len(data.sl) + len(data.vpr) - 1, 6))
    table.cell(1, 6).paragraphs[0].add_run('Повременная')


def add_employee_salary_table(data):
    table = add_table(
        [['Должность',
          'Число сотрудников',
          'Тарифная ставка, руб./год',
          'Надбавки, руб./год',
          'Стимулирующие выплаты, руб./год',
          'Итого, руб./год'
          ]] +
        [[e.name, str(e.amount), fn(e.data * 12, 0), 0, fn(e.data, 0), fn(e.amount * e.data * 13)] for e in data.sl.rows] +
        [[e.name, str(e.amount), fn(e.data * 12, 0), 0, fn(e.data, 0), fn(e.amount * e.data * 13)] for e in data.vpr.rows] +
        [['Рабочий', str(data.R_opr), fn(data.opr_salary * 12, 0), fn(data.opr_extra * 12, 0),
          fn(data.opr_salary, 0), fn(data.FOT_opr + data.FOT_opr_extra, 0)]],
        [Cm(3.7), Cm(2), Cm(2.15), Cm(2.15), Cm(2.5), Cm(3)],
        style=table_style_12
    )
    r = table.add_row()
    r.cells[0].paragraphs[0].add_run('Итого').bold = True
    r.cells[1].paragraphs[0].add_run(fn(data.sl.calc_sum(lambda e, _: e) + data.vpr.calc_sum(lambda e, _: e) + data.R_opr, 0)).bold = True
    r.cells[2].paragraphs[0].add_run(fn(data.sl.calc_sum(lambda e, c: e * c * 12) + data.vpr.calc_sum(lambda e, c: e * c * 12) + data.FOT_opr, 0)).bold = True
    r.cells[3].paragraphs[0].add_run(fn(data.R_opr * data.opr_extra * 12, 0)).bold = True
    r.cells[4].paragraphs[0].add_run(
        fn(data.stimulating_salary_percent * (
                data.sl.calc_sum(lambda e, c: e * c) +
                data.vpr.calc_sum(lambda e, c: e * c) +
                data.R_opr * data.opr_salary
        ), 0)).bold = True
    r.cells[5].paragraphs[0].add_run(fn(data.FOT.total)).bold = True


def add_production_calculation_table(direct: DirectCosts, indirect: WorkAndOther, kom):
    full = direct.direct + indirect.total + kom
    table = add_table([
        ['Статья затрат', 'Сумма, руб./шт.', '%'],
        ['Прямые затраты', None, None],
        ['    материалы и комплектующие изделия', fn(direct.materials_i_comp), fn(direct.materials_i_comp / full * 100, 1)],
        ['    заработная плата ОРП', fn(direct.FOT), fn(direct.FOT / full * 100, 1)],
        ['    страховые взносы', fn(direct.FOT_fee), fn(direct.FOT_fee / full * 100, 1)],
        ['    амортизация', fn(direct.amortisation), fn(direct.amortisation / full * 100, 1)],
        ['Итого прямые затраты', fn(direct.direct), fn(direct.direct / full * 100, 1)],
        [None, None, None],
        ['Косвенные затраты', None, None],
        ['    связанные с работой оборудования', fn(indirect.work), fn(indirect.work / full * 100, 1)],
        ['    не связанные с работой оборудования', fn(indirect.other), fn(indirect.other / full * 100, 1)],
        ['Итого косвенные затраты', fn(indirect.total), fn(indirect.total / full * 100, 1)],
        [None, None, None],
        ['Производственная себестоимость', fn(direct.direct + indirect.total), fn((direct.direct + indirect.total) / full * 100, 1)],
        ['Коммерческие затраты', fn(kom), fn(kom / full * 100, 1)],
        ['Полная себестоимость изделия', fn(full), '100']
    ], [Cm(9), Cm(5), Cm(2)], style=table_style_12_dense, first_bold=True)
    return table


def add_const_and_variable_costs_table(S_sum, FOT, FOT_fee, style=None):
    costs = S_sum['S proizv']
    S_kom = S_sum['S_kom']
    table_wdt = [Cm(1), Cm(4), Cm(3.5), Cm(1), Cm(4), Cm(3.5)]
    table = add_table([
        [None, None, None, None, None, None],
        ['№', 'Условно-постоянные затраты', 'Сумма, тыс.руб./год', '№', 'Переменные затраты', 'Сумма, тыс.руб./год']
    ], table_wdt, style=table_style_12 if style is None else style)
    table.cell(0, 0).merge(table.cell(0, 5)).text = f'Суммарные затраты, руб./год: {fn(S_sum.total)}'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(0, 0).paragraphs[0].runs[0].bold = True

    const = [
        (S_kom._display_name, S_kom.const),
        (costs['fuel non tech']._display_name, costs['fuel non tech'].const),
        ('З/п кроме ОПР', FOT.const),
        ('Страховые взносы кроме ОПР', FOT_fee.const),
        (costs['amortisation']._display_name, costs['amortisation'].const),
        (costs['inventory']._display_name, costs['inventory'].const),
        (costs['move save']._display_name, costs['move save'].const),
        (costs['extra']._display_name, costs['extra'].const)
    ]

    variable = [
        (S_kom._display_name, S_kom.variable),
        (costs['fuel tech']._display_name, costs['fuel tech'].variable),
        ('З/п ОПР', FOT.variable),
        ('Страховые взносы ОПР', FOT_fee.variable),
        (costs['material_main']._display_name, costs['material_main'].variable),
        (costs['helper']._display_name, costs['helper'].variable),
        (costs['move save']._display_name, costs['move save'].variable),
        (costs['extra']._display_name, costs['extra'].variable)
    ]

    for i, z in enumerate(zip(const, variable)):
        c = z[0]
        v = z[1]
        r = table.add_row()
        r.cells[0].paragraphs[0].add_run(str(i + 1))
        r.cells[1].paragraphs[0].add_run(c[0])
        r.cells[2].paragraphs[0].add_run(fn(c[1]))
        r.cells[3].paragraphs[0].add_run(str(i + 1))
        r.cells[4].paragraphs[0].add_run(v[0])
        r.cells[5].paragraphs[0].add_run(fn(v[1]))
        for j, e in enumerate(table_wdt):
            r.cells[j].width = e

    r = table.add_row()
    r.cells[0].merge(r.cells[1]).paragraphs[0].add_run('Итого').bold = True
    r.cells[2].paragraphs[0].add_run(fn(sum([e[1] for e in const])))
    r.cells[3].merge(r.cells[4]).paragraphs[0].add_run('Итого').bold = True
    r.cells[5].paragraphs[0].add_run(fn(sum([e[1] for e in variable])))


def fn(num, ln=2):
    return ('{:,.' + str(ln) + 'f}').format(num)


class InitialData:
    __slots__ = [
        'N_pl',
        'materials', 'accessories', 'operations',
        'materials_A', 'accessories_A', 'operations_A',
        'materials_B', 'accessories_B', 'operations_B',
        'materials_C', 'accessories_C', 'operations_C',
    ]

    def __init__(self, N_pl=45_000):
        self.N_pl = N_pl
        self.materials_B = self.materials = Table('name', 'cost', 'amount', 't_zap')

        self.materials.add_row('а', 60, 1, 30)
        self.materials.add_row('б', 150, 3, 40)
        self.materials.add_row('в', 350, 3, 60)
        self.materials.add_row('г', 60, 2, 50)

        self.accessories_B = self.accessories = Table('name', 'cost', 'amount', 't_zap')
        self.accessories.add_row('а', 50, 1, 35)
        self.accessories.add_row('б', 60, 3, 70)
        self.accessories.add_row('в', 70, 2, 45)

        self.operations_B = self.operations = Table('cost', 'time', 'name')
        self.operations.add_row(0, 0.4, '-ручная операция-')
        self.operations.add_row(500_000, 0.3, 'б')
        self.operations.add_row(600_000, 0.2, 'в')
        self.operations.add_row(700_000, 0.6, 'г')

        self.materials_A = Table('name', 'cost', 'amount', 't_zap')
        self.materials_A.add_row('-', 0, 0, 0)
        self.materials_A.add_row('б', 50, 1, 30)
        self.materials_A.add_row('в', 50, 8, 30)
        self.materials_A.add_row('г', 15, 8, 30)

        self.accessories_A = Table('name', 'cost', 'amount', 't_zap')
        self.accessories_A.add_row('-', 0, 0, 0)
        self.accessories_A.add_row('б', 70, 4, 30)
        self.accessories_A.add_row('в', 40, 2, 30)

        self.operations_A = Table('cost', 'time', 'name')
        self.operations_A.add_row(1_200_000, 0.2, 'а')
        self.operations_A.add_row(1_400_000, 0.1, 'б')
        self.operations_A.add_row(1_500_000, 0.3, 'в')
        self.operations_A.add_row(1_600_000, 0.5, 'г')
        self.operations_A.add_row(1_700_000, 0.5, 'д')
        self.operations_A.add_row(1_300_000, 1.0, 'е')
        self.operations_A.add_row(0, 0.3, '-ручная операция-')

        self.materials_C = Table('name', 'cost', 'amount', 't_zap')
        self.materials_C.add_row('а', 35, 1, 30)
        self.materials_C.add_row('б', 50, 1, 50)
        self.materials_C.add_row('-', 0, 0, 0)
        self.materials_C.add_row('г', 15, 4, 75)

        self.accessories_C = Table('name', 'cost', 'amount', 't_zap')
        self.accessories_C.add_row('а', 70, 1, 35)
        self.accessories_C.add_row('-', 0, 0, 0)
        self.accessories_C.add_row('в', 40, 2, 40)

        self.operations_C = Table('cost', 'time', 'name')
        self.operations_C.add_row(1_200_000, 0.2, 'а')
        self.operations_C.add_row(1_400_000, 0.1, 'б')
        self.operations_C.add_row(1_600_000, 1.0, 'г')
        self.operations_C.add_row(1_700_000, 0.5, 'д')


class Chapter_1:
    __slots__ = [
        'T_pl', 'B', 'C', 'D', 'O', 'H', 'gamma', 'F_ob_ef',
        'N_machine_work_percent', 'machines',
        'TO_perv', 'main_resources', 'S_os', 'S_os_amortisable'
    ]

    def __init__(self, initial_data: InitialData):
        def calc_1_1_n_ob_k(N, t, b, F, ceil=True):
            numerator = N * t
            denominator = b * F
            if ceil:
                return math.ceil(numerator / denominator)
            return numerator / denominator

        self.T_pl = 365
        self.B = 118
        self.C = 2
        self.D = 8
        self.O = 20
        self.H = 20
        self.gamma = 0.05
        self.F_ob_ef = (self.T_pl - self.B) * self.C * self.D * (1 - self.gamma)
        self.N_machine_work_percent = 0.75

        self.machines = Table('name', 'cost', 'n rasch', 'n fact', 'b_fact')

        self.TO_perv = 0

        for operation in initial_data.operations.rows:
            name = operation['name']

            if len(name) == 1:
                e0 = calc_1_1_n_ob_k(initial_data.N_pl, operation['time'], self.N_machine_work_percent, self.F_ob_ef, False)
                e = math.ceil(e0)
                self.machines.add_row(name, operation['cost'], e0, e, calc_1_1_n_ob_k(initial_data.N_pl, operation['time'], e, self.F_ob_ef, False))
                self.TO_perv += operation['cost'] * e

        self.main_resources = Table('n', 'name', '%', 'cost')
        mr = self.main_resources
        self.S_os = self.TO_perv / 0.4

        main_resources = [
            ['1', 'Земля', 0.14],
            ['2', 'Здания', 0.09],
            ['3', 'Сооружения', 0.07],
            ['4', 'Передаточные устройства', 0.06],
            ['5', 'Машины и оборудование, в т.ч.', 0.50],
            ['', '- технологическое оборудование', 0.40],
            ['', '- нетехнологические машины и оборудование', 0.10],
            ['7', 'Транспортные средства', 0.09],
            ['8', 'Инструменты и технологическая оснастка', 0.04],
            ['9', 'Производственный и хозяйственный инвентарь', 0.01],
        ]

        self.S_os_amortisable = round(self.S_os * (1.0 - 0.14), 2)

        for e in main_resources:
            mr.add_row(e[0], e[1], e[2], self.S_os * e[2])


class Chapter_2:
    __slots__ = [
        'F_rab_ef',
        'C_opr_mean', 'p_mean',
        'R_opr_raw', 'R_opr', 'FOT_opr', 'FOT_opr_extra', 'opr_extra', 'opr_salary',
        'vpr', 'R_vpr', 'FOT_vpr',
        'sl', 'R_sl', 'FOT_sl',
        'R_ppp', 'FOT',
        'insurance_fee', 'FOT_fee', 'FOT_with_fee',
        'stimulating_salary_percent'
    ]

    def __init__(self, initial_data: InitialData, chapter_1: Chapter_1, const=None):
        self.F_rab_ef = (chapter_1.T_pl - chapter_1.B - chapter_1.O - chapter_1.H) * chapter_1.D

        total_time = initial_data.operations.calculate_sum(lambda x: x['time'])

        self.R_opr_raw = initial_data.N_pl * total_time / self.F_rab_ef
        self.R_opr = int(math.ceil(self.R_opr_raw))

        self.vpr = PerPercentTable(self.R_opr, True, False)
        self.vpr.add_row('Настройщик оборудования', 0.05, 60_000)
        self.vpr.add_row('Складовщик', 0.07, 50_000)
        self.vpr.add_row('Уборщик', 0.05, 30_000)
        self.vpr.add_row('Контролёр ОТК', 0.07, 80_000)

        self.sl = PerPercentTable(self.R_opr, True, False)
        self.sl.add_row('Генеральный директор', 0, 120_000)
        self.sl.add_row('HR', 0, 70_000)
        self.sl.add_row('Менеджер по закупу', 0, 70_000)
        self.sl.add_row('Менеджер по производству', 0, 85_000)
        self.sl.add_row('Инженер', 2. / self.R_opr, 85_000)
        self.sl.add_row('Бухгалтер', 2. / self.R_opr, 80_000)
        self.sl.add_row('Охраниик', 4. / self.R_opr, 35_000)
        self.sl.add_row('Логист', 3. / self.R_opr, 50_000)
        self.sl.add_row('Курьер', 0.5 - 10. / self.R_opr, 45_000)

        self.R_vpr = self.vpr.total
        self.R_sl = self.sl.total

        self.opr_salary = 50_000
        self.opr_extra = 5_000

        self.stimulating_salary_percent = 1.0

        self.C_opr_mean = 12 * self.opr_salary / self.F_rab_ef
        self.p_mean = self.C_opr_mean * total_time / len(initial_data.operations.rows)

        self.FOT_opr = self.p_mean * initial_data.N_pl * len(initial_data.operations)
        self.FOT_opr_extra = self.R_opr * (self.opr_extra * 12 + self.stimulating_salary_percent * self.opr_salary)

        self.FOT_vpr = self.vpr.calc_sum(lambda amount, salary: amount * salary * (12 + self.stimulating_salary_percent))
        self.FOT_sl = self.sl.calc_sum(lambda amount, salary: amount * salary * (12 + self.stimulating_salary_percent))

        self.R_ppp = self.R_opr + self.R_vpr + self.R_sl
        if const is None or const['fot'] is None:
            self.FOT = Value('fot', self.FOT_vpr + self.FOT_sl, self.FOT_opr + self.FOT_opr_extra, 'Затраты на оплату труда')
        else:
            self.FOT = Value('fot', const['fot'].const, self.FOT_opr + self.FOT_opr_extra, 'Затраты на оплату труда')

        self.insurance_fee = PerPercentTable(self.FOT.total)
        self.insurance_fee.add_row('ОПФ', 0.22)
        self.insurance_fee.add_row('ФОМС', 0.051)
        self.insurance_fee.add_row('ФСС', 0.029)
        self.insurance_fee.add_row('Страхование от несчастных случаев на производстве и профессиональных заболеваний', 0.04)

        self.FOT_fee = Value('fot fee', round(self.FOT.const * 0.34, 2), round(self.FOT.variable * 0.34, 2), 'Страховые взносы')
        self.FOT_with_fee = Value('fot total', display_name='ФОТ')
        self.FOT_with_fee.add_child(self.FOT)
        self.FOT_with_fee.add_child(self.FOT_fee)


class Chapter_3:
    __slots__ = [
        'S_mat_i_comp',
        'main_materials',
        'help_materials_percent',
        'moving_save_percent',
        'move_save_const_percent',
        'inventory_percent',
        'fuel_percent',
        'fuel_tech_percent',
        'fuel_non_tech_percent',

        'OS_amortisation_percent',
        'OS_amortisation',

        'NMA',
        'NMA_amortisation_percent',
        'NMA_amortisation',

        'OS_fix_percent',
        'OS_fix',

        'extra_percent',

        'costs',
        'S_pr_tek_pl'
    ]

    def calc_mat_costs(self, n, fot, fot_safe, const=None):
        pass

    def __init__(self, initial_data: InitialData, chapter_1: Chapter_1, chapter_2: Chapter_2, const=None):
        self.S_mat_i_comp = initial_data.materials.calculate_sum(lambda x: x['cost'] * x['amount']) + initial_data.accessories.calculate_sum(lambda x: x['cost'] * x['amount'])

        self.costs = Value('S proizv', display_name='Затраты')
        self.main_materials = Value('material', variable=initial_data.N_pl * self.S_mat_i_comp, display_name='Материальные затраты')
        m_base = self.main_materials.total

        self.main_materials.add_child(Value('material_main', 0, m_base, display_name='Основные материалы'))

        self.help_materials_percent = 0.05
        self.main_materials.add_child(Value('helper', 0, round(m_base * self.help_materials_percent, 2), display_name='Вспомогательные материалы'))

        self.moving_save_percent = 0.12
        self.move_save_const_percent = 0.3
        ms = round(m_base * self.moving_save_percent, 2)
        msc = round(ms * self.move_save_const_percent, 2) if const is None or const['move save'] is None else const['move save'].const
        msv = round(ms - msc, 2) if const is None or const['move save'] is None else round((1 - self.move_save_const_percent) * ms, 2)
        self.main_materials.add_child(Value('move save', msc, msv, display_name='Транспортно-заготовительные расходы'))

        self.inventory_percent = 0.03
        self.main_materials.add_child(
            Value('inventory', const=round(m_base * self.inventory_percent, 2), display_name='Инструменты, инвентарь')
            if const is None or const['inventory'] is None else const['inventory'])

        self.fuel_percent = 0.55
        fuelt = round(m_base * self.fuel_percent, 2)

        self.fuel_tech_percent = 0.7
        self.fuel_non_tech_percent = 1.0 - self.fuel_tech_percent

        fuel_energy_costs = self.main_materials.add_child(Value('fuel total', display_name='Топливо и энергия'))
        ft = fuel_energy_costs.add_child(Value('fuel tech', variable=round(fuelt * self.fuel_tech_percent, 2), display_name='Технологическое топливо и энергия'))
        fuel_energy_costs.add_child(
            Value('fuel non tech', const=fuelt - ft.total, display_name='Нетехнологическое топливо и энергия')
            if const is None or const['fuel non tech'] is None else const['fuel non tech'])
        self.costs.add_child(self.main_materials)

        self.costs.add_child(chapter_2.FOT)
        self.costs.add_child(chapter_2.FOT_fee)

        self.OS_amortisation_percent = 0.1
        self.OS_amortisation = round(self.OS_amortisation_percent * chapter_1.S_os_amortisable, 2)

        self.NMA = 3_000_000
        self.NMA_amortisation_percent = 0.1
        self.NMA_amortisation = round(self.NMA_amortisation_percent * self.NMA, 2)

        aos = Value('amortisation OS', const=self.OS_amortisation, display_name='Амортизация ОС')
        anma = Value('amortisation NMA', const=self.NMA_amortisation, display_name='Амортизация НМА')
        a = Value('amortisation', display_name='Амортизация ОС и НМА')
        a.add_child(aos)
        a.add_child(anma)
        self.costs.add_child(a)

        self.OS_fix_percent = 0.06
        self.OS_fix = round(self.OS_fix_percent * chapter_1.S_os_amortisable, 2)

        self.extra_percent = 0.05
        self.costs.add_child(Value('extra', const=self.costs.const * self.extra_percent + self.OS_fix, variable=self.costs.variable * self.extra_percent, display_name='Прочие затраты'))

        self.S_pr_tek_pl = self.costs.total


class Chapter_4:
    __slots__ = [
        'N_pl_values',
        'S_b_proizv',
        'S_b_poln',
        'S_kom_percent',
        'S_kom_const_percent',
        'S_kom',
        'S_sum',
        'ct1'
    ]

    def calc_n(self, n):
        fake_initial = InitialData(n)
        fake_chapter_2 = Chapter_2(fake_initial, chapter_1, chapter_2.FOT)
        fake_chapter_3 = Chapter_3(fake_initial, chapter_1, fake_chapter_2, const=chapter_3.costs)
        fake_chapter_4 = Chapter_4(fake_initial, fake_chapter_3, self.S_kom)
        return fake_chapter_4.S_sum.head()

    def __init__(self, initial_data_fm: InitialData, chapter_3: Chapter_3, const=None):
        self.N_pl_values = [450, 2700, 7200, 18900, 33750, 45000]
        self.S_b_proizv = chapter_3.S_pr_tek_pl / initial_data_fm.N_pl
        self.S_kom_percent = 0.04
        self.S_kom_const_percent = 0.6
        s = round(self.S_b_proizv * self.S_kom_percent, 0) * initial_data_fm.N_pl
        sc = round(s * self.S_kom_const_percent, 2) if const is None or const['S_kom'] is None else const['S_kom'].const
        sv = round(s - sc, 2) if const is None or const['S_kom'] is None else s * (1 - self.S_kom_const_percent) * initial_data_fm.N_pl / initial_data.N_pl

        self.S_sum = Value('S_sum', display_name='Суммарные затраты')
        self.S_sum.add_child(chapter_3.costs)
        self.S_kom = Value('S_kom', sc, sv, 'Коммерческие затраты')
        self.S_sum.add_child(self.S_kom)
        self.S_b_poln = Value('S_b_poln', self.S_sum.const / initial_data_fm.N_pl, self.S_sum.variable / initial_data_fm.N_pl)

        if const is None:
            self.ct1 = CalculateTable(self.N_pl_values, Chapter_4.calc_n, self)
        else:
            self.ct1 = None


class Chapter_5:
    __slots__ = [
        'K_ob_sr_mk',
        'k_ob_sr_percent',
        'K_ob_sr_pr_zap',

        'k_nz',
        'gamma_cycle',
        'T_cycle',
        'K_ob_nez_pr',

        't_real',
        'K_ob_got_prod',

        'gamma_ob',
        'K_ob_extra',
        'K_ob_sum',
    ]

    def __init__(self, initial_data: InitialData, chapter_1: Chapter_1, chapter_3: Chapter_3, chapter_4: Chapter_4):
        mz = round(initial_data.materials.calculate_sum(lambda x: x['amount'] * x['cost'] * initial_data.N_pl / chapter_1.T_pl * x['t_zap']), 2)
        cz = round(initial_data.accessories.calculate_sum(lambda x: x['amount'] * x['cost'] * initial_data.N_pl / chapter_1.T_pl * x['t_zap']), 2)
        self.K_ob_sr_mk = round(mz + cz, 2)
        self.k_ob_sr_percent = 0.4
        self.K_ob_sr_pr_zap = round((1 + self.k_ob_sr_percent) * self.K_ob_sr_mk, 2)
        self.k_nz = (chapter_3.S_mat_i_comp + chapter_4.S_b_proizv) / (chapter_4.S_b_proizv * 2)
        self.gamma_cycle = 50

        self.T_cycle = round(initial_data.operations.calculate_sum(lambda x: x['time']) * self.gamma_cycle /
                             (chapter_1.C * chapter_1.D) * chapter_1.T_pl / (chapter_1.T_pl - chapter_1.B), 3)

        self.K_ob_nez_pr = chapter_4.S_b_proizv * initial_data.N_pl / chapter_1.T_pl * self.k_nz * self.T_cycle

        self.t_real = 10
        self.K_ob_got_prod = round(chapter_4.S_b_proizv * initial_data.N_pl / chapter_1.T_pl * self.t_real, 2)

        self.gamma_ob = 0.6
        a = round(self.K_ob_sr_pr_zap + self.K_ob_nez_pr + self.K_ob_got_prod, 2)
        self.K_ob_sum = round(a / self.gamma_ob, 2)
        self.K_ob_extra = round(self.K_ob_sum - a, 2)


class Chapter_6:
    __slots__ = [
        'k_ob_RPB_percent',
        'ustavnoy_capital_percent',
        'doldosroch_zaemn_sredstva_percent',
        'kratkosroch_zaemn_sredstva_percent',
        'active_passive'
    ]

    def __init__(self, chapter_1: Chapter_1, chapter_5: Chapter_5):
        self.active_passive = ActivePassive()

        self.active_passive.NMA = chapter_3.NMA
        self.active_passive.OS = chapter_1.S_os

        self.k_ob_RPB_percent = 0.3
        self.active_passive.K_ob_RBP = round(chapter_5.K_ob_extra * self.k_ob_RPB_percent, 2)
        self.active_passive.K_ob_sr_pr_zap = chapter_5.K_ob_sr_pr_zap
        self.active_passive.K_ob_ds = chapter_5.K_ob_sum - (chapter_5.K_ob_sr_pr_zap + self.active_passive.K_ob_RBP)
        self.active_passive.K_ob_extra = chapter_5.K_ob_extra - self.active_passive.K_ob_RBP

        self.ustavnoy_capital_percent = 0.8
        self.active_passive.ustavnoy_kapital = round(self.active_passive.active * self.ustavnoy_capital_percent, 2)

        S_summ_passiv_left = self.active_passive.active - self.active_passive.ustavnoy_kapital

        self.doldosroch_zaemn_sredstva_percent = 0.6
        self.active_passive.doldosroch_zaemn_sredstva = round(S_summ_passiv_left * self.doldosroch_zaemn_sredstva_percent, 2)

        self.kratkosroch_zaemn_sredstva_percent = 0.25
        self.active_passive.kratkosroch_zaem_sredstva = round(S_summ_passiv_left * self.kratkosroch_zaemn_sredstva_percent, 2)
        self.active_passive.kratkosroch_prochee = round(S_summ_passiv_left - self.active_passive.doldosroch_zaemn_sredstva - self.active_passive.kratkosroch_zaem_sredstva, 2)


class Chapter_7:
    __slots__ = [
        'tax',
        'net_profit_percent',
        'net_profit',
        'profit_before_tax',
        'k_nats',
        'P_b_poln',
        'P_b_perem',
        'P_proizv_plan',
        'price_fact_percent',
        'P_fact'
    ]

    def __init__(self, chapter_4: Chapter_4, chapter_6: Chapter_6):
        self.tax = 0.2
        self.net_profit_percent = 0.6
        self.net_profit = round(chapter_6.active_passive.ustavnoy_kapital * self.net_profit_percent, 2)
        self.profit_before_tax = round(self.net_profit / (1 - self.tax), 2)
        self.k_nats = self.profit_before_tax / chapter_4.S_sum.total
        self.P_b_poln = round(chapter_4.S_b_poln.total * (1 + self.k_nats), 2)
        self.P_b_perem = round(chapter_4.S_b_poln.variable * (1 + (self.profit_before_tax + chapter_4.S_sum.const) / chapter_4.S_sum.variable), 2)
        self.P_proizv_plan = max(self.P_b_poln, self.P_b_perem)
        self.price_fact_percent = 0.94
        self.P_fact = round(self.P_proizv_plan * self.price_fact_percent, 2)


class Chapter_8:
    __slots__ = [
        'N_fact_percent',
        'N_fact',
        'N_ost',
        'Q_plan', 'Q_fact',
        'K_ob_got_prod_plan', 'K_ob_got_prod_fact',
        'S_pr_got_pr_plan', 'S_pr_got_pr_fact',
        'S_valovaya_plan', 'S_valovaya_fact',
        'kom_percent',
        'S_kom_plan', 'S_kom_fact',
        'P_pr_plan', 'P_pr_fact',
        'P_pr_do_nalogov_plan', 'P_pr_do_nalogov_fact',
        'pr_dir_fact_percent',
        'S_prochie_dohidy_i_rashody_plan', 'S_prochie_dohidy_i_rashody_fact',
        'nalog_na_pribil_plan', 'nalog_na_pribil_fact',
        'P_chistaya_plan', 'P_chistaya_fact'
    ]

    def __init__(self, initial_data: InitialData, chapter_3: Chapter_3, chapter_4: Chapter_4, chapter_5: Chapter_5, chapter_7: Chapter_7):
        self.N_fact_percent = 0.95
        self.N_fact = int(self.N_fact_percent * initial_data.N_pl)
        self.N_ost = initial_data.N_pl - self.N_fact

        self.Q_plan = chapter_7.P_proizv_plan * initial_data.N_pl
        self.Q_fact = chapter_7.P_fact * self.N_fact

        self.K_ob_got_prod_plan = chapter_5.K_ob_got_prod
        self.K_ob_got_prod_fact = chapter_5.K_ob_got_prod + round(chapter_4.S_b_proizv * self.N_ost, 2)
        self.S_pr_got_pr_plan = chapter_3.S_pr_tek_pl - chapter_5.K_ob_nez_pr - self.K_ob_got_prod_plan
        self.S_pr_got_pr_fact = chapter_3.S_pr_tek_pl - chapter_5.K_ob_nez_pr - self.K_ob_got_prod_fact

        self.S_valovaya_plan = self.Q_plan - self.S_pr_got_pr_plan
        self.S_valovaya_fact = self.Q_fact - self.S_pr_got_pr_fact

        self.kom_percent = 0.94
        self.S_kom_plan = chapter_4.S_kom.total
        self.S_kom_fact = self.kom_percent * chapter_4.S_kom.total

        self.P_pr_plan = self.S_valovaya_plan - self.S_kom_plan
        self.P_pr_fact = self.S_valovaya_fact - self.S_kom_fact

        self.P_pr_do_nalogov_plan = chapter_7.profit_before_tax
        self.pr_dir_fact_percent = 0.93
        self.S_prochie_dohidy_i_rashody_plan = self.P_pr_plan - self.P_pr_do_nalogov_plan
        self.S_prochie_dohidy_i_rashody_fact = round(self.S_prochie_dohidy_i_rashody_plan * self.pr_dir_fact_percent, 2)
        self.P_pr_do_nalogov_fact = self.P_pr_fact - self.S_prochie_dohidy_i_rashody_fact
        self.nalog_na_pribil_plan = self.P_pr_do_nalogov_plan - chapter_7.net_profit
        self.nalog_na_pribil_fact = round(self.P_pr_do_nalogov_fact * chapter_7.tax, 2)

        self.P_chistaya_plan = chapter_7.net_profit
        self.P_chistaya_fact = self.P_pr_do_nalogov_fact - self.nalog_na_pribil_fact


class Chapter_9:
    __slots__ = [
        'K_den_sr_plan', 'K_den_sr_fact',
        'K_den_sr_konez_plan', 'K_den_sr_konez_fact',
        'S_kratkosroch_zaem_sredstva_konez_plan',
        'S_kratkosroch_zaem_sredstva_konez_fact',
        'active_passive_plan', 'active_passive_fact',
        'valid_to_cope_kz_plan', 'valid_to_cope_kz_fact',
    ]

    def __init__(self):
        amortisation = chapter_3.costs['amortisation'].total

        self.K_den_sr_plan = chapter_6.active_passive.K_ob_ds + amortisation + chapter_8.P_chistaya_plan - (chapter_5.K_ob_nez_pr + chapter_8.K_ob_got_prod_plan)
        if self.K_den_sr_plan - 500_000 > chapter_6.active_passive.kratkosroch_zaem_sredstva:
            self.K_den_sr_konez_plan = self.K_den_sr_plan - chapter_6.active_passive.kratkosroch_zaem_sredstva
            self.S_kratkosroch_zaem_sredstva_konez_plan = 0
            self.valid_to_cope_kz_plan = 'full'
        elif self.K_den_sr_plan > 500_000:
            self.K_den_sr_konez_plan = 500_000
            self.S_kratkosroch_zaem_sredstva_konez_plan = chapter_6.active_passive.kratkosroch_zaem_sredstva - (self.K_den_sr_plan - 500_000)
            self.valid_to_cope_kz_plan = 'part'
        else:
            self.S_kratkosroch_zaem_sredstva_konez_plan = chapter_6.active_passive.kratkosroch_zaem_sredstva
            self.K_den_sr_konez_plan = self.K_den_sr_plan
            self.valid_to_cope_kz_plan = 'none'

        self.K_den_sr_fact = chapter_6.active_passive.K_ob_ds + amortisation + chapter_8.P_chistaya_fact - (chapter_5.K_ob_nez_pr + chapter_8.K_ob_got_prod_fact)
        if self.K_den_sr_fact - 500_000 > chapter_6.active_passive.kratkosroch_zaem_sredstva:
            self.K_den_sr_konez_fact = self.K_den_sr_fact - chapter_6.active_passive.kratkosroch_zaem_sredstva
            self.S_kratkosroch_zaem_sredstva_konez_fact = 0
            self.valid_to_cope_kz_fact = 'full'
        elif self.K_den_sr_fact > 500_000:
            self.K_den_sr_konez_fact = 500_000
            self.S_kratkosroch_zaem_sredstva_konez_fact = chapter_6.active_passive.kratkosroch_zaem_sredstva - (self.K_den_sr_fact - 500_000)
            self.valid_to_cope_kz_fact = 'part'
        else:
            self.S_kratkosroch_zaem_sredstva_konez_fact = chapter_6.active_passive.kratkosroch_zaem_sredstva
            self.K_den_sr_konez_fact = self.K_den_sr_fact
            self.valid_to_cope_kz_fact = 'none'

        p = self.active_passive_plan = ActivePassive()
        f = self.active_passive_fact = ActivePassive()

        f.NMA = p.NMA = chapter_3.NMA - chapter_3.costs['amortisation NMA'].total
        f.OS = p.OS = chapter_1.S_os - chapter_3.costs['amortisation OS'].total

        f.K_ob_sr_pr_zap = p.K_ob_sr_pr_zap = chapter_5.K_ob_sr_pr_zap
        f.K_ob_nez_pr = p.K_ob_nez_pr = chapter_5.K_ob_nez_pr
        p.K_ob_got_prod = chapter_5.K_ob_got_prod
        f.K_ob_RBP = p.K_ob_RBP = chapter_6.active_passive.K_ob_RBP
        f.K_ob_extra = p.K_ob_extra = chapter_6.active_passive.K_ob_extra
        p.K_ob_ds = self.K_den_sr_konez_plan

        p.ustavnoy_kapital = chapter_6.active_passive.ustavnoy_kapital
        p.neraspred_pribil = chapter_8.P_chistaya_plan

        p.doldosroch_zaemn_sredstva = chapter_6.active_passive.doldosroch_zaemn_sredstva

        p.kratkosroch_zaem_sredstva = 0
        p.kratkosroch_prochee = chapter_6.active_passive.kratkosroch_prochee

        f.K_ob_got_prod = chapter_8.K_ob_got_prod_fact
        f.K_ob_ds = self.K_den_sr_konez_fact

        f.ustavnoy_kapital = chapter_6.active_passive.ustavnoy_kapital
        f.neraspred_pribil = chapter_8.P_chistaya_fact

        f.doldosroch_zaemn_sredstva = chapter_6.active_passive.doldosroch_zaemn_sredstva

        f.kratkosroch_zaem_sredstva = self.S_kratkosroch_zaem_sredstva_konez_fact
        f.kratkosroch_prochee = chapter_6.active_passive.kratkosroch_prochee


class Chapter_10:
    __slots__ = [
        'k_sob_ob_sr_plan', 'k_sob_ob_sr_fact',
        'k_obespech_sob_sr_plan', 'k_obespech_sob_sr_fact',
        'k_abs_likvid_plan', 'k_abs_likvid_fact',
        'k_tek_likvid_plan', 'k_tek_likvid_fact',
        'V',
        'OS_year_mean',
        'k_FO_plan', 'k_FO_fact',
        'k_FE_plan', 'k_FE_fact',
        'Z_ob_sr_year_mean_plan', 'Z_ob_sr_year_mean_fact',
        'Z_ob_plan', 'Z_ob_fact',
        'S_sobstv_cap_year_mean_plan', 'S_sobstv_cap_year_mean_fact',
        'k_oborach_sobstv_capital_plan', 'k_oborach_sobstv_capital_fact',
        'R_production_plan', 'R_production_fact',
        'R_sell_plan', 'R_sell_fact',
        'R_active_plan', 'R_active_fact',
        'R_sobstv_capital_plan', 'R_sobstv_capital_fact',

        'N_kr',
        'Q_kr',
        'k_pokr',
        'Q_fin_pr_plan', 'Q_fin_pr_fact',
        'proizv_richag_plan', 'proizv_richag_fact'
    ]

    @staticmethod
    def calc_n(n):
        fake_initial = InitialData(n)
        fake_chapter_2 = Chapter_2(fake_initial, chapter_1, chapter_2.FOT)
        fake_chapter_3 = Chapter_3(fake_initial, chapter_1, fake_chapter_2, const=chapter_3.costs)
        fake_chapter_4 = Chapter_4(fake_initial, fake_chapter_3, chapter_4.S_kom)
        return fake_chapter_4.S_sum

    @staticmethod
    def bin_search():
        N_kr_left = 0
        N_kr_right = initial_data.N_pl

        while N_kr_right - N_kr_left > 1:
            N_kr_mean = round((N_kr_right + N_kr_left) / 2)

            s = Chapter_10.calc_n(N_kr_mean)

            if N_kr_mean * chapter_7.P_proizv_plan > s['S_sum'].total:
                N_kr_right = N_kr_mean
            else:
                N_kr_left = N_kr_mean

        return N_kr_left

    @staticmethod
    def calc_k_pokr(n):
        fake_initial = InitialData(n)
        fake_chapter_2 = Chapter_2(fake_initial, chapter_1, const=chapter_2.FOT)
        fake_chapter_3 = Chapter_3(fake_initial, chapter_1, fake_chapter_2, const=chapter_3.costs)
        fake_chapter_4 = Chapter_4(fake_initial, fake_chapter_3, const=chapter_4.S_kom)

        k_pokr = (chapter_7.P_proizv_plan - fake_chapter_4.S_b_poln.variable) / chapter_7.P_proizv_plan

        return k_pokr

    def __init__(self, initial_data: InitialData,
                 chapter_1: Chapter_1, chapter_2: Chapter_2, chapter_3: Chapter_3, chapter_4: Chapter_4,
                 chapter_6: Chapter_6, chapter_7: Chapter_7, chapter_8: Chapter_8, chapter_9: Chapter_9):
        self.k_sob_ob_sr_plan = chapter_9.active_passive_plan.r2 - chapter_9.active_passive_plan.r5
        self.k_sob_ob_sr_fact = chapter_9.active_passive_fact.r2 - chapter_9.active_passive_fact.r5

        self.k_obespech_sob_sr_plan = self.k_sob_ob_sr_plan / chapter_9.active_passive_plan.r2
        self.k_obespech_sob_sr_fact = self.k_sob_ob_sr_fact / chapter_9.active_passive_fact.r2

        self.k_abs_likvid_plan = chapter_9.active_passive_plan.K_ob_ds / chapter_9.active_passive_plan.r5
        self.k_abs_likvid_fact = chapter_9.active_passive_fact.K_ob_ds / chapter_9.active_passive_fact.r5

        self.k_tek_likvid_plan = chapter_9.active_passive_plan.r2 / chapter_9.active_passive_plan.r5
        self.k_tek_likvid_fact = chapter_9.active_passive_fact.r2 / chapter_9.active_passive_fact.r5

        self.V = initial_data.N_pl / chapter_2.R_ppp

        self.OS_year_mean = round(chapter_1.S_os_amortisable - chapter_3.costs['amortisation OS'].total * 0.5, 2)
        self.k_FO_plan = chapter_8.Q_plan / self.OS_year_mean
        self.k_FO_fact = chapter_8.Q_fact / self.OS_year_mean

        self.k_FE_plan = 1 / self.k_FO_plan
        self.k_FE_fact = 1 / self.k_FO_fact

        self.Z_ob_sr_year_mean_plan = round((chapter_6.active_passive.r2 + chapter_9.active_passive_plan.r2) * 0.5, 2)
        self.Z_ob_sr_year_mean_fact = round((chapter_6.active_passive.r2 + chapter_9.active_passive_fact.r2) * 0.5, 2)
        self.Z_ob_plan = chapter_8.Q_plan / self.Z_ob_sr_year_mean_plan
        self.Z_ob_fact = chapter_8.Q_fact / self.Z_ob_sr_year_mean_fact

        self.S_sobstv_cap_year_mean_plan = round((chapter_6.active_passive.r3 + chapter_9.active_passive_plan.r3) * 0.5, 2)
        self.S_sobstv_cap_year_mean_fact = round((chapter_6.active_passive.r3 + chapter_9.active_passive_fact.r3) * 0.5, 2)
        self.k_oborach_sobstv_capital_plan = chapter_8.Q_plan / self.S_sobstv_cap_year_mean_plan
        self.k_oborach_sobstv_capital_fact = chapter_8.Q_fact / self.S_sobstv_cap_year_mean_fact

        self.R_production_plan = chapter_8.P_pr_plan / chapter_4.S_sum.total
        self.R_production_fact = chapter_8.P_pr_fact / chapter_4.S_sum.total

        self.R_sell_plan = chapter_8.P_chistaya_plan / chapter_8.Q_plan
        self.R_sell_fact = chapter_8.P_chistaya_fact / chapter_8.Q_fact

        self.R_active_plan = chapter_8.P_chistaya_plan / chapter_9.active_passive_plan.active
        self.R_active_fact = chapter_8.P_chistaya_fact / chapter_9.active_passive_fact.active

        self.R_sobstv_capital_plan = chapter_8.P_chistaya_plan / self.S_sobstv_cap_year_mean_plan
        self.R_sobstv_capital_fact = chapter_8.P_chistaya_fact / self.S_sobstv_cap_year_mean_fact

        self.N_kr = Chapter_10.bin_search()
        self.Q_kr = self.N_kr * chapter_7.P_proizv_plan
        self.k_pokr = CalculateTable(chapter_4.N_pl_values, Chapter_10.calc_k_pokr)

        self.Q_fin_pr_plan = (chapter_8.Q_plan - self.Q_kr) / chapter_8.Q_plan
        self.Q_fin_pr_fact = (chapter_8.Q_fact - self.Q_kr) / chapter_8.Q_fact

        self.proizv_richag_plan = (chapter_8.Q_plan - chapter_4.S_sum.variable) / chapter_8.P_pr_plan
        self.proizv_richag_fact = (chapter_8.Q_fact - chapter_4.S_b_poln.variable * chapter_8.N_fact) / chapter_8.P_pr_fact


class Chapter_2_1:
    __slots__ = [
        'A_percent', 'B_percent', 'C_percent',
        'N_pl_A', 'N_pl_B', 'N_pl_C',
        'T_pl', 'B', 'C', 'D', 'O', 'H', 'gamma', 'F_ob_ef'
    ]

    def __init__(self, initial_data: InitialData):
        self.A_percent = 0.25
        self.B_percent = 0.5
        self.C_percent = 0.25

        self.N_pl_A = round(initial_data.N_pl * self.A_percent)
        self.N_pl_C = round(initial_data.N_pl * self.C_percent)
        self.N_pl_B = initial_data.N_pl - self.N_pl_A - self.N_pl_C

        self.T_pl = 365
        self.B = 118
        self.C = 2
        self.D = 8
        self.O = 20
        self.H = 20
        self.gamma = 0.05

        self.F_ob_ef = (self.T_pl - self.B) * self.C * self.D * (1 - self.gamma)


class Chapter_2_2:
    __slots__ = [
        'b_norm',
        'machines',
        'main_resources',
        'TO_perv',
        'S_os',
        'S_os_amortisable',
        'new_machines_cost',
    ]

    def __init__(self, initial_data: InitialData, chapter_1: Chapter_1, chapter_2_1: Chapter_2_1, chapter_3: Chapter_3):
        self.b_norm = 0.7
        F_ob_ef = chapter_2_1.F_ob_ef

        machines = {}

        for operations, N in [
            [initial_data.operations_A.rows, chapter_2_1.N_pl_A],
            [initial_data.operations_C.rows, chapter_2_1.N_pl_C],
            [initial_data.operations_B.rows, chapter_2_1.N_pl_B]]:
            for operation in operations:
                name = operation['name']
                if len(name) == 1:
                    e = N * operation['time'] / (self.b_norm * F_ob_ef)
                    if name in machines:
                        machines[name][0] += e
                    else:
                        machines[name] = [e, operation['cost']]

        self.machines = Table('name', 'stock', 'need_rasch', 'need_fact', 'need_new', 'b_fact', 'cost')

        for name, machine_and_cost in machines.items():
            stock_row = chapter_1.machines.find('name', name)
            stock = 0
            if stock_row is not None:
                stock = stock_row['n fact']

            machine = machine_and_cost[0]

            need_total = math.ceil(machine - 0.05)
            need_new = max(0, need_total - stock)
            ct = machine * (self.b_norm * F_ob_ef)
            try:
                b_fact = ct / (need_total * F_ob_ef)
            except:
                b_fact = 0
            self.machines.add_row(name, stock, machine, need_total, need_new, b_fact, machine_and_cost[1])

        self.new_machines_cost = self.machines.calculate_sum(lambda x: x['need_new'] * x['cost'])

        self.main_resources = Table('n', 'name', '%', 'cost I', 'amortisation I', 'cost II begin', 'delta', 'cost II')
        S_os_delta = round(self.new_machines_cost / 0.4 * 0.5, 2)

        self.S_os = 0
        self.S_os_amortisable = 0

        for old in chapter_1.main_resources.rows:
            if old['n'] != '1':
                amortisation = round(old['cost'] * chapter_3.OS_amortisation_percent, 2)
            else:
                amortisation = 0
            cost_2_begin = old['cost'] - amortisation
            if old['name'] == '- технологическое оборудование':
                need = self.new_machines_cost
            elif old['name'] == 'Машины и оборудование, в т.ч.':
                need = self.new_machines_cost + round(S_os_delta * chapter_1.main_resources.find('name', '- нетехнологические машины и оборудование')['%'], 2)
            else:
                need = S_os_delta * old['%']
            need = round(need, 2)
            cost_2 = cost_2_begin + need
            if len(old['n']) > 0:
                self.S_os += cost_2
                if old['n'] != '1':
                    self.S_os_amortisable += cost_2
            self.main_resources.add_row(old['n'], old['name'], old['%'], old['cost'], amortisation, cost_2_begin, need, cost_2)

        self.TO_perv = chapter_1.TO_perv + self.new_machines_cost


class Chapter_2_3:
    __slots__ = [
        'F_rab_ef', 'C_opr_mean',
        'total_time_A',
        'total_time_B',
        'total_time_C',
        'total_time',
        'opr_salary', 'opr_extra',
        'stimulating_salary_percent',
        'R_opr_raw', 'R_opr',
        'vpr', 'R_vpr',
        'sl', 'R_sl',
        'R_ppp',
        'FOT_opr', 'FOT_opr_extra',
        'FOT_vpr', 'FOT_sl',
        'FOT', 'FOT_fee', 'insurance_fee', 'FOT_with_fee',
    ]

    def __init__(self, chapter_2: Chapter_2, chapter_2_1: Chapter_2_1):
        self.F_rab_ef = (chapter_2_1.T_pl - chapter_2_1.B - chapter_2_1.O - chapter_2_1.H) * chapter_2_1.D

        self.total_time_A = initial_data.operations_A.calculate_sum(lambda x: x['time'])
        self.total_time_B = initial_data.operations_B.calculate_sum(lambda x: x['time'])
        self.total_time_C = initial_data.operations_C.calculate_sum(lambda x: x['time'])

        self.total_time = chapter_2_1.N_pl_A * self.total_time_A + chapter_2_1.N_pl_B * self.total_time_B + chapter_2_1.N_pl_C * self.total_time_C

        self.opr_salary = chapter_2.opr_salary
        self.opr_extra = chapter_2.opr_extra
        self.stimulating_salary_percent = chapter_2.stimulating_salary_percent

        self.R_opr_raw = self.total_time / self.F_rab_ef
        self.R_opr = int(math.ceil(self.R_opr_raw))

        self.vpr = chapter_2.vpr.clone(self.R_opr)

        self.sl = chapter_2.sl.clone(self.R_opr)
        self.R_vpr = self.vpr.total
        self.R_sl = self.sl.total

        self.C_opr_mean = 12 * self.opr_salary / self.F_rab_ef

        self.FOT_opr = self.C_opr_mean * self.total_time
        self.FOT_opr_extra = self.R_opr * (self.opr_extra * 12 + self.stimulating_salary_percent * self.opr_salary)

        self.FOT_vpr = self.vpr.calc_sum(lambda amount, salary: amount * salary * (12 + self.stimulating_salary_percent))
        self.FOT_sl = self.sl.calc_sum(lambda amount, salary: amount * salary * (12 + self.stimulating_salary_percent))

        self.R_ppp = self.R_opr + self.R_vpr + self.R_sl
        self.FOT = Value('fot', self.FOT_vpr + self.FOT_sl, self.FOT_opr + self.FOT_opr_extra, 'Затраты на оплату труда')

        self.insurance_fee = PerPercentTable(self.FOT.total)
        self.insurance_fee.add_row('ОПФ', 0.22)
        self.insurance_fee.add_row('ФОМС', 0.051)
        self.insurance_fee.add_row('ФСС', 0.029)
        self.insurance_fee.add_row('Страхование от несчастных случаев на производстве и профессиональных заболеваний', 0.04)

        self.FOT_fee = Value('fot fee', round(self.FOT.const * 0.34, 2), round(self.FOT.variable * 0.34, 2), 'Страховые взносы')
        self.FOT_with_fee = Value('fot total', display_name='ФОТ')
        self.FOT_with_fee.add_child(self.FOT)
        self.FOT_with_fee.add_child(self.FOT_fee)


class Chapter_2_4:
    __slots__ = [
        'S_materials_i_comp_A',
        'S_materials_i_comp_B',
        'S_materials_i_comp_C',

        'main_materials',

        'OS_amortisation',

        'NMA',
        'NMA_amortisation',

        'OS_fix',

        'costs',
        'S_pr_tek_pl',
        'S_kom',
        'S_sum'
    ]

    def __init__(self, initial_data: InitialData, chapter_3: Chapter_3, chapter_4: Chapter_4, chapter_2_1: Chapter_2_1, chapter_2_3: Chapter_2_3, const=None):
        self.S_materials_i_comp_B = chapter_3.S_mat_i_comp

        A_m = initial_data.materials_A.calculate_sum(lambda x: x['amount'] * x['cost'])
        A_a = initial_data.accessories_A.calculate_sum(lambda x: x['amount'] * x['cost'])
        self.S_materials_i_comp_A = A_m + A_a

        C_m = initial_data.materials_C.calculate_sum(lambda x: x['amount'] * x['cost'])
        C_a = initial_data.accessories_C.calculate_sum(lambda x: x['amount'] * x['cost'])
        self.S_materials_i_comp_C = C_m + C_a

        main_materials = chapter_2_1.N_pl_A * self.S_materials_i_comp_A + chapter_2_1.N_pl_B * self.S_materials_i_comp_B + chapter_2_1.N_pl_C * self.S_materials_i_comp_C

        self.costs = Value('S proizv', display_name='Затраты')
        self.main_materials = Value('material', variable=main_materials, display_name='Материальные затраты')
        m_base = self.main_materials.total

        self.main_materials.add_child(Value('material_main', 0, m_base, display_name='Основные материалы'))

        self.main_materials.add_child(Value('helper', 0, round(m_base * chapter_3.help_materials_percent, 2), display_name='Вспомогательные материалы'))

        ms = round(m_base * chapter_3.moving_save_percent, 2)
        msc = round(ms * chapter_3.move_save_const_percent, 2) if const is None or const['move save'] is None else const['move save'].const
        msv = round(ms - msc, 2) if const is None or const['move save'] is None else round((1 - chapter_3.move_save_const_percent) * ms, 2)
        self.main_materials.add_child(Value('move save', msc, msv, display_name='Транспортно-заготовительные расходы'))

        self.main_materials.add_child(Value('inventory', const=round(m_base * chapter_3.inventory_percent, 2), display_name='Инструменты, инвентарь')
                                      if const is None or const['inventory'] is None else const['inventory'])

        fuelt = round(m_base * chapter_3.fuel_percent, 2)

        fuel_energy_costs = self.main_materials.add_child(Value('fuel total', display_name='Топливо и энергия'))
        ft = fuel_energy_costs.add_child(Value('fuel tech', variable=round(fuelt * chapter_3.fuel_tech_percent, 2), display_name='Технологическое топливо и энергия'))
        fuel_energy_costs.add_child(Value('fuel non tech', const=fuelt - ft.total, display_name='Нетехнологическое топливо и энергия')
                                    if const is None or const['fuel non tech'] is None else const['fuel non tech'])
        self.costs.add_child(self.main_materials)

        self.costs.add_child(chapter_2_3.FOT)
        self.costs.add_child(chapter_2_3.FOT_fee)

        self.OS_amortisation = round(chapter_3.OS_amortisation_percent * chapter_2_2.S_os_amortisable, 2)

        self.NMA_amortisation = chapter_3.NMA_amortisation
        self.NMA = round(chapter_3.NMA - chapter_3.NMA_amortisation, 2)

        aos = Value('amortisation OS', const=self.OS_amortisation, display_name='Амортизация ОС')
        anma = Value('amortisation NMA', const=self.NMA_amortisation, display_name='Амортизация НМА')
        a = Value('amortisation', display_name='Амортизация ОС и НМА')
        a.add_child(aos)
        a.add_child(anma)
        self.costs.add_child(a)

        self.OS_fix = round(chapter_3.OS_fix_percent * chapter_2_2.S_os_amortisable, 2)

        ex = Value('extra', display_name='Прочие затраты')
        ex.add_child(Value('OS fix', const=self.OS_fix, display_name='Средства на ремонт ОС'))
        ex.add_child(Value('pure extra', self.costs.const * chapter_3.extra_percent, self.costs.variable * chapter_3.extra_percent))
        self.costs.add_child(ex)

        self.S_pr_tek_pl = self.costs.total

        s = round(self.S_pr_tek_pl * chapter_4.S_kom_percent // 10_000, 0) * 10_000
        sc = round(s * chapter_4.S_kom_const_percent, 2) if const is None or const['S_kom'] is None else const['S_kom'].const
        sv = round(s - sc, 2) if const is None or const['S_kom'] is None else s * (1 - chapter_4.S_kom_const_percent)

        self.S_sum = Value('S_sum', display_name='Суммарные затраты')
        self.S_sum.add_child(self.costs)
        self.S_kom = Value('S_kom', sc, sv, 'Коммерческие затраты')
        self.S_sum.add_child(self.S_kom)


class Chapter_2_5:
    __slots__ = [
        'direct_A',
        'direct_B',
        'direct_C',

        'direct_total',
        'total_machine_time',
        'indirect',

        'machine_time_A',
        'machine_time_B',
        'machine_time_C',
        'total_machine_time',
        'S_m_ch',

        'S_rab_ob_A',
        'S_rab_ob_B',
        'S_rab_ob_C',

        'S_sv_s_rab_ob_A',
        'S_sv_s_rab_ob_B',
        'S_sv_s_rab_ob_C',

        'S_ne_sv_s_rab_ob_A',
        'S_ne_sv_s_rab_ob_B',
        'S_ne_sv_s_rab_ob_C',

        'k_kosv',

        'S_A_proizv',
        'S_B_proizv',
        'S_C_proizv',
    ]

    def __init__(self, initial_data: InitialData, chapter_3: Chapter_3, chapter_2_1: Chapter_2_1, chapter_2_2: Chapter_2_2, chapter_2_3: Chapter_2_3, chapter_2_4: Chapter_2_4):
        percent_A = chapter_2_3.total_time_A / chapter_2_3.total_time
        percent_B = chapter_2_3.total_time_B / chapter_2_3.total_time
        percent_C = chapter_2_3.total_time_C / chapter_2_3.total_time

        OS_manufacture = chapter_2_2.TO_perv * chapter_3.OS_amortisation_percent
        OS_not_manufacture = (chapter_2_2.S_os_amortisable - chapter_2_2.TO_perv) * chapter_3.OS_amortisation_percent + chapter_2_4.NMA_amortisation

        self.direct_A = DirectCosts(percent_A, chapter_2_4.S_materials_i_comp_A, chapter_2_3.FOT.variable, chapter_2_3.FOT_fee.variable, OS_manufacture)
        self.direct_B = DirectCosts(percent_B, chapter_2_4.S_materials_i_comp_B, chapter_2_3.FOT.variable, chapter_2_3.FOT_fee.variable, OS_manufacture)
        self.direct_C = DirectCosts(percent_C, chapter_2_4.S_materials_i_comp_C, chapter_2_3.FOT.variable, chapter_2_3.FOT_fee.variable, OS_manufacture)

        self.direct_total = chapter_2_4.costs['material_main'].total + chapter_2_3.FOT.variable + chapter_2_3.FOT_fee.variable + OS_manufacture

        self.indirect = Table('n', 'name', 'cost')

        self.indirect.add_row('1', 'Вспомогательные материалы', WorkAndOther(chapter_2_4.costs['helper'].total * 0.6, chapter_2_4.costs['helper'].total * 0.4))
        self.indirect.add_row('2', 'Транспортно-заготовительные расходы', WorkAndOther(other=chapter_2_4.costs['move save'].total))
        self.indirect.add_row('3', 'Инструменты, инвентарь, хозяйственные принадлежности',
                              WorkAndOther(chapter_2_4.costs['inventory'].total * 0.8, chapter_2_4.costs['inventory'].total * 0.2))
        self.indirect.add_row('4', 'Топливо и энергия', WorkAndOther(chapter_2_4.costs['fuel tech'].total, chapter_2_4.costs['fuel non tech'].total))
        self.indirect.add_row('', '    на технологические цели', WorkAndOther(chapter_2_4.costs['fuel tech'].total))
        self.indirect.add_row('', '    на нетехнологические цели', WorkAndOther(other=chapter_2_4.costs['fuel non tech'].total))
        self.indirect.add_row('5', 'Заработная плата служащих и вспомогательных рабочих', WorkAndOther(other=chapter_2_3.FOT.const))
        self.indirect.add_row('6', 'Страховые взносы (на указанную заработную плату)', WorkAndOther(other=chapter_2_3.FOT_fee.const))
        self.indirect.add_row('7', 'Амортизация основных средств (не используемая при производстве продукции) и нематериальных активов', WorkAndOther(other=OS_not_manufacture))
        self.indirect.add_row('8', 'Затраты на ремонт оборудования', WorkAndOther(chapter_2_4.costs['OS fix'].total))
        self.indirect.add_row('9', 'Прочие расходы', WorkAndOther(other=chapter_2_4.costs['pure extra'].total))

        self.machine_time_A = initial_data.operations_A.filter_table(lambda x: len(x['name']) == 1).calculate_sum(lambda x: x['time'])
        self.machine_time_B = initial_data.operations_B.filter_table(lambda x: len(x['name']) == 1).calculate_sum(lambda x: x['time'])
        self.machine_time_C = initial_data.operations_C.filter_table(lambda x: len(x['name']) == 1).calculate_sum(lambda x: x['time'])

        self.total_machine_time = chapter_2_1.N_pl_A * self.machine_time_A + chapter_2_1.N_pl_B * self.machine_time_B + chapter_2_1.N_pl_C * self.machine_time_C

        total_indirect_work = self.indirect.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['cost'].work)
        total_indirect_other = self.indirect.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['cost'].other)
        self.S_m_ch = total_indirect_work / self.total_machine_time
        self.S_sv_s_rab_ob_A = self.machine_time_A * chapter_2_1.N_pl_A * self.S_m_ch
        self.S_sv_s_rab_ob_B = self.machine_time_B * chapter_2_1.N_pl_B * self.S_m_ch
        self.S_sv_s_rab_ob_C = self.machine_time_C * chapter_2_1.N_pl_C * self.S_m_ch

        self.k_kosv = total_indirect_other / (total_indirect_work + chapter_2_3.FOT.variable)

        self.S_ne_sv_s_rab_ob_A = (chapter_2_1.N_pl_A * self.direct_A.FOT + self.S_sv_s_rab_ob_A) * self.k_kosv
        self.S_ne_sv_s_rab_ob_B = (chapter_2_1.N_pl_B * self.direct_B.FOT + self.S_sv_s_rab_ob_B) * self.k_kosv
        self.S_ne_sv_s_rab_ob_C = (chapter_2_1.N_pl_C * self.direct_C.FOT + self.S_sv_s_rab_ob_C) * self.k_kosv

        self.S_rab_ob_A = WorkAndOther(self.machine_time_A * self.S_m_ch, self.S_ne_sv_s_rab_ob_A / max(1, chapter_2_1.N_pl_A))
        self.S_rab_ob_B = WorkAndOther(self.machine_time_B * self.S_m_ch, self.S_ne_sv_s_rab_ob_B / max(1, chapter_2_1.N_pl_B))
        self.S_rab_ob_C = WorkAndOther(self.machine_time_C * self.S_m_ch, self.S_ne_sv_s_rab_ob_C / max(1, chapter_2_1.N_pl_C))

        self.S_A_proizv = self.direct_A.direct + self.S_rab_ob_A.total
        self.S_B_proizv = self.direct_B.direct + self.S_rab_ob_B.total
        self.S_C_proizv = self.direct_C.direct + self.S_rab_ob_C.total


class Chapter_2_6:
    __slots__ = [
        'K_ob_sr_mk',
        'k_ob_sr_percent',
        'K_ob_sr_pr_zap',

        'k_nz_A',
        'k_nz_B',
        'k_nz_C',
        'gamma_cycle',
        'T_cycle_A',
        'T_cycle_B',
        'T_cycle_C',
        'K_ob_nez_pr_A', 'K_ob_nez_pr_B', 'K_ob_nez_pr_C',
        'K_ob_nez_pr',

        't_real',
        'K_ob_got_prod_A', 'K_ob_got_prod_B', 'K_ob_got_prod_C',
        'K_ob_got_prod',

        'gamma_ob',
        'K_ob_extra',
        'K_ob_RBP',
        'K_ob_ds',
        'K_ob_sum',
    ]

    def __init__(self, chapter_5: Chapter_5, chapter_2_1: Chapter_2_1, chapter_2_5: Chapter_2_5):
        mz = 0
        for i, n in [
            (initial_data.materials_A, chapter_2_1.N_pl_A),
            (initial_data.materials_B, chapter_2_1.N_pl_B),
            (initial_data.materials_C, chapter_2_1.N_pl_C)
        ]:
            mz += i.calculate_sum(lambda x: x['amount'] * x['cost'] * n / chapter_2_1.T_pl * x['t_zap'])

        cz = 0
        for i, n in [
            (initial_data.accessories_A, chapter_2_1.N_pl_A),
            (initial_data.accessories_B, chapter_2_1.N_pl_B),
            (initial_data.accessories_C, chapter_2_1.N_pl_C)
        ]:
            cz += i.calculate_sum(lambda x: x['amount'] * x['cost'] * n / chapter_2_1.T_pl * x['t_zap'])

        self.K_ob_sr_mk = round(mz + cz, 2)
        self.k_ob_sr_percent = chapter_5.k_ob_sr_percent
        self.K_ob_sr_pr_zap = round((1 + self.k_ob_sr_percent) * self.K_ob_sr_mk, 2)
        self.k_nz_A = (chapter_2_4.S_materials_i_comp_A + chapter_2_5.S_A_proizv) / (chapter_2_5.S_A_proizv * 2)
        self.k_nz_B = (chapter_2_4.S_materials_i_comp_B + chapter_2_5.S_B_proizv) / (chapter_2_5.S_B_proizv * 2)
        self.k_nz_C = (chapter_2_4.S_materials_i_comp_C + chapter_2_5.S_C_proizv) / (chapter_2_5.S_C_proizv * 2)

        self.gamma_cycle = chapter_5.gamma_cycle

        ml = self.gamma_cycle / (chapter_2_1.C * chapter_2_1.D) * chapter_2_1.T_pl / (chapter_2_1.T_pl - chapter_2_1.B)

        self.T_cycle_A = round(initial_data.operations_A.calculate_sum(lambda x: x['time']) * ml, 3)
        self.T_cycle_B = round(initial_data.operations_B.calculate_sum(lambda x: x['time']) * ml, 3)
        self.T_cycle_C = round(initial_data.operations_C.calculate_sum(lambda x: x['time']) * ml, 3)

        self.K_ob_nez_pr_A = chapter_2_5.S_A_proizv * chapter_2_1.N_pl_A / chapter_2_1.T_pl * self.k_nz_A * self.T_cycle_A
        self.K_ob_nez_pr_B = chapter_2_5.S_B_proizv * chapter_2_1.N_pl_B / chapter_2_1.T_pl * self.k_nz_B * self.T_cycle_B
        self.K_ob_nez_pr_C = chapter_2_5.S_C_proizv * chapter_2_1.N_pl_C / chapter_2_1.T_pl * self.k_nz_C * self.T_cycle_C
        self.K_ob_nez_pr = round(self.K_ob_nez_pr_A + self.K_ob_nez_pr_B + self.K_ob_nez_pr_C, 2)

        self.t_real = chapter_5.t_real
        self.K_ob_got_prod_A = chapter_2_5.S_A_proizv * chapter_2_1.N_pl_A / chapter_2_1.T_pl * self.t_real
        self.K_ob_got_prod_B = chapter_2_5.S_B_proizv * chapter_2_1.N_pl_B / chapter_2_1.T_pl * self.t_real
        self.K_ob_got_prod_C = chapter_2_5.S_C_proizv * chapter_2_1.N_pl_C / chapter_2_1.T_pl * self.t_real
        self.K_ob_got_prod = round(self.K_ob_got_prod_A + self.K_ob_got_prod_B + self.K_ob_got_prod_C, 2)

        self.gamma_ob = chapter_5.gamma_ob
        a = round(self.K_ob_sr_pr_zap + self.K_ob_nez_pr + self.K_ob_got_prod, 2)
        self.K_ob_sum = round(a / self.gamma_ob, 2)
        self.K_ob_extra = round(self.K_ob_sum - a, 2)

        self.K_ob_RBP = chapter_6.active_passive.K_ob_RBP
        self.K_ob_ds = self.K_ob_sum - (self.K_ob_sr_pr_zap + self.K_ob_RBP)


class Chapter_2_7:
    __slots__ = [
        'TS_real_ost',
        'P_real_ost',
        'S_nalogi',
        'P_chistaya_real_ost',
        'to_sell',
        'S_sell_OS',
        'adding',
        'subbing',

        'active_passive'
    ]

    def __init__(self, chapter_1: Chapter_1, chapter_3: Chapter_3, chapter_4: Chapter_4,
                 chapter_7: Chapter_7, chapter_8: Chapter_8, chapter_9: Chapter_9, chapter_2_2: Chapter_2_2, chapter_2_6: Chapter_2_6):
        self.TS_real_ost = round(chapter_4.S_b_poln.total + (chapter_7.P_fact - chapter_4.S_b_poln.total) * 0.7, 2)
        self.P_real_ost = chapter_8.N_ost * self.TS_real_ost
        self.S_nalogi = round(self.P_real_ost * chapter_7.tax, 2)
        self.P_chistaya_real_ost = self.P_real_ost - self.S_nalogi
        self.to_sell = Table('name', 'extra', 'cost')
        s = chapter_2_2.machines.filter_table(lambda x: x['stock'] > x['need_fact'])
        costs = 0
        for i in s.rows:
            r = initial_data.operations_B.find('name', i['name'])
            if r:
                costs += r['cost'] * (i['stock'] - i['need_fact'])
                self.to_sell.add_row(r['name'], i['stock'] - i['need_fact'], round(r['cost'] * (1 - chapter_3.OS_amortisation_percent), 2))
        self.S_sell_OS = round(costs * (1 - chapter_3.OS_amortisation_percent), 2)

        self.adding = self.P_chistaya_real_ost + self.S_sell_OS
        self.subbing = chapter_2_2.S_os - chapter_1.S_os + chapter_3.costs["amortisation OS"].total

        ap = self.active_passive = ExtendedActivePassive('end I', 'buy TO', 'buy OS', 'sell ost', 'sell TO', 'begin II')

        end_1 = ap['end I']
        end_1.set(chapter_9.active_passive_fact, chapter_9.active_passive_fact)

        buy_TO = ap['buy TO']
        buy_TO.set(end_1.active, end_1.passive)
        free_ds = buy_TO.active.K_ob_ds - 500_000
        need_ds = chapter_2_2.new_machines_cost
        trata = min(need_ds, free_ds)
        need_ds -= trata
        free_ds -= trata
        zaem = need_ds
        buy_TO.active.OS += chapter_2_2.new_machines_cost
        buy_TO.active.K_ob_ds = free_ds + 500_000
        buy_TO.passive.kratkosroch_zaem_sredstva += zaem

        buy_OS = ap['buy OS']
        buy_OS.set(buy_TO.active, buy_TO.passive)
        free_ds = buy_OS.active.K_ob_ds - 500_000
        need_ds = chapter_2_2.S_os - chapter_1.S_os + chapter_3.costs["amortisation OS"].total - chapter_2_2.new_machines_cost
        trata = min(need_ds, free_ds)
        need_ds -= trata
        free_ds -= trata
        zaem = need_ds
        buy_OS.active.OS += chapter_2_2.S_os - chapter_1.S_os + chapter_3.costs["amortisation OS"].total - chapter_2_2.new_machines_cost
        buy_OS.active.K_ob_ds = free_ds + 500_000
        buy_OS.passive.kratkosroch_zaem_sredstva += zaem

        sell_ost = ap['sell ost']
        sell_ost.set(buy_OS.active, buy_OS.passive)
        sell_ost.passive.neraspred_pribil += self.P_chistaya_real_ost
        sbpr_nost = chapter_4.S_b_proizv * chapter_8.N_ost
        sell_ost.active.K_ob_ds += sbpr_nost
        sell_ost.active.K_ob_got_prod -= sbpr_nost
        sell_ost.active.K_ob_ds += self.P_chistaya_real_ost
        free_ds = max(sell_ost.active.K_ob_ds - 500_000, 0)
        zaem_left = max(0, sell_ost.passive.kratkosroch_zaem_sredstva - free_ds)
        sell_ost.active.K_ob_ds -= max(0, sell_ost.passive.kratkosroch_zaem_sredstva - zaem_left)
        sell_ost.passive.kratkosroch_zaem_sredstva = zaem_left

        sell_TO = ap['sell TO']
        sell_TO.set(sell_ost.active, sell_ost.passive)
        sell_TO.active.K_ob_ds += self.S_sell_OS
        sell_TO.active.OS -= self.S_sell_OS
        free_ds = max(sell_TO.active.K_ob_ds - 500_000, 0)
        zaem_left = max(0, sell_TO.passive.kratkosroch_zaem_sredstva - free_ds)
        sell_TO.active.K_ob_ds -= max(0, sell_TO.passive.kratkosroch_zaem_sredstva - zaem_left)
        sell_TO.passive.kratkosroch_zaem_sredstva = zaem_left

        begin_2 = ap['begin II']
        begin_2.set(sell_TO.active, sell_TO.passive)
        begin_2.active.K_ob_ds += begin_2.active.K_ob_got_prod + begin_2.active.K_ob_nez_pr
        begin_2.active.K_ob_got_prod = 0
        begin_2.active.K_ob_nez_pr = 0
        delta_K_ob_extra = chapter_2_6.K_ob_extra - chapter_9.active_passive_fact.K_ob_extra
        delta_K_ob_sr_pr_zap = chapter_2_6.K_ob_sr_pr_zap - chapter_9.active_passive_fact.K_ob_sr_pr_zap
        begin_2.active.K_ob_sr_pr_zap = chapter_2_6.K_ob_sr_pr_zap
        begin_2.active.K_ob_extra = chapter_2_6.K_ob_extra
        total_delta = delta_K_ob_extra + delta_K_ob_sr_pr_zap
        if total_delta > 0:
            free_ds = begin_2.active.K_ob_ds - 500_000
            need_ds = total_delta
            trata = min(need_ds, free_ds)
            need_ds -= trata
            free_ds -= trata
            zaem = need_ds
            begin_2.active.K_ob_ds = free_ds + 500_000
            begin_2.passive.kratkosroch_zaem_sredstva += zaem
        else:
            begin_2.active.K_ob_ds -= total_delta


class Chapter_2_8:
    __slots__ = [
        'parametric_data',
        'k_param_price_A', 'k_param_price_C',
        'param_price_A', 'param_price_B', 'param_price_C',
        'P_prodaj',
        'k_nats',
        'total_cost_price_A', 'total_cost_price_B', 'total_cost_price_C',
        'k_nats_variable',
        'costs_A', 'costs_B', 'costs_C',
        'S_A_poln', 'S_B_poln', 'S_C_poln',
        'S_A_sum', 'S_B_sum', 'S_C_sum',
        'variable_cost_price_A', 'variable_cost_price_B', 'variable_cost_price_C',
        'TS_A_plan', 'TS_B_plan', 'TS_C_plan'
    ]

    @staticmethod
    def calc_costs(mic, helper, move_save, inventory, fuel_tech, fuel_non_tech, FOT, FOT_fee, amortisation, OS_fix, extra):
        costs = Value('S proizv', display_name='Затраты')
        main_materials = Value('material', display_name='Материальные затраты')

        main_materials.add_child(Value('material_main', 0, mic, display_name='Основные материалы'))
        main_materials.add_child(Value('helper', 0, helper, display_name='Вспомогательные материалы'))

        ms = move_save
        msc = ms * chapter_3.move_save_const_percent
        msv = (ms - msc)
        main_materials.add_child(Value('move save', msc, msv, display_name='Транспортно-заготовительные расходы'))

        main_materials.add_child(Value('inventory', const=inventory, display_name='Инструменты, инвентарь'))

        fuel_energy_costs = main_materials.add_child(Value('fuel total', display_name='Топливо и энергия'))
        fuel_energy_costs.add_child(Value('fuel tech', variable=fuel_tech, display_name='Технологическое топливо и энергия'))
        fuel_energy_costs.add_child(Value('fuel non tech', const=fuel_non_tech, display_name='Нетехнологическое топливо и энергия'))
        costs.add_child(main_materials)

        costs.add_child(FOT)
        costs.add_child(FOT_fee)

        costs.add_child(Value('amortisation', const=amortisation, display_name='Амортизация ОС и НМА'))

        ec = costs.const / costs.total

        ex = costs.add_child(Value('extra', display_name='Прочие затраты'))
        ex.add_child(Value('pure extra', const=extra * ec, variable=extra * (1 - ec)))
        ex.add_child(Value('OS fix', const=OS_fix))
        return costs

    @staticmethod
    def calc_for_product(N, S_materials_i_comp, direct, indirect, costs, w, o):
        A = Chapter_2_8.calc_costs(
            S_materials_i_comp * N,
            costs['helper'].total * (0.6 * w + 0.4 * o),
            costs['move save'].total * o,
            costs['inventory'].total * (0.8 * w + 0.2 * o),
            costs['fuel tech'].total * w,
            costs['fuel non tech'].total * o,
            Value('fot', chapter_2_3.FOT.const * o, direct.FOT * N),
            Value('fot fee', chapter_2_3.FOT_fee.const * o, direct.FOT_fee * N),
            direct.amortisation * N + indirect.find('n', '7')['cost'].total * o,
            costs['OS fix'].total * w,
            indirect.find('n', '9')['cost'].total * o
        )
        return A

    def __init__(self):
        pa = self.parametric_data = Table('name', 'A', 'B', 'C', 'importance A', 'importance C', 'power')
        pa.add_row('X1', 20, 20, 20, 0.1, 0.2, 0)
        pa.add_row('X2', 15, 12, 10, 0.5, 0.4, 1)
        pa.add_row('X3', 48, 60, 75, 0.4, 0.4, -1)

        pa2 = pa.filter_table(lambda x: x['name'] in ['X2', 'X3'])
        self.k_param_price_A = pa2.calculate_sum(lambda x: (x['A'] / x['B']) ** x['power'] * x['importance A'])
        self.k_param_price_C = pa2.calculate_sum(lambda x: (x['C'] / x['B']) ** x['power'] * x['importance C'])
        self.param_price_A = round(chapter_7.P_fact * self.k_param_price_A, 2)
        self.param_price_B = round(chapter_7.P_fact, 2)
        self.param_price_C = round(chapter_7.P_fact * self.k_param_price_C, 2)

        self.P_prodaj = chapter_2_4.S_sum.total * chapter_10.R_production_fact
        self.k_nats = self.P_prodaj / chapter_2_4.S_sum.total

        s_kom_unit = chapter_2_4.S_kom.total / initial_data.N_pl

        self.total_cost_price_A = round((chapter_2_5.S_A_proizv + s_kom_unit) * (1 + self.k_nats), 2)
        self.total_cost_price_B = round((chapter_2_5.S_B_proizv + s_kom_unit) * (1 + self.k_nats), 2)
        self.total_cost_price_C = round((chapter_2_5.S_C_proizv + s_kom_unit) * (1 + self.k_nats), 2)

        const_S_kom_unit = s_kom_unit * chapter_4.S_kom_const_percent
        variable_S_kom_unit = s_kom_unit * (1 - chapter_4.S_kom_const_percent)

        S_sum_a = Value('S sum A')
        S_sum_a.add_child(Value('S kom', const_S_kom_unit * chapter_2_1.N_pl_A, variable_S_kom_unit * chapter_2_1.N_pl_A))

        tmt = chapter_2_5.S_sv_s_rab_ob_A + chapter_2_5.S_sv_s_rab_ob_B + chapter_2_5.S_sv_s_rab_ob_C
        teb = chapter_2_5.S_ne_sv_s_rab_ob_A + chapter_2_5.S_ne_sv_s_rab_ob_B + chapter_2_5.S_ne_sv_s_rab_ob_C

        self.costs_A = Chapter_2_8.calc_for_product(chapter_2_1.N_pl_A, chapter_2_4.S_materials_i_comp_A, chapter_2_5.direct_A,
                                                    chapter_2_5.indirect, chapter_2_4.costs, chapter_2_5.S_sv_s_rab_ob_A / tmt, chapter_2_5.S_ne_sv_s_rab_ob_A / teb)
        self.costs_B = Chapter_2_8.calc_for_product(chapter_2_1.N_pl_B, chapter_2_4.S_materials_i_comp_B, chapter_2_5.direct_B,
                                                    chapter_2_5.indirect, chapter_2_4.costs, chapter_2_5.S_sv_s_rab_ob_B / tmt, chapter_2_5.S_ne_sv_s_rab_ob_B / teb)
        self.costs_C = Chapter_2_8.calc_for_product(chapter_2_1.N_pl_C, chapter_2_4.S_materials_i_comp_C, chapter_2_5.direct_C,
                                                    chapter_2_5.indirect, chapter_2_4.costs, chapter_2_5.S_sv_s_rab_ob_C / tmt, chapter_2_5.S_ne_sv_s_rab_ob_C / teb)

        s_kom_A = chapter_2_4.S_kom.total * chapter_2_1.A_percent / (max(1, chapter_2_1.N_pl_A))
        s_kom_B = chapter_2_4.S_kom.total * chapter_2_1.B_percent / (max(1, chapter_2_1.N_pl_B))
        s_kom_C = chapter_2_4.S_kom.total * chapter_2_1.C_percent / (max(1, chapter_2_1.N_pl_C))

        kp = chapter_4.S_kom_const_percent

        self.S_A_sum = Value('S A sum')
        self.S_A_sum.add_child(self.costs_A)
        self.S_A_sum.add_child(Value('S_kom', s_kom_A * chapter_2_1.N_pl_A * kp, s_kom_A * chapter_2_1.N_pl_A * (1 - kp), 'Коммерческие затраты'))

        self.S_B_sum = Value('S B sum')
        self.S_B_sum.add_child(self.costs_B)
        self.S_B_sum.add_child(Value('S_kom', s_kom_B * chapter_2_1.N_pl_B * kp, s_kom_B * chapter_2_1.N_pl_B * (1 - kp), 'Коммерческие затраты'))

        self.S_C_sum = Value('S C sum')
        self.S_C_sum.add_child(self.costs_C)
        self.S_C_sum.add_child(Value('S_kom', s_kom_C * chapter_2_1.N_pl_C * kp, s_kom_C * chapter_2_1.N_pl_C * (1 - kp), 'Коммерческие затраты'))

        self.S_A_poln = Value('S A poln')
        self.S_A_poln.add_child(Value('proizv', self.costs_A.const / chapter_2_1.N_pl_A, self.costs_A.variable / chapter_2_1.N_pl_A))
        self.S_A_poln.add_child(Value('S_kom', const=s_kom_A * kp, variable=s_kom_A * (1 - kp)))

        self.S_B_poln = Value('S B poln')
        self.S_B_poln.add_child(Value('proizv', self.costs_B.const / chapter_2_1.N_pl_B, self.costs_B.variable / chapter_2_1.N_pl_B))
        self.S_B_poln.add_child(Value('S_kom', const=s_kom_B * kp, variable=s_kom_B * (1 - kp)))

        self.S_C_poln = Value('S C poln')
        self.S_C_poln.add_child(Value('proizv', self.costs_C.const / chapter_2_1.N_pl_C, self.costs_C.variable / chapter_2_1.N_pl_C))
        self.S_C_poln.add_child(Value('S_kom', const=s_kom_C * kp, variable=s_kom_C * (1 - kp)))
        self.k_nats_variable = (self.P_prodaj + chapter_2_4.S_sum.const) / chapter_2_4.S_sum.variable

        self.variable_cost_price_A = round(self.S_A_poln.variable * (1 + self.k_nats_variable), 2)
        self.variable_cost_price_B = round(self.S_B_poln.variable * (1 + self.k_nats_variable), 2)
        self.variable_cost_price_C = round(self.S_C_poln.variable * (1 + self.k_nats_variable), 2)

        self.TS_A_plan = round(sum([self.param_price_A, self.total_cost_price_A, self.variable_cost_price_A]) / 3, 2)
        self.TS_B_plan = round(sum([self.param_price_B, self.total_cost_price_B, self.variable_cost_price_B]) / 3, 2)
        self.TS_C_plan = round(sum([self.param_price_C, self.total_cost_price_C, self.variable_cost_price_C]) / 3, 2)


class Chapter_2_9:
    __slots__ = [
        'Q_plan_A', 'Q_plan_B', 'Q_plan_C',
        'Q_plan',
        'K_ob_got_prod_plan',
        'S_pr_got_pr_plan',
        'S_valovaya_plan',
        'S_kom_plan',
        'P_pr_plan_A', 'P_pr_plan_B', 'P_pr_plan_C',
        'P_pr_plan',
        'P_pr_do_nalogov_plan',
        'S_prochie_dohidy_i_rashody_plan',
        'nalog_na_pribil_plan',
        'P_chistaya_plan'
    ]

    def __init__(self):
        self.Q_plan_A = chapter_2_8.TS_A_plan * chapter_2_1.N_pl_A
        self.Q_plan_B = chapter_2_8.TS_B_plan * chapter_2_1.N_pl_B
        self.Q_plan_C = chapter_2_8.TS_C_plan * chapter_2_1.N_pl_C
        self.Q_plan = self.Q_plan_A + self.Q_plan_B + self.Q_plan_C

        self.K_ob_got_prod_plan = chapter_2_6.K_ob_got_prod
        self.S_pr_got_pr_plan = chapter_2_4.S_pr_tek_pl - chapter_2_6.K_ob_nez_pr - self.K_ob_got_prod_plan

        self.S_valovaya_plan = self.Q_plan - self.S_pr_got_pr_plan

        self.S_kom_plan = chapter_2_4.S_kom.total

        self.P_pr_plan_A = (
                                   self.Q_plan_A - chapter_2_8.costs_A.total + chapter_2_6.K_ob_nez_pr_A + chapter_2_6.K_ob_got_prod_A) - chapter_2_4.S_kom.total * chapter_2_1.N_pl_A / initial_data.N_pl
        self.P_pr_plan_B = (
                                   self.Q_plan_B - chapter_2_8.costs_B.total + chapter_2_6.K_ob_nez_pr_B + chapter_2_6.K_ob_got_prod_B) - chapter_2_4.S_kom.total * chapter_2_1.N_pl_B / initial_data.N_pl
        self.P_pr_plan_C = (
                                   self.Q_plan_C - chapter_2_8.costs_C.total + chapter_2_6.K_ob_nez_pr_C + chapter_2_6.K_ob_got_prod_C) - chapter_2_4.S_kom.total * chapter_2_1.N_pl_C / initial_data.N_pl

        self.P_pr_plan = self.S_valovaya_plan - self.S_kom_plan

        self.S_prochie_dohidy_i_rashody_plan = chapter_8.S_prochie_dohidy_i_rashody_plan * 1.04
        self.P_pr_do_nalogov_plan = self.P_pr_plan - self.S_prochie_dohidy_i_rashody_plan
        self.nalog_na_pribil_plan = round(self.P_pr_do_nalogov_plan * chapter_7.tax, 2)

        self.P_chistaya_plan = self.P_pr_do_nalogov_plan - self.nalog_na_pribil_plan


class Chapter_2_10:
    __slots__ = [
        'K_den_sr_plan',
        'K_den_sr_konez_plan',
        'S_kratkosroch_zaem_sredstva_konez_plan',
        'active_passive_plan',
        'valid_to_cope_kz_plan',
    ]

    def __init__(self):
        amortisation = chapter_2_4.costs['amortisation'].total

        self.K_den_sr_plan = chapter_2_7.active_passive['begin II'].active.K_ob_ds + amortisation + chapter_2_9.P_chistaya_plan - (chapter_2_6.K_ob_nez_pr + chapter_2_9.K_ob_got_prod_plan)
        if self.K_den_sr_plan - 500_000 > chapter_2_7.active_passive['begin II'].passive.kratkosroch_zaem_sredstva:
            self.K_den_sr_konez_plan = self.K_den_sr_plan - chapter_2_7.active_passive['begin II'].passive.kratkosroch_zaem_sredstva
            self.S_kratkosroch_zaem_sredstva_konez_plan = 0
            self.valid_to_cope_kz_plan = 'full'
        elif self.K_den_sr_plan > 500_000:
            self.K_den_sr_konez_plan = 500_000
            self.S_kratkosroch_zaem_sredstva_konez_plan = chapter_2_7.active_passive['begin II'].passive.kratkosroch_zaem_sredstva - (self.K_den_sr_plan - 500_000)
            self.valid_to_cope_kz_plan = 'part'
        else:
            self.S_kratkosroch_zaem_sredstva_konez_plan = chapter_2_7.active_passive['begin II'].passive.kratkosroch_zaem_sredstva
            self.K_den_sr_konez_plan = self.K_den_sr_plan
            self.valid_to_cope_kz_plan = 'none'

        p = self.active_passive_plan = ActivePassive()

        p.NMA = chapter_2_7.active_passive['begin II'].active.NMA - chapter_2_4.costs['amortisation NMA'].total
        p.OS = chapter_2_2.S_os - chapter_2_4.costs['amortisation OS'].total

        p.K_ob_sr_pr_zap = chapter_2_6.K_ob_sr_pr_zap
        p.K_ob_nez_pr = chapter_2_6.K_ob_nez_pr
        p.K_ob_got_prod = chapter_2_6.K_ob_got_prod
        p.K_ob_RBP = chapter_2_7.active_passive['begin II'].active.K_ob_RBP
        p.K_ob_extra = chapter_2_7.active_passive['begin II'].active.K_ob_extra
        p.K_ob_ds = self.K_den_sr_konez_plan

        p.ustavnoy_kapital = chapter_2_7.active_passive['begin II'].passive.ustavnoy_kapital
        p.neraspred_pribil = chapter_2_7.active_passive['begin II'].passive.neraspred_pribil + chapter_2_9.P_chistaya_plan

        p.doldosroch_zaemn_sredstva = chapter_2_7.active_passive['begin II'].passive.doldosroch_zaemn_sredstva

        p.kratkosroch_zaem_sredstva = self.S_kratkosroch_zaem_sredstva_konez_plan
        p.kratkosroch_prochee = chapter_2_7.active_passive['begin II'].passive.kratkosroch_prochee


class Chapter_2_11:
    __slots__ = [
        'k_sob_ob_sr_plan',
        'k_obespech_sob_sr_plan',
        'k_abs_likvid_plan',
        'k_tek_likvid_plan',
        'V',
        'OS_year_mean',
        'k_FO_plan',
        'k_FE_plan',
        'Z_ob_sr_year_mean_plan',
        'Z_ob_plan',
        'S_sobstv_cap_year_mean_plan',
        'k_oborach_sobstv_capital_plan',
        'R_production_plan',
        'R_sell_plan',
        'R_active_plan',
        'R_sobstv_capital_plan',

        'N_kr_A', 'Q_kr_A', 'k_pokr_A', 'Q_fin_pr_A', 'proizv_richag_A',
        'N_kr_B', 'Q_kr_B', 'k_pokr_B', 'Q_fin_pr_B', 'proizv_richag_B',
        'N_kr_C', 'Q_kr_C', 'k_pokr_C', 'Q_fin_pr_C', 'proizv_richag_C',
    ]

    def __init__(self):
        self.k_sob_ob_sr_plan = chapter_2_10.active_passive_plan.r2 - chapter_2_10.active_passive_plan.r5

        self.k_obespech_sob_sr_plan = self.k_sob_ob_sr_plan / chapter_2_10.active_passive_plan.r2

        self.k_abs_likvid_plan = chapter_2_10.active_passive_plan.K_ob_ds / chapter_2_10.active_passive_plan.r5

        self.k_tek_likvid_plan = chapter_2_10.active_passive_plan.r2 / chapter_2_10.active_passive_plan.r5

        self.V = initial_data.N_pl / chapter_2_3.R_ppp

        self.OS_year_mean = round(chapter_2_2.S_os_amortisable - chapter_2_4.costs['amortisation OS'].total * 0.5, 2)
        self.k_FO_plan = chapter_2_9.Q_plan / self.OS_year_mean

        self.k_FE_plan = 1 / self.k_FO_plan

        self.Z_ob_sr_year_mean_plan = round((chapter_2_7.active_passive['begin II'].active.r2 + chapter_2_10.active_passive_plan.r2) * 0.5, 2)
        self.Z_ob_plan = chapter_2_9.Q_plan / self.Z_ob_sr_year_mean_plan

        self.S_sobstv_cap_year_mean_plan = round((chapter_2_7.active_passive['begin II'].passive.r3 + chapter_2_10.active_passive_plan.r3) * 0.5, 2)
        self.k_oborach_sobstv_capital_plan = chapter_2_9.Q_plan / self.S_sobstv_cap_year_mean_plan

        self.R_production_plan = chapter_2_9.P_pr_plan / chapter_2_4.S_sum.total
        self.R_sell_plan = chapter_2_9.P_chistaya_plan / chapter_2_9.Q_plan
        self.R_active_plan = chapter_2_9.P_chistaya_plan / chapter_2_10.active_passive_plan.active
        self.R_sobstv_capital_plan = chapter_2_9.P_chistaya_plan / self.S_sobstv_cap_year_mean_plan

        self.N_kr_A = math.ceil(chapter_2_8.S_A_sum.const / (chapter_2_8.TS_A_plan - chapter_2_8.S_A_poln.variable))
        self.Q_kr_A = self.N_kr_A * chapter_2_8.TS_A_plan
        self.k_pokr_A = (chapter_2_8.TS_A_plan - chapter_2_8.S_A_poln.variable) / chapter_2_8.TS_A_plan
        self.Q_fin_pr_A = (chapter_2_9.Q_plan_A - self.Q_kr_A) / chapter_2_9.Q_plan_A
        self.proizv_richag_A = (chapter_2_9.Q_plan_A - chapter_2_8.S_A_sum.variable) / chapter_2_9.P_pr_plan_A

        self.N_kr_B = math.ceil(chapter_2_8.S_B_sum.const / (chapter_2_8.TS_B_plan - chapter_2_8.S_B_poln.variable))
        self.Q_kr_B = self.N_kr_B * chapter_2_8.TS_B_plan
        self.k_pokr_B = (chapter_2_8.TS_B_plan - chapter_2_8.S_B_poln.variable) / chapter_2_8.TS_B_plan
        self.Q_fin_pr_B = (chapter_2_9.Q_plan_B - self.Q_kr_B) / chapter_2_9.Q_plan_B
        self.proizv_richag_B = (chapter_2_9.Q_plan_B - chapter_2_8.S_B_sum.variable) / chapter_2_9.P_pr_plan_B

        self.N_kr_C = math.ceil(chapter_2_8.S_C_sum.const / (chapter_2_8.TS_C_plan - chapter_2_8.S_C_poln.variable))
        self.Q_kr_C = self.N_kr_C * chapter_2_8.TS_C_plan
        self.k_pokr_C = (chapter_2_8.TS_C_plan - chapter_2_8.S_C_poln.variable) / chapter_2_8.TS_C_plan
        self.Q_fin_pr_C = (chapter_2_9.Q_plan_C - self.Q_kr_C) / chapter_2_9.Q_plan_C
        self.proizv_richag_C = (chapter_2_9.Q_plan_C - chapter_2_8.S_C_sum.variable) / chapter_2_9.P_pr_plan_C


initial_data = InitialData()
chapter_1 = Chapter_1(initial_data)
chapter_2 = Chapter_2(initial_data, chapter_1)
chapter_3 = Chapter_3(initial_data, chapter_1, chapter_2)
chapter_4 = Chapter_4(initial_data, chapter_3)
chapter_5 = Chapter_5(initial_data, chapter_1, chapter_3, chapter_4)
chapter_6 = Chapter_6(chapter_1, chapter_5)
chapter_7 = Chapter_7(chapter_4, chapter_6)
chapter_8 = Chapter_8(initial_data, chapter_3, chapter_4, chapter_5, chapter_7)
chapter_9 = Chapter_9()
chapter_10 = Chapter_10(initial_data, chapter_1, chapter_2, chapter_3, chapter_4, chapter_6, chapter_7, chapter_8, chapter_9)

chapter_2_1 = Chapter_2_1(initial_data)
chapter_2_2 = Chapter_2_2(initial_data, chapter_1, chapter_2_1, chapter_3)
chapter_2_3 = Chapter_2_3(chapter_2, chapter_2_1)
chapter_2_4 = Chapter_2_4(initial_data, chapter_3, chapter_4, chapter_2_1, chapter_2_3)
chapter_2_5 = Chapter_2_5(initial_data, chapter_3, chapter_2_1, chapter_2_2, chapter_2_3, chapter_2_4)
chapter_2_6 = Chapter_2_6(chapter_5, chapter_2_1, chapter_2_5)
chapter_2_7 = Chapter_2_7(chapter_1, chapter_3, chapter_4, chapter_7, chapter_8, chapter_9, chapter_2_2, chapter_2_6)
chapter_2_8 = Chapter_2_8()
chapter_2_9 = Chapter_2_9()
chapter_2_10 = Chapter_2_10()
chapter_2_11 = Chapter_2_11()


def gen_first_list():
    dp('todo').add_run().add_break(WD_BREAK.PAGE)
    document.add_section()

    dp('Содержание')
    p = document.add_paragraph()
    document.add_page_break()

    return p


def gen_introduction():
    document.add_paragraph('Введение', style=title_text)

    paragraphs = [
        'Рассматривается деятельность условного предприятия на протяжение двух периодов хозяйственной деятельности.',

        'В первом разделе условное предприятие начинает свою деятельность с производства лишь одного изделия Б. '
        'Рыночные условия формируются исполнителем путем задания фактических объемов продаж и цен, отличных от плановых значений.',

        'В первом периоде на базе исходных данных производится расчёт основных и оборотных средств; численности персонала и фонда оплаты труда. '
        'Происходит формирование сметы затрат на производство, после чего рассчитывается полная себестоимость продукции. Потом составляется '
        'баланс хозяйственных средств предприятия на начало периода. Рассчитываются плановая цена и, имеющая место на реальном рынке, фактическая '
        'цена. Составляется отчет о прибылях и убытках по плановым и фактическим данным и плановый и фактический баланса хозяйственных средств на '
        'конец периода. По завершении года рассчитываются показатели хозяйственной деятельности и делаются выводы об эффективности работы предприятия.',

        'Во втором разделе предполагается, что анализ рыночного спроса показал, что часть потребителей не удовлетворена качеством выпускаемого изделия '
        'Б. Причем одна группа потребителей готова приобретать аналог А более высокого качества даже по более высокой цене; другая группа потребителей '
        'готова приобретать изделие-аналог В более низкого качества и за более умеренную цену. При этом основной задачей остаётся определить, '
        'в каком количестве производить различные виды продукции.',

        'Во втором разделе также производятся плановые расчёты, аналогичные первому пункту, производится расчёт показателей хозяйственной деятельности '
        'и итогам анализа делаются выводы об эффективности деятельности предприятия и прогнозируется его дальнейшая деятельность.'
    ]

    lr = None
    for text in paragraphs:
        lr = document.add_paragraph(text, style=main_text).runs[0]

    lr.add_break(WD_BREAK.PAGE)


def gen_initial_data():
    def build_consumable_table(name, data):
        table = add_table(
            [[name, 'Стоимость, руб./ед.изм', 'Норма расхода, шт.']] +
            [[str(i + 1), str(e['cost']), str(e['amount'])] for i, e in enumerate(data.rows)],
            [Cm(5), Cm(6), Cm(5)], True)
        table_last_row = table.add_row()
        table_last_row.cells[0].paragraphs[0].add_run('Итого, руб.').bold = True
        table_last_row.cells[1].merge(table_last_row.cells[2])
        table_last_row.cells[1].paragraphs[0].add_run(str(data.calculate_sum(lambda x: x['cost'] * x['amount']))).bold = True
        table_last_row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    dp('Исходные данные 623', title_text)

    dp('Таблица 1, объём производства изделия Б', table_name_text)
    add_table([
        ['Объём производства издения Б, тыс.шт./год', str(initial_data.N_pl // 1000)]
    ], [Cm(11), Cm(5)])
    dp()

    dp('Таблица 2, потребность изделия Б в материалах', table_name_text)
    build_consumable_table('Вид материала', initial_data.materials)
    dp()

    dp('Таблица 3, потребность изделия Б в комплектующих', table_name_text)
    build_consumable_table('Вид комплектующих изделий', initial_data.accessories)
    dp()

    dp('Таблица 4, технологическая трудоёмкость изделия Б', table_name_text)
    table_4 = add_table(
        [['Номер технологической операции', 'Используемое оборудование', 'Первоначальная стоимость, тыс.руб./ед', 'Технологическая трудоёмкость, час./шт.']] +
        [[str(i + 1), e['name'], str(e['cost'] // 1000), str(e['time'])] for i, e in enumerate(initial_data.operations.rows)], first_bold=True)
    table_4_lr = table_4.add_row()
    table_4_lr.cells[0].merge(table_4_lr.cells[2])
    table_4_lr.cells[0].paragraphs[0].add_run('Итого, час.').bold = True
    table_4_lr.cells[3].paragraphs[0].add_run(str(initial_data.operations.calculate_sum(lambda x: x['time']))).bold = True
    table_4_lr.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()


def gen_1_1():
    chapter = chapter_1
    dp('1. Расчет потребности в основных средствах', title_text)
    dp('1.1. Расчет потребности в технологическом оборудовании', subtitle_text)
    dp('Найдём эффективный фонд времени работы оборудования')
    add_formula_with_description('F_{об.эф.} = (A - B) \\cdot C \\cdot D \\cdot (1-γ)', [
        ['A', f'число календарных дней в рассматриваемом периоде = {chapter.T_pl}'],
        ['B', f'число выходных и праздничных дней = {chapter.B}'],
        ['C', f'число смен в сутки = {chapter.C}'],
        ['D', f'число часов в одной смене = {chapter.D}'],
        ['γ', f'планируемые простои оборудования в долях единицы = {chapter.gamma}'],
    ])

    add_formula(
        f'F_{{об.эф.}} = ({chapter.T_pl} - {chapter.B}) \\cdot {chapter.C} \\cdot '
        f'{chapter.D} \\cdot (1-{chapter.gamma}) = {fn(chapter.F_ob_ef)} час./период')
    dp()

    dp('После этого найдём само число оборудования', main_text)
    add_formula_with_description('n_{об.k_{расч.}} = \\frac{N_{пл} t_j}{\\beta_{норм} F_{об.эф.}}', [
        ['t_j', 'время обработки изделия или услуги на i-том оборудовании, час/шт'],
        ['N_{пл}', 'планируемый объем производства в рассматриваемом периоде, шт./период'],
        ['\\beta', 'коэффициент загрузки оборудования, 0.75'],
        ['F_{об.эф.}', 'эффективный фонд времени работы оборудования'],
    ])

    dp('Таблица 1.1, Расчет потребности в технологическом оборудовании', table_name_text)
    add_table(
        [['Операция', 'Расчётное число оборудования', 'Принятое число оборудования', 'Фактический коэффициент нагрузки оборудования']] +
        [[e['name'], fn(e['n rasch']), fn(e['n fact'], 0), fn(e['b_fact'])] for e in chapter_1.machines.rows],
        [Cm(4), Cm(4), Cm(4), Cm(4)], True)
    dp()
    dp(
        'Следует отметить, что можно было взять 3 единицы оборудования для операции «в», однако в таком случае '
        'фактический коэффициент нагрузки возрос бы до 0.83, что негативно бы сказалось на сроке службы оборудования. '
        'Поэтому заложим большую надёжность и возьмём 4 единицы.')

    dp('Найдём суммарную первоначальную стоимость технологического оборудования, [тыс. руб.]')

    sm = []
    for e in chapter_1.machines.rows:
        sm.append(f"{e['cost'] // 1000} \\cdot {e['n fact']}")

    add_formula('ТО_{перв} = \\sum^{m}_{i}{ТО_{перв i} n_{об i_{прин}}} = ' + ' + '.join(sm) + f' = {fn(chapter.TO_perv // 1000)} [тыс.руб.]', style=formula_style_12)

    dp('1.2 Стоимостная структура основных средств', title_text)
    dp('Таблица 1.2, стоимостная структура основных средств', table_name_text)
    table = add_table(
        [['№', 'Название', '%', 'Стоимость, руб.']] + [[str(e['n']), e['name'], fn(e['%'] * 100, 0), fn(e['cost'], 0)] for e in chapter_1.main_resources.rows],
        [Cm(1), Cm(8), Cm(1.5), Cm(4)], True
    )
    tr = table.add_row()
    tr.cells[0].merge(tr.cells[1])
    tr.cells[0].paragraphs[0].add_run('Итого').bold = True
    tr.cells[2].paragraphs[0].add_run('100').bold = True
    tr.cells[3].paragraphs[0].add_run(fn(chapter_1.S_os, 0)).bold = True
    dp()
    document.add_page_break()


def gen_1_2():
    dp('2. Расчет численности и фонда з/п персонала предприятия', title_text)
    dp('2.1 Расчет численности ППП', subtitle_text)
    dp('Условно примем что непромышленный персонал отсутствует. Тогда численность ППП складывается из '
       'двух категорий – рабочих и служащих. Рабочие подразделяются на ОПР и ВПР.')
    dp('Чтобы найти численность ОПР найдём в начале эффективный фонд времени одного работающего:')
    add_formula_with_description(
        'F_{раб_{эф}}=(T_{пл}-B-O-H) \\cdot D = ' +
        f"({chapter_1.T_pl}-{chapter_1.B}-{chapter_1.O}-{chapter_1.H}) \\cdot {chapter_1.D} = {chapter_2.F_rab_ef}\\ [час/год\\ чел.]",
        [
            ['O', f'продолжительность отпуска = {chapter_1.H} раб. дн.'],
            ['H', f'число планируемых невыходов = {chapter_1.O} раб. дн.'],
        ],
        style=formula_style_12)
    dp('Примем величину планируемых невыходов за 20 (производство изделий на токарных станках негативно сказывается на здоровье за счёт мелкодисперсной '
       'металлической стружки, переменных магнитных полей электродвигателей и так далее), отпуск 20 рабочих дней (28 – 8 нерабочих дней за 4 недели).')

    dp('Подставим полученное значение в формулу численности ОПР:')

    add_formula_with_description(
        'R_{ОПР}=\\frac{N_{пл} \\sum_{i}^{m}{t_{техн i}}}{F_{раб_{эф}} k_{вн}} = \\frac{' +
        f'{initial_data.N_pl} \\cdot ({"+".join([str(e["time"]) for e in initial_data.operations.rows])})}}{{ {chapter_2.F_rab_ef} \\cdot 1 }} = {chapter_2.R_opr} [чел.]',
        [['t_{техн i}', 'трудоёмкость i-той операции']], style=formula_style_12
    )

    dp()
    p = dp(f'Численность ВРП примем за {fn(chapter_2.R_vpr / chapter_2.R_opr)} R')
    p.add_run('ОПР').font.subscript = True
    p.add_run(', а численность служащих за 0.6R')
    p.add_run('ОПР').font.subscript = True
    p.add_run(':')

    add_formula('R_{ВПР} = ' + fn(chapter_2.R_vpr / chapter_2.R_opr) + 'R_{ОПР} = ' + f'{chapter_2.R_vpr}')
    add_formula('R_{СЛ} = ' + fn(chapter_2.R_sl / chapter_2.R_opr) + 'R_{ОПР} = ' + f'{chapter_2.R_sl}')

    dp('Численность ППП:')
    p = add_formula('R_{ППП} = R_{ОПР} + R_{ВПР} + R_{СЛ} = ' + f'{chapter_2.R_opr} + {chapter_2.R_vpr} + {chapter_2.R_sl} = {chapter_2.R_ppp}')
    p.runs[0].add_break(WD_BREAK.PAGE)

    dp('2.2 Формирование штатного расписания', subtitle_text)
    dp(f'Рабочие ручной операции и весь обслуживающий персонал технологических операций «б», «в», «г» работают по сдельно-премиальной системе. '
       f'При выполнении ими плана, этим работникам положена надбавка в размере {fn(chapter_2.opr_extra, 0)} рублей в месяц.')
    dp('Остальной персонал работает по повременной оплате труда.')
    dp(f'Все работники получают стимулирующие выплаты раз в год в размере {fn(chapter_2.stimulating_salary_percent * 100)}% заработной платы.')

    dp(f'Средняя тарифная ставка ОПР: {fn(chapter_2.opr_salary)} [руб./чел.мес.]')
    add_formula('C_{тар.ст.} =\\frac{12 L_{ОПР ст}}{F_{раб_{эф}}} = \\frac{12 * ' + f'{fn(chapter_2.opr_salary, 0)} }}{{ {chapter_2.F_rab_ef} }} = {fn(chapter_2.C_opr_mean)}')
    add_formula('P_{ср} = C_{тар.ст.} \\cdot t_{ср} = ' + f'{fn(chapter_2.C_opr_mean)}' +
                '\\cdot \\frac{' + '+'.join([str(e['time']) for e in initial_data.operations.rows]) + '}{' + str(len(initial_data.operations)) + f'}} = {fn(chapter_2.p_mean)}')

    document.add_page_break()

    dp('Таблица 2.2.1, состав, структура и заработная плата персонала', table_name_text)
    add_employee_structure_table(chapter_2)

    document.add_page_break()

    dp('Таблица 2.2.2, суммарные заработные платы персонала за год', table_name_text)
    add_employee_salary_table(chapter_2)

    dp()
    dp(f'Отметим, что ОПР получают сдельную зарплату за {fn(chapter_2.R_opr_raw)} чел., но на предприятии их трудится {chapter_2.R_opr}, поэтому каждый из них получит меньше, '
       f'но их суммарная з/п будет равна ФОТ ОПР.').add_run().add_break(WD_BREAK.PAGE)

    dp('2.3 Расчет фонда оплаты труда ППП, величины страховых взносов', subtitle_text)
    dp('Найдём ФОТ ОПР на год:')
    add_formula('ФОТ_{ОПР без надб.} = p_{ср} N_{пл} m = ' + f'{fn(chapter_2.p_mean)} \\cdot {initial_data.N_pl} \\cdot {len(initial_data.operations)} = {fn(chapter_2.FOT_opr)} [руб./год]')

    dp('Учтём надбавки ОПР:')
    add_formula(
        'ФОТ_{ОПР} = ФОТ_{ОПР} + R_{ОПР} \\cdot (ОПР_{надб.} \\cdot 12 + ОПР_{тариф.ст.}) = ' +
        f'{chapter_2.R_opr} \\cdot ({fn(chapter_2.opr_extra)} \\cdot 12 + {fn(chapter_2.opr_salary)} = {fn(chapter_2.FOT_opr + chapter_2.FOT_opr_extra)} [руб./год]')
    dp('ФОТ ВПР и служащих:')
    add_formula_with_description('ФОТ_{ВПР+сл} = \\sum_{i}^{n}{(ТС_i \\cdot N_i \\cdot 12 + ТС_i)} = ' + f'{fn(chapter_2.FOT_vpr + chapter_2.FOT_sl)}', [
        ['n', 'число ОПР и служащих'],
        ['ТС_i', 'тарифная ставка'],
        ['N_i', 'численнойсть']
    ])

    dp('Общий ФОТ:')
    add_formula('ФОТ_{общ} = ФОТ_{ОПР} + ФОТ_{ВПР+сл} = ' + f'{fn(chapter_2.FOT.total)}')

    dp('Таблица 2.3, страховые взносы', table_name_text)
    table = add_table([['Взнос', 'Величина, %', 'Сумма, руб./год']] + [[e.name, fn(e.percent * 100, 1), fn(e.amount)] for e in chapter_2.insurance_fee.rows])
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    r.cells[0].paragraphs[0].add_run('Итого').bold = True
    r.cells[2].paragraphs[0].add_run(fn(chapter_2.FOT_fee.total)).bold = True
    document.add_page_break()


def gen_1_3():
    dp('3. Формирование сметы затрат на производство', title_text)
    dp('Стоимость основных материалов на единицу изделия:')
    add_formula('S_{ом.ед} = (S_м + S_k) = ' + f'{fn(chapter_3.S_mat_i_comp)} руб./шт.')
    dp('1. Стоимость основных материалов')
    add_formula(
        'S_{ом} = S_{ом.ед} \\cdot N_{пл} = ' + f'{fn(chapter_3.S_mat_i_comp)} \\cdot {fn(initial_data.N_pl, 0)} = {fn(chapter_3.costs["material_main"].total)}  [руб./год]')

    dp(f'2. Стоимость вспомогательных материалов (принято за {fn(chapter_3.help_materials_percent * 100)}% от *)')
    add_formula('S_{вм} = S_{ом} \\cdot k_{вм} = ' + f'{fn(chapter_3.costs["helper"].total)}  [руб./год]')

    dp(f'3. Транспортно-заготовительные расходы (принято за {fn(chapter_3.moving_save_percent * 100)}% от *)')
    add_formula('S_{т-з} = S_{ом} \\cdot k_{т-з} = ' + f'{fn(chapter_3.costs["move save"].total)}  [руб./год]')
    dp(f'Из них {fn(chapter_3.move_save_const_percent * 100, 0)}% - постоянные затраты {fn(chapter_3.costs["move save"].const)} [руб./год]')

    dp(f'4. Инструменты, инвентарь (принято за {fn(chapter_3.inventory_percent * 100)}% от *)')
    add_formula('S_{инстр} = S_{ом} \\cdot k_{интср} = ' + f'{fn(chapter_3.costs["inventory"].total)}  [руб./год]')

    dp(f'5. Топливо и энергия (принято за {fn(chapter_3.fuel_percent * 100)}% от *)')
    add_formula('S_{топл +эн} = S_{ом} \\cdot k_{топл+эн} = ' + f'{fn(chapter_3.costs["fuel total"].total)}  [руб./год]')

    dp(f'5.1 Технологическое ({fn(chapter_3.fuel_tech_percent * 100)}% от топлива и энергии)')
    add_formula('S_{тех. топл + эн} = S_{топл + эн} \\cdot k_{техн} = ' + f'{fn(chapter_3.costs["fuel tech"].total)}  [руб./год]')
    dp(f'5.2 Нетехнологическое ({fn(chapter_3.fuel_non_tech_percent * 100)}% от топлива и энергии)')
    add_formula('S_{тех. топл + эн} = S_{топл + эн} \\cdot (1 - k_{тех.}) = ' + f'{fn(chapter_3.costs["fuel non tech"].total)}  [руб./год]')

    dp('* - от стоимости основных материалов и комплектующих')

    dp(f'Норму амортизации основных средств примем за {fn(chapter_3.OS_amortisation_percent * 100)}%:')
    add_formula('A_{ос} = k_{ам.ос.} \\cdot (S_{осн.ср.} - S_{земл.}) = ' + f'{fn(chapter_3.OS_amortisation)} [руб./год]')

    dp(f'НМА = {fn(chapter_3.NMA)} руб.')
    dp(f'Норма амортизации НМА = {fn(chapter_3.NMA_amortisation_percent * 100)}%:')
    add_formula('A_{НМА} = k_{ам. НМА} \\cdot НМА = ' + f'{fn(chapter_3.NMA_amortisation)} [руб./год]')

    dp(f'Планируемые расходы на ремонт основных средств примем за {fn(chapter_3.OS_fix_percent * 100)}% от стоимости ОС:')
    add_formula('S_{рем.ос} = k_{рем ос} \\cdot (S_{осн.ср.} - S_{земл.}) = ' + f'{fn(chapter_3.OS_fix)} [руб./год]')

    dp('1. Материальные затраты:')
    add_formula('S_{мат.зат} = S_{ом} + S_{вм} + S_{т-з} S_{инстр} S_{топл.+эн.} = ' + f'{fn(chapter_3.costs["material"].total)} [руб./год]')

    dp('2. Затраты на оплату труда:')
    add_formula('S_{ФОТ} = ' + f'{fn(chapter_2.FOT.total)} [руб./год]')

    dp('3. Страховые взносы:')
    add_formula('S_{страх.вз} = ' + f'{fn(chapter_2.FOT_fee.total)} [руб./год]')

    dp('4. Амортизация основных средств и нематериальных активов:')
    add_formula('A_{ОС+НМА} = A_{ос} + A_{НМА}' + f'{fn(chapter_3.costs["amortisation"].total)} [руб./год]')

    dp(f'5. Прочие затраты (примем за {fn(chapter_3.extra_percent * 100)}% от первых 4 пунктов + планируемые расходы на ремонт ОС):')
    add_formula('S_{проч.зат.} = k_{проч.зат.} \\cdot (S_{мат.зат.} + S_{ФОТ} + S_{страх.вз} + A_{ОС+НМА}) + S_{рем.ос} = ' + f'{fn(chapter_3.costs["extra"].total)} [руб./год]')

    costs = [
        chapter_3.costs['material'],
        chapter_3.costs['fot'],
        chapter_3.costs['fot fee'],
        chapter_3.costs['amortisation'],
        chapter_3.costs['extra'],
    ]

    dp('Таблица 3, смета затрат', table_name_text)
    table = add_table(
        [['№', 'Элемент сметы', 'Сумма, руб/год', '%']] +
        [[str(i + 1), e._display_name, fn(e.total), fn(e.total / chapter_3.costs.total * 100)] for i, e in enumerate(costs)],
        [Cm(1), Cm(8), Cm(4), Cm(2)]
    )
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    r.cells[0].paragraphs[0].add_run('Итого: ').bold = True
    r.cells[0].paragraphs[0].add_run('затраты на производство в текущем периоде ')
    add_formula('S_{пр.тек._{пл.}} ', r.cells[0].paragraphs[0])
    r.cells[2].paragraphs[0].add_run(fn(chapter_3.costs.total)).bold = True
    r.cells[3].paragraphs[0].add_run('100').bold = True
    document.add_page_break()


def gen_1_4():
    dp('4. Расчет себестоимости изделия Б, величины условно-постоянных и переменных затрат', title_text)
    dp('Рассчитаем производственную и полную себестоимости изделия Б')

    dp('4.1 Расчет производственной себестоимости изделия Б', subtitle_text)
    dp('Производственную себестоимость определим по формуле:')
    add_formula('S_{Б\\ произв} = \\frac{S_{пр.тек.пл.}}{N_{пл}} = \\frac{' +
                f'{fn(chapter_3.S_pr_tek_pl)} }}{{ {fn(initial_data.N_pl)} }} = {fn(chapter_4.S_b_proizv)}\\ [руб./шт.]')

    dp('4.2 Расчет полной себестоимости изделия Б', subtitle_text)
    dp('Полную себестоимость определим по формуле:')
    add_formula(
        'S_{Б\\ полн} = \\frac{S_{пр.тек.пл.} + S_{ком}}{N_{пл}} = \\frac{S_{Б сум}}{N_{пл}} = \\frac{' +
        f'{fn(chapter_4.S_sum.total)} }}{{ {fn(initial_data.N_pl)} }} = {fn(chapter_4.S_b_poln.total)}\\ [руб./шт.]')
    dp('Коммерческие затраты связаны с реализацией продукции, включают расходы на тару и упаковку изделий на складах готовой продукции; расходы по '
       'доставке продукции на станцию отправления; комиссионные сборы (отчисления), уплачиваемые сбытовым и другим посредническим предприятиям; '
       'расходы по содержанию помещений для хранения продукции в местах ее реализации; рекламные и представительские расходы)')
    dp(f'Примем коммерческие затраты равными {fn(chapter_4.S_kom_percent * 100)}% от величины затрат на производство в текущем периоде:')
    add_formula('S_{ком} = k_{ком} \\cdot S_{пр.тек.пл.} = ' + f'{fn(chapter_4.S_kom.total)} [руб./год]')
    dp(f'Из них {fn(chapter_4.S_kom_const_percent * 100)}% составляют постоянные затраты, {fn(chapter_4.S_kom.const)}\\ [руб./год]').add_run().add_break(WD_BREAK.PAGE)

    dp('4.3. Построение графических зависимостей (условно-постоянных и переменных затрат)', subtitle_text)

    dp('Таблица 4.3.1, условно-постоянные и переменные затраты', table_name_text)
    table_wdt = [Cm(1), Cm(4), Cm(3.5), Cm(1), Cm(4), Cm(3.5)]
    table = add_table([
        [None, None, None, None, None, None],
        ['№', 'Условно-постоянные затраты', 'Сумма, тыс.руб./год', '№', 'Переменные затраты', 'Сумма, тыс.руб./год']
    ], table_wdt, style=table_style_12)
    table.cell(0, 0).merge(table.cell(0, 5)).text = f'Суммарные затраты, руб./год: {fn(chapter_4.S_sum.total)}'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(0, 0).paragraphs[0].runs[0].bold = True

    const = [
        (chapter_4.S_kom._display_name, chapter_4.S_kom.const),
        (chapter_3.costs['fuel non tech']._display_name, chapter_3.costs['fuel non tech'].const),
        ('З/п кроме ОПР', chapter_2.FOT.const),
        ('Страховые взносы кроме ОПР', chapter_2.FOT_fee.const),
        (chapter_3.costs['amortisation']._display_name, chapter_3.costs['amortisation'].const),
        (chapter_3.costs['inventory']._display_name, chapter_3.costs['inventory'].const),
        (chapter_3.costs['move save']._display_name, chapter_3.costs['move save'].const),
        (chapter_3.costs['extra']._display_name, chapter_3.costs['extra'].const)
    ]

    variable = [
        (chapter_4.S_kom._display_name, chapter_4.S_kom.variable),
        (chapter_3.costs['fuel tech']._display_name, chapter_3.costs['fuel tech'].variable),
        ('З/п ОПР', chapter_2.FOT.variable),
        ('Страховые взносы ОПР', chapter_2.FOT_fee.variable),
        (chapter_3.costs['material_main']._display_name, chapter_3.costs['material_main'].variable),
        (chapter_3.costs['helper']._display_name, chapter_3.costs['helper'].variable),
        (chapter_3.costs['move save']._display_name, chapter_3.costs['move save'].variable),
        (chapter_3.costs['extra']._display_name, chapter_3.costs['extra'].variable)
    ]

    for i, z in enumerate(zip(const, variable)):
        c = z[0]
        v = z[1]
        r = table.add_row()
        r.cells[0].paragraphs[0].add_run(str(i + 1))
        r.cells[1].paragraphs[0].add_run(c[0])
        r.cells[2].paragraphs[0].add_run(fn(c[1]))
        r.cells[3].paragraphs[0].add_run(str(i + 1))
        r.cells[4].paragraphs[0].add_run(v[0])
        r.cells[5].paragraphs[0].add_run(fn(v[1]))
        for j, e in enumerate(table_wdt):
            r.cells[j].width = e

    r = table.add_row()
    r.cells[0].merge(r.cells[1]).paragraphs[0].add_run('Итого').bold = True
    r.cells[2].paragraphs[0].add_run(fn(sum([e[1] for e in const])))
    r.cells[3].merge(r.cells[4]).paragraphs[0].add_run('Итого').bold = True
    r.cells[5].paragraphs[0].add_run(fn(sum([e[1] for e in variable])))

    add_formula('S_{Б\\ сум} = S_{пост} + S_{перем} = ' + f'{fn(chapter_4.S_sum.total)}')

    ct1 = chapter_4.ct1

    S_const_costs = [e.const for e in ct1.output_data]
    S_variable_costs = [e.variable for e in ct1.output_data]
    S_total_costs = [e.total for e in ct1.output_data]

    B_const_costs = [e.const / i for i, e in ct1.items]
    B_variable_costs = [e.variable / i for i, e in ct1.items]
    B_total_costs = [e.total / i for i, e in ct1.items]

    N_pl_values = chapter_4.N_pl_values
    dp()
    dp('Таблица 4.3.2 зависимость общей суммы затрат на производство и реализацию продукции от величины объема производства за планируемый период', table_name_text)
    add_table([
        ['N, шт./год'] + [fn(e, 0) for e in N_pl_values],
        ['S у-п, руб./год'] + [fn(e) for e in S_const_costs],
        ['S перем, руб./год'] + [fn(e) for e in S_variable_costs],
        ['S сум, руб./год'] + [fn(e) for e in S_total_costs],
        ], style=table_style_10)
    document.add_page_break()

    dp('Таблица 4.3.3, зависимость себестоимости единицы продукции от величины объема производства за планируемый период', table_name_text)
    add_table([
        ['N, шт./год'] + [fn(e, 0) for e in N_pl_values],
        ['B у-п, руб./год'] + [fn(e) for e in B_const_costs],
        ['B перем, руб./год'] + [fn(e) for e in B_variable_costs],
        ['B сум, руб./год'] + [fn(e) for e in B_total_costs],
        ], style=table_style_10)

    dp('График 4.3.4, зависимости общей суммы и себестоимости единицы продукции от величины объёма производства за планируемый период', table_name_text)
    plt.figure(figsize=(10, 4))
    plt.subplot(1, 2, 1)

    plt.title('S(N)')
    plt.xlabel('N, штук')
    plt.ylabel('S, млн. руб./год')
    plt.xticks(N_pl_values, [str(i) for i in N_pl_values], rotation=45)
    plt.yticks(np.linspace(min(S_const_costs + S_variable_costs + S_total_costs) / 1e6, max(S_const_costs + S_variable_costs + S_total_costs) / 1e6, 7))
    plt.grid(True)
    plt.plot(N_pl_values, np.array(S_const_costs) / 1e6, label='S_const')
    plt.plot(N_pl_values, np.array(S_variable_costs) / 1e6, label='S_variable')
    plt.plot(N_pl_values, np.array(S_total_costs) / 1e6, label='S_total')
    plt.legend()

    ax = plt.subplot(1, 2, 2)
    plt.title('B(N)')
    plt.xlabel('N, штук')
    plt.ylabel('B, руб./шт.')
    plt.yscale('log')
    # plt.yticks([0, max(B_const_costs)])
    ticker = matplotlib.ticker.ScalarFormatter()
    ax.yaxis.set_major_formatter(ticker)
    plt.yticks(np.logspace(np.log10(min(B_const_costs + B_variable_costs + B_total_costs)), np.log10(max(B_const_costs + B_variable_costs + B_total_costs)), 7))
    plt.xticks(N_pl_values, [str(i) for i in N_pl_values], rotation=45)
    plt.grid(True)
    plt.plot(N_pl_values, np.array(B_const_costs), label='B_const')
    plt.plot(N_pl_values, np.array(B_variable_costs), label='B_variable')
    plt.plot(N_pl_values, np.array(B_total_costs), label='B_total')
    plt.legend()
    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile)

    picP = document.add_paragraph()
    picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picP.add_run().add_picture(memfile, width=Cm(17))
    picP.add_run().add_break(WD_BREAK.PAGE)


def gen_1_5():
    dp('5. Расчет потребности в оборотных средствах', title_text)
    dp('Оборотные средства участвуют в одном производственном цикле и полностью переносят свою стоимость на готовую продукцию, и '
       'состоят из оборотных производственных фондов и фондов обращения. Сумма оборотных производственных фондов и фондов обращения '
       'представляют собой оборотные средства предприятия')
    dp()

    dp('5.1 Оборотные средства в производственных запасах', subtitle_text)

    dp('Таблица 5.1.1, норма запаса материалов', table_name_text)
    add_table([['Материал', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.materials.rows])

    dp()

    dp('Таблица 5.1.2 норма запаса комплектующих', table_name_text)
    add_table([['Комплектующие', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.accessories.rows])

    document.add_page_break()

    dp('Оборотные средства в производственных запасах материалов и комплектующих изделий определим по формуле:')
    add_formula_with_description(
        'K_{об.ср.\\ мат\\ и\\ комп.} = \\sum^{n_{мат}}_{i=1}{\\frac{M_{мат\\ i}N_{пл}}{Т_{пл}}}t_{зап.\\ мат.i} + '
        '\\sum^{n_{комп}}_{j=1}{\\frac{К_{комп\\ j}N_{пл}}{Т_{пл}}}t_{зап.\\ комп.j}', [
            ['М_{мат\\ i},\\ К_{комп\\ j}', 'норма расхода i-го материала и j-го вида комплектующих изделий на одно изготавливаемое изделие в стоимостном выражении, руб./шт.'],
            ['t_{зап.\\ мат.i},\\ t_{зап.\\ комп.j}', 'нормы запаса материалов и комплектующих изделий в календарных днях']
        ]
    )

    add_formula('K_{об.ср.\\ мат\\ и\\ комп.} = ' + f'{fn(chapter_5.K_ob_sr_mk)}\\ [руб.]')

    dp()

    dp(f'Рассчитанную величину оборотных средств в запасах материалов и комплектующих изделий увеличим на {fn(chapter_5.k_ob_sr_percent * 100, 0)}%. '
       f'Общая сумма оборотных средств в производственных запасах составит:')
    add_formula('K_{об.ср._{пр.зап.}} = ' + f'{fn(1 + chapter_5.k_ob_sr_percent)} \\cdot K_{{ об.ср.\\ мат.и\\ комп}} = {fn(chapter_5.K_ob_sr_pr_zap)}\\ [руб.]')
    dp()

    dp('5.2. Оборотные средства в незавершенном производстве', subtitle_text)

    dp('Затраты на материалы и комплектующие изделия:')
    add_formula('S_{мат\\ и\\ комп} =' + f'{fn(chapter_3.S_mat_i_comp)}\\ [руб.]')

    dp('Коэффициент нарастания затрат (условно принимается равномерное нарастание затрат)')
    add_formula(
        'k_{нз} = \\frac{S_{мат\\ и\\ комп} + S_{Б_{произв}}}{2 S_{Б_{произв}}} = \\frac{' +
        f'{fn(chapter_3.S_mat_i_comp)} + {fn(chapter_4.S_b_proizv)} }}{{ 2 \\cdot {fn(chapter_4.S_b_proizv)} }} = {fn(chapter_5.k_nz)}')

    document.add_page_break()

    dp('Производственный цикл (кален. дни) – отрезок времени между началом и окончанием производственного процесса '
       'изготовления одного изделия, включающий время технологических операций; время подготовительно-заключительных операций; '
       'длительность естественных процессов и вспомогательных операций; время межоперационных и междусменных перерывов; '
       'время ожидания обработки при передаче изделий на рабочие места по партиям')

    add_formula_with_description('T_ц = \\frac{\\sum^{m}_{i=1}{t_{техн\\ i}}\\gamma_ц}{C D}\\frac{T_{пл}}{Т_{пл} - B}', [
        ['\\gamma_ц', f'соотношение между производственным циклом и суммарной технологической трудоемкостью изготовления изделия, принято за {fn(chapter_5.gamma_cycle, 0)}']
    ])

    dp('Оборотные средства, находящиеся в незавершенном производстве, руб. рассчитаем по формуле:')
    add_formula('K_{об._{нез.пр.}} = \\frac{S_{Б_{произв}} N_{пл}}{Т_{пл}}k_{нз}T_ц')
    add_formula(f'T_ц = {fn(chapter_5.T_cycle, 3)}\\ [дн.]')
    add_formula(
        'K_{об._{нез.пр.}} = \\frac{' +
        f'{fn(chapter_4.S_b_proizv)} \\cdot {fn(initial_data.N_pl, 0)} }}{{ {chapter_1.T_pl} }} '
        f'\\cdot {fn(chapter_5.k_nz, 2)} \\cdot {fn(chapter_5.T_cycle, 3)} = {fn(chapter_5.K_ob_nez_pr)}\\ [руб.]')

    dp()
    dp('5.3. Оборотные средства в готовой продукции', subtitle_text)
    dp('Время нахождения на складе:')
    add_formula('t_{реал.} = ' + f'{chapter_5.t_real}\\ [дн.]')
    dp('Оборотные средства, находящиеся в готовой продукции:')
    add_formula(
        'K_{об._{гот.прод.}} = \\frac{S_{Б_{произв}} N_{пл}}{Т_{пл}}t_{реал} = \\frac{' +
        f'{fn(chapter_4.S_b_proizv)} \\cdot {fn(initial_data.N_pl)} }}{{ {chapter_1.T_pl} }} \\cdot {chapter_5.t_real} = {fn(chapter_5.K_ob_got_prod)}\\ [руб.]')

    document.add_page_break()

    dp('5.4. Суммарная потребность в оборотных средствах', subtitle_text)
    dp('Оборотные средства включают в себе не только оборотные средства в производственных запасах, незавершенном производстве и готовой продукции, '
       'а также в расходах будущих периодов, дебиторской задолженности, краткосрочных финансовых вложениях, денежных средствах и т.п. (т.е. прочие оборотные средства).')
    dp('Для упрощения расчетов в курсовой работе зададим удельный вес стоимости производственных запасов, незавершенного производства и готовой продукции в '
       'общей сумме оборотных средств:')

    add_formula('\\gamma_{об} = ' + f'{chapter_5.gamma_ob}')
    dp('Суммарные оборотные средства:')
    add_formula('K_{об_{сум}} = \\frac{K_{об.ср._{пр.зап.}} + K_{об._{нез.пр.}} + K_{об._{гот.прод.}}}{\\gamma_{об}}')
    add_formula('K_{об_{сум}} = \\frac{' + f'{fn(chapter_5.K_ob_sr_pr_zap)} + {fn(chapter_5.K_ob_nez_pr)} + {fn(chapter_5.K_ob_got_prod)} }}'
                                           f'{{ {chapter_5.gamma_ob} }} = {fn(chapter_5.K_ob_sum)}\\ [руб.]')
    dp('Прочие оборотные средства:')
    add_formula('K_{об_{проч}} = (1 - \\gamma_{об}) \\cdot K_{об_{сум}} = ' + f'{fn(chapter_5.K_ob_extra)}\\ [руб.]')

    document.add_page_break()


def gen_1_6():
    dp('6. Бухгалтерский баланс на начало деятельности условного предприятия', title_text)
    dp('Условно в работе вступительный баланс (составляемый на момент возникновения предприятия) совпадает с текущим балансом, составляемым на начало отчетного периода.')
    dp('На начало хозяйственной деятельности в бухгалтерском балансе отсутствуют: затраты в незавершенном производстве; готовая продукция и товары для перепродажи; '
       'дебиторская задолженность. Поэтому их значения на начало деятельности принять равными нулю.')
    dp('Величина сырья и материалов – исходя из рассчитанного норматива оборотных средств в производственных запасов.')
    dp('Значение прочих запасов и затрат определяется как разница между нормативом оборотных средств в производственных запасах с учетом прочих элементов '
       'производственных запасов + РБП, и нормативом оборотных средств в производственных запасов')

    dp('Денежные средства определяются:')
    add_formula(
        'K_{об_{ДС}} = K_{об_{сум}} - (K_{об.ср._{пр.зап.}}) + K_{об._{РБП}}) = ' +
        f'{fn(chapter_5.K_ob_sum)} - ({fn(chapter_5.K_ob_sr_pr_zap)} + {fn(chapter_6.active_passive.K_ob_RBP)}) = {fn(chapter_6.active_passive.K_ob_ds)}\\ [руб.]',
        style=formula_style_12)

    dp()

    dp(f'Удельный вес уставного капитала принят за {fn(chapter_6.ustavnoy_capital_percent * 100)}% от общей суммы пассива баланса.')
    dp('Распределение заёмного капитала:').paragraph_format.first_line_indent = 0
    dp(f'Долгосрочные заёмные средства: {fn(chapter_6.doldosroch_zaemn_sredstva_percent * 100)}%')
    dp(f'Краткосрочные заёмные средства: {fn(chapter_6.kratkosroch_zaemn_sredstva_percent * 100)}%')
    g = 1 - chapter_6.kratkosroch_zaemn_sredstva_percent - chapter_6.doldosroch_zaemn_sredstva_percent
    dp(f'Прочие краткосрочные обязательства: {fn(g * 100)}%').add_run().add_break(WD_BREAK.PAGE)

    dp('Таблица 6, бухгалтерский баланс на начало деятельности условного предприятия', table_name_text)
    add_active_passive_table(chapter_6.active_passive)

    document.add_page_break()


def gen_1_7():
    dp('7. Планирование цены изделия Б, определение фактической ', title_text)
    dp('Расчет планируемой цены изделия Б произведем с помощью методов полных и переменных затрат.')
    dp()

    dp('7. 1 Расчет плановой цены изделия Б методом полных затрат', subtitle_text)
    dp('Расчет осуществим по формуле:')
    add_formula_with_description('Ц_{Б_{полн.(пл)}} = S_{Б\\ полн.} (1 + \\frac{П_{продаж}}{S_{Б\\ сум.}}) = S_{Б\\ полн.}(1 + k_{нац})', [
        ['П_{продаж}', 'прибыль от продаж, условно примем равной прибыли до налогообложения, руб./год'],
        ['k_{нац}', 'коэффициент наценки']
    ])

    dp()
    dp(f'Прибыль до налогообложения определим через удельный вес чистой прибыли в общей сумме прибыли до налогообложения {fn(1 - chapter_7.tax)}')

    dp(f'Условно примем желаемую чистую прибыль за {fn(chapter_7.net_profit_percent * 100)}% от собственного капитала')

    add_formula('П_{чист} = S_{соб.капитал} \\cdot ' + f'{fn(chapter_7.net_profit_percent)} = {fn(chapter_7.net_profit)}\\ [тыс.руб/год]')

    dp('Прибыль от реализации = прибыль до налогообложения:')

    add_formula('П_{реал} = \\frac{П_{чист}}{' + f'{fn(1 - chapter_7.tax)} }} = {fn(chapter_7.profit_before_tax)}\\ [тыс.руб/год]')
    add_formula('k_{нац} = \\frac{' + f'{fn(chapter_7.profit_before_tax)} }}{{ {fn(chapter_4.S_sum.total)} }} = {fn(chapter_7.k_nats)}')
    add_formula(
        'Ц_{Б_{полн.(пл)}} = ' +
        f'{fn(chapter_4.S_b_poln.total)} \\cdot (1 + {fn(chapter_7.k_nats)}) = {fn(chapter_7.P_b_poln)}\\ [руб.]')
    document.add_page_break()

    dp('7.2 Расчет плановой цены изделия Б методом переменных затрат', subtitle_text)
    dp('Расчет осуществим по формуле:')
    add_formula('Ц_{Б_{перем.(пл)}} = S_{Б\\ перем.} (1 + \\frac{П_{продаж} + S_{усл.-пост.}}{S_{перем.}})')
    add_formula(
        'Ц_{Б_{перем.(пл)}} = ' +
        f'{fn(chapter_4.S_b_poln.variable)} (1 + \\frac{{ {fn(chapter_7.profit_before_tax)} + '
        f'{fn(chapter_4.S_sum.const)} }}{{ {fn(chapter_4.S_sum.variable)} }}) = {fn(chapter_7.P_b_perem)}\\ [руб.]')

    dp('Цены должны быть одинаковыми, если не округлять до копеек дельные затраты и коэффициент наценки. При вычислении использованы более точные цифры, '
       'чем те, которые указаны в уравнении, поэтому разница минимальна, или отстутсвует вовсе.')
    dp()

    dp(f'В качестве плановой цены возьмём большую из двух: {fn(chapter_7.P_proizv_plan)} [руб.]')
    dp()

    dp('7.3. Расчет рыночной (фактической) цены изделия Б', subtitle_text)
    dp('Реальные условия сложились таким образом, что изделие Б смогли реализовать по цене ниже, чем запланировали:')
    add_formula(
        'Ц_{Б\\ факт} = Ц_{Б\\ произв\\ пл} \\cdot k_{Ц\\ факт} = ' +
        f' {fn(chapter_7.P_proizv_plan)} \\cdot {fn(chapter_7.price_fact_percent)} = {fn(chapter_7.P_fact)}\\ [руб.]')

    document.add_page_break()


def gen_1_8():
    dp('8. Отчет о финансовых результатах на конец первого периода', title_text)
    dp(f'Примем что реализовать удалось только {fn(chapter_8.N_fact_percent * 100)}% продукции:')
    add_formula('N_{факт} = ' + f'{fn(chapter_8.N_fact_percent)} \\cdot N_{{пл}} = {chapter_8.N_fact}\\ [шт.]')

    dp('Определим выручку:')
    add_formula('Q_{план} = Ц_{Б\\ произв\\ пл} \\cdot N_{пл} = ' + f'{fn(chapter_8.P_pr_plan)} \\cdot {fn(initial_data.N_pl)} = {fn(chapter_8.Q_plan)}')
    add_formula('Q_{факт} = Ц_{Б\\ произв\\ факт} \\cdot N_{факт} = ' + f'{fn(chapter_8.P_pr_fact)} \\cdot {fn(chapter_8.N_fact)} = {fn(chapter_8.Q_fact)}')
    dp()

    dp('Определим себестоимость проданной готовой продукции:')
    add_formula(
        'S_{пр.гот.пр_{план}} = S_{пр.тек_{пл.}} - K_{об_{нез.пр.}} - K_{об._{гот.прод.}} = ' +
        f'{fn(chapter_3.S_pr_tek_pl)} - {fn(chapter_5.K_ob_nez_pr)} - {fn(chapter_5.K_ob_got_prod)} = {fn(chapter_8.S_pr_got_pr_plan)}', style=formula_style_12)
    add_formula(
        'S_{пр.гот.пр_{факт}} = S_{пр.тек_{пл.}} - K_{об_{нез.пр.}} - K_{об._{гот.прод._{факт}}} = '
        'S_{пр.тек_{пл.}} - K_{об_{нез.пр.}} - S_{Б_{произв}} N_{ост.} - K_{об._{гот.прод.}} = ' +
        f'{fn(chapter_3.S_pr_tek_pl)} - {fn(chapter_5.K_ob_nez_pr)} - ({fn(chapter_4.S_b_poln.total)}'
        f' \\cdot {fn(chapter_8.N_ost, 0)}) = {fn(chapter_8.S_pr_got_pr_fact)}', style=formula_style_12)

    dp()
    dp('Определим валовую прибыль:')
    add_formula(
        'S_{валовая\\ план} = Q_{план} - S_{пр.гот.пр_{план}} = ' +
        f'{fn(chapter_8.Q_plan)} - {fn(chapter_8.S_pr_got_pr_plan)} = {fn(chapter_8.S_valovaya_plan)}', style=formula_style_12)
    add_formula(
        'S_{валовая\\ факт} = Q_{факт} - S_{пр.гот.пр_{факт}} = ' +
        f'{fn(chapter_8.Q_fact)} - {fn(chapter_8.S_pr_got_pr_fact)} = {fn(chapter_8.S_valovaya_fact)}', style=formula_style_12)

    dp()
    dp(f'Фактические коммерческие расходы определим как {fn(chapter_8.kom_percent * 100, 0)}% от планируемых')
    add_formula('K_{ком\\ план} = ' + f'{fn(chapter_8.S_kom_plan)}\\ (см\\ п.\\ 4.2)')
    add_formula('K_{ком\\ факт} = K_{ком\\ план} \\cdot ' + f'{fn(chapter_8.kom_percent)} = {fn(chapter_8.S_kom_fact)}')

    document.add_page_break()
    dp('Определим прибыль от продаж:')
    add_formula('П_{пр_{план}} = S_{валовая\\ план} - K_{ком.пл.} = ' +
                f'{fn(chapter_8.S_valovaya_plan)} - {fn(chapter_8.S_kom_plan)} = {fn(chapter_8.P_pr_plan)}', style=formula_style_12)
    add_formula('П_{пр_{факт}} = S_{валовая\\ факт} - K_{ком.факт.} = ' +
                f'{fn(chapter_8.S_valovaya_fact)} - {fn(chapter_8.S_kom_fact)} = {fn(chapter_8.P_pr_fact)}', style=formula_style_12)

    dp('Определим прочие расходы (прочие доходы примем равными нулю):')
    dp(f'Для этого условно вычтем из прибыли от продаж прибыль до налогообложения (см. п. 7.1), '
       f'фактические прочие расходы определим как {fn(chapter_8.pr_dir_fact_percent * 100, 0)}% от плановых.')
    add_formula('S_{прочие\\ план.} = П_{пр_{план}} - П_{продаж.} = ' +
                f'{fn(chapter_8.P_pr_plan)} - {fn(chapter_8.P_pr_do_nalogov_plan)} = {fn(chapter_8.S_prochie_dohidy_i_rashody_plan)}', style=formula_style_12)
    add_formula('S_{прочие\\ факт.} = S_{прочие план.} \\cdot ' +
                f'{fn(chapter_8.pr_dir_fact_percent)}  = {fn(chapter_8.S_prochie_dohidy_i_rashody_fact)}', style=formula_style_12)

    dp()
    dp('Определим прибыль до налогообложения:')
    add_formula(
        'П_{до\\ налогообл.\\ план.} = П_{пр_{план}} - S_{прочие\\ план.} = ' +
        f'{fn(chapter_8.P_pr_plan)} - {fn(chapter_8.S_prochie_dohidy_i_rashody_plan)} = {fn(chapter_8.P_pr_do_nalogov_plan)}', style=formula_style_12)
    add_formula(
        'П_{до\\ налогообл.\\ факт.} = П_{пр_{факт}} - S_{прочие\\ факт.} = ' +
        f'{fn(chapter_8.P_pr_fact)} - {fn(chapter_8.S_prochie_dohidy_i_rashody_fact)} = {fn(chapter_8.P_pr_do_nalogov_fact)}', style=formula_style_12)

    dp()
    dp(f'Определим налог на прибыль (для ООО по общей системе налогообложения составляет {fn(chapter_7.tax * 100)}%):')
    add_formula('S_{налог\\ план.} = П_{до\\ налогообл.\\ план.} \\cdot ' + f'{fn(chapter_7.tax)} = {fn(chapter_8.nalog_na_pribil_plan)}')
    add_formula('S_{налог\\ факт.} = П_{до\\ налогообл.\\ факт.} \\cdot ' + f'{fn(chapter_7.tax)} = {fn(chapter_8.nalog_na_pribil_fact)}')

    dp()
    dp('Определим чистую прибыль:')
    add_formula(
        'П_{чист.\\ план.} = П_{до\\ налогообл.\\ план.} - S_{налог\\ план.} = ' +
        f'{fn(chapter_8.P_pr_do_nalogov_plan)} - {fn(chapter_8.nalog_na_pribil_plan)} = {fn(chapter_8.P_chistaya_plan)}', style=formula_style_12)
    add_formula(
        'П_{чист.\\ факт.} = П_{до\\ налогообл.\\ факт.} - S_{налог\\ факт.} = ' +
        f'{fn(chapter_8.P_pr_do_nalogov_fact)} - {fn(chapter_8.nalog_na_pribil_fact)} = {fn(chapter_8.P_chistaya_fact)}', style=formula_style_12)
    document.add_page_break()

    dp('Таблица 8, отчёт о финансовых результатах на конец первого периода', table_name_text)
    table = add_table([
        ['Наименование показателя', 'Сумма, руб/год', None],
        [None, 'план', 'факт'],
        ['Выручка', fn(chapter_8.Q_plan), fn(chapter_8.Q_fact)],
        ['Себестоимость продаж (проданной готовой продукции)', fn(chapter_8.S_pr_got_pr_plan), fn(chapter_8.S_pr_got_pr_fact)],
        ['Валовая прибыль (убыток)', fn(chapter_8.S_valovaya_plan), fn(chapter_8.S_valovaya_fact)],
        ['Коммерческие расходы', fn(chapter_8.S_kom_plan), fn(chapter_8.S_kom_fact)],
        ['Прибыль (убыток) от продаж', fn(chapter_8.P_pr_plan), fn(chapter_8.P_pr_fact)],
        ['Прочие доходы ', fn(0), fn(0)],
        ['Прочие расходы', fn(chapter_8.S_prochie_dohidy_i_rashody_plan), fn(chapter_8.S_prochie_dohidy_i_rashody_fact)],
        ['Прибыль (убыток) до налогообложения', fn(chapter_8.P_pr_do_nalogov_plan), fn(chapter_8.P_pr_do_nalogov_fact)],
        ['Налог на прибыль', fn(chapter_8.nalog_na_pribil_plan), fn(chapter_8.nalog_na_pribil_fact)],
        ['Чистая прибыль (убыток)', fn(chapter_8.P_chistaya_plan), fn(chapter_8.P_chistaya_fact)],

    ], [Cm(9), Cm(4), Cm(4)])
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 2))
    document.add_page_break()


def gen_1_9():
    dp('9. Плановый и фактический бухгалтерский баланс на конец периода', title_text)

    dp(f'Амортизация основных средств: {fn(chapter_3.costs["amortisation OS"].total)} [руб./год] (см п.3)')
    dp(f'Амортизация НМА: {fn(chapter_3.costs["amortisation NMA"].total)} [руб./год] (см п.3)')
    dp('Рассчитаем оборотные средства в незавершённом производстве и готовой продукции:')
    add_formula('K_{об_{нез.пр.\\ план}} = ' + f'{fn(chapter_5.K_ob_nez_pr)}\\ [руб.]\\ (см.\\ п.\\ 5.2)')
    add_formula('K_{об_{гот.пр.\\ план}} = ' + f'{fn(chapter_5.K_ob_got_prod)}\\ [руб.]\\ (см.\\ п.\\ 5.2)')
    add_formula('K_{об_{гот.под.\\ факт}} = K_{об._{гот.прод.}} + S_{Б_{произв}} N_{ост.} = ' + f'{fn(chapter_8.K_ob_got_prod_fact)}\\ [руб.]')

    dp()
    dp('Рассчитаем денежные средства:')
    add_formula('K_{ден.ср.\\ план} = K_{ден.ср.нач.период.} - A_{НМА} - А_{ОС} + П_{чист.\\ план} - (K_{об_{нез.пр.\\ план}} + K_{об_{гот.пр.\\ план}}) = ' +
                f'{fn(chapter_6.active_passive.K_ob_ds)} - {fn(chapter_3.costs["amortisation NMA"].total)} - {fn(chapter_3.costs["amortisation OS"].total)} + '
                f'{fn(chapter_8.P_chistaya_plan)} - ({fn(chapter_5.K_ob_nez_pr)} + {fn(chapter_8.K_ob_got_prod_plan)}) = {fn(chapter_9.K_den_sr_plan)}', style=formula_style_12)
    add_formula('K_{ден.ср.\\ факт} = K_{ден.ср.нач.период.} - A_{НМА} - А_{ОС} + П_{чист.\\ факт} - (K_{об_{нез.пр.\\ факт}} + K_{об_{гот.пр.\\ факт}}) = ' +
                f'{fn(chapter_6.active_passive.K_ob_ds)} - {fn(chapter_3.costs["amortisation NMA"].total)} - {fn(chapter_3.costs["amortisation OS"].total)} + '
                f'{fn(chapter_8.P_chistaya_fact)} - ({fn(chapter_5.K_ob_nez_pr)} + {fn(chapter_8.K_ob_got_prod_fact)}) = {fn(chapter_9.K_den_sr_fact)}', style=formula_style_12)

    dp()
    if chapter_9.valid_to_cope_kz_plan == 'full':
        dp('Возможно полное погашение краткосрочных заёмных средств для плановой суммы денежных средств:')
        add_formula('K_{ден.ср.конец\\ план} = K_{ден.ср.\\ план} - S_{кр.заёмн.ср.} = ' +
                    f'{fn(chapter_9.K_den_sr_plan)} - {fn(chapter_6.active_passive.kratkosroch_zaem_sredstva)} = {fn(chapter_9.K_den_sr_konez_plan)}\\ [руб.]',
                    style=formula_style_12)
    elif chapter_9.valid_to_cope_kz_plan == 'part':
        dp('Возможно частичное погашение краткосрочных заёмных средств для плановой суммы '
           'денежных средств (дальнейшая генерация неверна):').runs[0].font.color.rgb = RGBColor(255, 0, 0)
        add_formula('K_{ден.ср.конец\\ план} = K_{ден.ср.\\ план} - S_{кр.заёмн.ср.} = ' + f'{fn(500_000)} \\ [руб.]')
        add_formula('S_{кр.заёмн.ср.\\ план} = ' + f'{fn(chapter_9.S_kratkosroch_zaem_sredstva_konez_plan)} \\ [руб.]')
    else:
        dp('Погашение краткосрочных заёмных средств невозможно для плановой суммы денежных средств, '
           'дальнейшая генерация неверна.').runs[0].font.color.rgb = RGBColor(255, 0, 0)

    if chapter_9.valid_to_cope_kz_fact == 'full':
        dp('Возможно полное погашение краткосрочных заёмных средств для фактической суммы денежных средств:')
        add_formula('K_{ден.ср.конец\\ факт} = K_{ден.ср.\\ факт} - S_{кр.заёмн.ср.} = ' +
                    f'{fn(chapter_9.K_den_sr_fact)} - {fn(chapter_6.active_passive.kratkosroch_zaem_sredstva)} = {fn(chapter_9.K_den_sr_konez_fact)}\\ [руб.]',
                    style=formula_style_12)
    elif chapter_9.valid_to_cope_kz_fact == 'part':
        dp('Возможно частичное погашение краткосрочных заёмных средств для фактической суммы '
           'денежных средств (дальнейшая генерация неверна):').runs[0].font.color.rgb = RGBColor(255, 0, 0)
        add_formula('K_{ден.ср.конец\\ факт} = K_{ден.ср.\\ факт} - S_{кр.заёмн.ср.} = ' + f'{fn(500_000)} \\ [руб.]')
        add_formula('S_{кр.заёмн.ср.\\ факт} = ' + f'{fn(chapter_9.S_kratkosroch_zaem_sredstva_konez_fact)} \\ [руб.]')
    else:
        dp('Погашение краткосрочных заёмных средств невозможно для фактической суммы денежных средств, '
           'дальнейшая генерация неверна.').runs[0].font.color.rgb = RGBColor(255, 0, 0)

    document.add_page_break()

    dp('Таблица 9.1, плановый бухгалтерский баланс на конец периода', table_name_text)
    add_active_passive_table(chapter_9.active_passive_plan)

    document.add_page_break()

    dp('Таблица 9.2, фактический бухгалтерский баланс на конец периода', table_name_text)
    add_active_passive_table(chapter_9.active_passive_fact)
    document.add_page_break()


def gen_1_10():
    dp('10. Анализ деятельности условного предприятия', title_text)
    dp('Анализ деятельности условного предприятия осуществляется на основе основных плановых и фактических '
       'показателей хозяйственной деятельности предприятия и графика рентабельности.')

    def gen_1_10_1():
        dp()
        dp('10.1 Основные показатели хозяйственной деятельности предприятия', subtitle_text)

        dp()
        dp('Сумма хозяйственных средств:')
        add_formula('K_{хс.\\ план} = ' + f'{fn(chapter_9.active_passive_plan.active)}\\ [руб.]', style=formula_style_12)
        add_formula('K_{хс.\\ факт} = ' + f'{fn(chapter_9.active_passive_fact.active)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Собственные оборотные средства:')
        add_formula('k_{соб.об.ср.} = Оборотные\\ активы - Краткосрочные\\ обязательства')
        add_formula('k_{соб.об.ср.\\ план} = ' +
                    f'{fn(chapter_9.active_passive_plan.r2)} - {fn(chapter_9.active_passive_plan.r5)} = {fn(chapter_10.k_sob_ob_sr_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('k_{соб.об.ср.\\ факт} = ' +
                    f'{fn(chapter_9.active_passive_fact.r2)} - {fn(chapter_9.active_passive_fact.r5)} = {fn(chapter_10.k_sob_ob_sr_fact)}\\ [руб.]', style=formula_style_12)
        dp()
        dp('Коэффициент обеспеченности собственными средствами:')
        add_formula('k_{обеспеч.соб.ср.} = \\frac{Оборотные\\ активы - Краткосрочные\\ обязательства}{Оборотные\\ активы}')
        add_formula(
            'k_{обеспеч.соб.ср.\\ план} = \\frac{' +
            f'{fn(chapter_10.k_sob_ob_sr_plan)} }}{{ {fn(chapter_9.active_passive_plan.r2)} }} = {fn(chapter_10.k_obespech_sob_sr_plan)}', style=formula_style_12)
        add_formula(
            'k_{обеспеч.соб.ср.\\ факт} = \\frac{' +
            f'{fn(chapter_10.k_sob_ob_sr_fact)}  }}{{ {fn(chapter_9.active_passive_fact.r2)} }} = {fn(chapter_10.k_obespech_sob_sr_fact)}', style=formula_style_12)

        document.add_page_break()

        dp('Коэффициент абсолютной ликвидности:')
        add_formula('k_{абс.ликв.} = \\frac{Абсолютно\\ ликвидныке\\ активы}{Краткосрочные\\ обязательства}')
        add_formula('k_{абс.ликв.\\ план} = \\frac{' +
                    f'{fn(chapter_9.active_passive_plan.K_ob_ds)} }}{{ {fn(chapter_9.active_passive_plan.r5)} }} = {fn(chapter_10.k_abs_likvid_plan)}', style=formula_style_12)
        add_formula('k_{абс.ликв.\\ факт} = \\frac{' +
                    f'{fn(chapter_9.active_passive_fact.K_ob_ds)} }}{{ {fn(chapter_9.active_passive_fact.r5)} }} = {fn(chapter_10.k_abs_likvid_fact)}', style=formula_style_12)

        dp()
        dp('Коэффициент текущей ликвидности (или коэффициент покрытия баланса):')
        add_formula('k_{тек.ликв.} = \\frac{Сумма\\ оборотных\\ активов}{Краткосрочные\\ обязательства}')
        add_formula('k_{тек.ликв.\\ план} = \\frac{' +
                    f'{fn(chapter_9.active_passive_plan.r2)} }}{{ {fn(chapter_9.active_passive_plan.r5)} }} = {fn(chapter_10.k_tek_likvid_plan)}', style=formula_style_12)
        add_formula('k_{тек.ликв.\\ факт} = \\frac{' +
                    f'{fn(chapter_9.active_passive_fact.r2)} }}{{ {fn(chapter_9.active_passive_fact.r5)} }} = {fn(chapter_10.k_tek_likvid_fact)}', style=formula_style_12)

        dp()
        dp('Выручка от продажи продукции:')
        add_formula('Q_{план} = ' + f'{fn(chapter_8.Q_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('Q_{факт} = ' + f'{fn(chapter_8.Q_fact)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Нераспределенная прибыль:')
        add_formula('П_{нерасп.\\ план} = ' + f'{fn(chapter_9.active_passive_plan.neraspred_pribil)}\\ [руб.]', style=formula_style_12)
        add_formula('П_{нерасп.\\ факт} = ' + f'{fn(chapter_9.active_passive_fact.neraspred_pribil)}\\ [руб.]', style=formula_style_12)

        document.add_page_break()

        dp('Выработка продукции на одного работника:')
        add_formula('V = \\frac{Объём\\ продукции}{Среднесписочное\\ кол-во\\ ППП} = \\frac{' +
                    f'{fn(initial_data.N_pl, 0)} }}{{ {fn(chapter_2.R_ppp, 0)} }} = {fn(chapter_10.V)}\\ [шт./чел.год]', style=formula_style_12)

        dp()
        dp('Среднегодовая стоимость ОПФ:')
        add_formula(
            'S_{ср.год.ст.ОПФ} = S_{ОПФ\\ нач.пер.} - А_{ОПФ} \\cdot 0.5 = ' +
            f'{fn(chapter_1.S_os_amortisable)} - {fn(chapter_3.costs["amortisation OS"].total)} = {fn(chapter_10.OS_year_mean)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Коэффициент фондоотдачи:')
        add_formula('k_{ФО\\ план} = \\frac{Q_{план}}{Среднегодовая\\ стоимость\\ ОПФ} = \\frac{' +
                    f'{fn(chapter_8.Q_plan)} }}{{ {fn(chapter_10.OS_year_mean)} }} = {fn(chapter_10.k_FO_plan)}', style=formula_style_12)
        add_formula('k_{ФО\\ факт} = \\frac{Q_{факт}}{Среднегодовая\\ стоимость\\ ОПФ} = \\frac{' +
                    f'{fn(chapter_8.Q_fact)} }}{{ {fn(chapter_10.OS_year_mean)} }} = {fn(chapter_10.k_FO_fact)}', style=formula_style_12)

        dp()
        dp('Коэффициент фондоемкости:')
        add_formula('k_{ФЕ\\ план} = k_{ФО\\ план}^{-1} = ' + f'{fn(chapter_10.k_FO_plan)} ^ {{-1}} = {fn(chapter_10.k_FE_plan)}', style=formula_style_12)
        add_formula('k_{ФЕ\\ факт} = k_{ФО\\ факт}^{-1} = ' + f'{fn(chapter_10.k_FO_fact)} ^ {{-1}} = {fn(chapter_10.k_FE_fact)}', style=formula_style_12)

        dp()
        dp('Число оборотов оборотных средств:')
        add_formula('Ср.сумм.исп.об.ср._{план} = ' + f'{fn(chapter_10.Z_ob_sr_year_mean_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('Ср.сумм.исп.об.ср._{факт} = ' + f'{fn(chapter_10.Z_ob_sr_year_mean_fact)}\\ [руб.]', style=formula_style_12)
        add_formula('Z_{об} = \\frac{Выручка\\ от\\ реализации}{Средняя\\ сумма\\ используемых\\ обороных\\ средств}')
        add_formula('Z_{об\\ план} = \\frac{' +
                    f'{fn(chapter_8.Q_plan)} }}{{ {fn(chapter_10.Z_ob_sr_year_mean_plan)} }} = {fn(chapter_10.Z_ob_sr_year_mean_plan)} [раз/год]', style=formula_style_12)
        add_formula('Z_{об\\ факт} = \\frac{' +
                    f'{fn(chapter_8.Q_fact)} }}{{ {fn(chapter_10.Z_ob_sr_year_mean_fact)} }} = {fn(chapter_10.Z_ob_sr_year_mean_fact)} [раз/год]', style=formula_style_12)

        document.add_page_break()

        dp('Оборачиваемость собственного капитала:')
        add_formula('Ср.год.собств.кап_{план} = ' + f'{fn(chapter_10.S_sobstv_cap_year_mean_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('Ср.год.собств.кап._{факт} = ' + f'{fn(chapter_10.S_sobstv_cap_year_mean_fact)}\\ [руб.]', style=formula_style_12)
        add_formula('k_{об.собств.кап.} = \\frac{Выручка\\ от\\ реализации}{Ср.год.собств.кап}')
        add_formula('k_{об.собств.кап.\\ план} = \\frac{' +
                    f'{fn(chapter_8.Q_plan)} }}{{ {fn(chapter_10.S_sobstv_cap_year_mean_plan)} }} = {fn(chapter_10.k_oborach_sobstv_capital_plan)}', style=formula_style_12)
        add_formula('k_{об.собств.кап.\\ факт} = \\frac{' +
                    f'{fn(chapter_8.Q_fact)} }}{{ {fn(chapter_10.S_sobstv_cap_year_mean_fact)} }} = {fn(chapter_10.k_oborach_sobstv_capital_fact)}', style=formula_style_12)

        dp()
        dp('Рентабельность продукции:')
        add_formula('R_{продукции} = \\frac{Прибыль\\ от\\ продаж}{Себестоимость\\ продаж}')
        add_formula('R_{продукции\\ план} = \\frac{' +
                    f'{fn(chapter_8.P_pr_plan)} }}{{ {fn(chapter_4.S_sum.total)} }} = {fn(chapter_10.R_production_plan)}', style=formula_style_12)
        add_formula('R_{продукции\\ факт} = \\frac{' +
                    f'{fn(chapter_8.P_pr_fact)} }}{{ {fn(chapter_4.S_sum.total)} }} = {fn(chapter_10.R_production_fact)}', style=formula_style_12)

        dp()
        dp('Рентабельность продаж:')
        add_formula('R_{продаж} = \\frac{Чистая\\ прибыль}{Выручка}')
        add_formula('R_{продаж\\ план} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_plan)} }}{{ {fn(chapter_8.Q_plan)} }} = {fn(chapter_10.R_sell_plan)}', style=formula_style_12)
        add_formula('R_{продаж\\ факт} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_fact)} }}{{ {fn(chapter_8.Q_fact)} }} = {fn(chapter_10.R_sell_fact)}', style=formula_style_12)

        document.add_page_break()
        dp('Рентабельность активов:')
        add_formula('R_{активов} = \\frac{Чистая\\ прибыль}{Актив}')
        add_formula('R_{активов\\ план} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_plan)} }}{{ {fn(chapter_9.active_passive_plan.active)} }} = {fn(chapter_10.R_active_plan)}', style=formula_style_12)
        add_formula('R_{активов\\ факт} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_fact)} }}{{ {fn(chapter_9.active_passive_fact.active)} }} = {fn(chapter_10.R_active_fact)}', style=formula_style_12)

        dp()
        dp('Рентабельность собственного капитала:')
        add_formula('R_{собств.кап.} = \\frac{Чистая\\ прибыль}{Актив}')
        add_formula('R_{собств.кап.\\ план} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_plan)} }}{{ {fn(chapter_10.S_sobstv_cap_year_mean_plan)} }} = {fn(chapter_10.R_sobstv_capital_plan)}', style=formula_style_12)
        add_formula('R_{собств.кап.\\ факт} = \\frac{' +
                    f'{fn(chapter_8.P_chistaya_fact)} }}{{ {fn(chapter_10.S_sobstv_cap_year_mean_plan)} }} = {fn(chapter_10.R_sobstv_capital_fact)}', style=formula_style_12)

        document.add_page_break()

        dp('Таблица 10, Плановые и фактические значения основных показателей хозяйственной деятельности предприятия в I периоде', table_name_text)
        add_table([
            ['Наименование показателя и его размерность', 'План', 'Факт'],
            ['Сумма хозяйственных средств, [руб.]', fn(chapter_9.active_passive_plan.active), fn(chapter_9.active_passive_fact.active)],
            ['Собственные оборотные средства, [руб.]', fn(chapter_10.k_sob_ob_sr_plan), fn(chapter_10.k_sob_ob_sr_fact)],
            ['Коэффициент обеспеченности собственными средствами', fn(chapter_10.k_obespech_sob_sr_plan), fn(chapter_10.k_obespech_sob_sr_fact)],
            ['Коэффициент абсолютной ликвидности', fn(chapter_10.k_abs_likvid_plan), fn(chapter_10.k_abs_likvid_fact)],
            ['Коэффициент текущей ликвидности', fn(chapter_10.k_tek_likvid_plan), fn(chapter_10.k_tek_likvid_fact)],
            ['Выручка от продажи продукции, [руб.]', fn(chapter_8.Q_plan), fn(chapter_8.Q_fact)],
            ['Нераспределенная прибыль, [руб.]', fn(chapter_9.active_passive_plan.neraspred_pribil), fn(chapter_9.active_passive_fact.neraspred_pribil)],
            ['Выработка продукции на одного работника [шт./чел.год]', fn(chapter_10.V), fn(chapter_10.V)],
            ['Среднегодовая стоимость ОПФ, [руб.]', fn(chapter_10.OS_year_mean), fn(chapter_10.OS_year_mean)],
            ['Коэффициент фондоотдачи [1/руб.]', fn(chapter_10.k_FO_plan), fn(chapter_10.k_FO_fact)],
            ['Коэффициент фондоемкости [руб.]', fn(chapter_10.k_FE_plan), fn(chapter_10.k_FE_fact)],
            ['Число оборотов оборотных средств, [раз/год]', fn(chapter_10.Z_ob_plan), fn(chapter_10.Z_ob_fact)],
            ['Оборачиваемость собственного капитала', fn(chapter_10.k_oborach_sobstv_capital_plan), fn(chapter_10.k_oborach_sobstv_capital_fact)],
            ['Рентабельность продукции', fn(chapter_10.R_production_plan), fn(chapter_10.R_production_fact)],
            ['Рентабельность продаж', fn(chapter_10.R_sell_plan), fn(chapter_10.R_sell_fact)],
            ['Рентабельность активов [1/год]', fn(chapter_10.R_active_plan), fn(chapter_10.R_active_fact)],
            ['Рентабельность собственного капитала [1/год]', fn(chapter_10.R_sobstv_capital_plan), fn(chapter_10.R_sobstv_capital_fact)],
        ], [Cm(9.25), Cm(3.75), Cm(3.75)], True, style=table_style_12)

        document.add_page_break()

    gen_1_10_1()
    dp('10.2 График рентабельности изделия Б', subtitle_text)
    dp('Построим график рентабельности в соответствии с полученными значениями, рассчитаем точку безубыточности, '
       'коэффициент покрытия, запас финансовой прочности и величину операционного рычага для плановых условий.')

    dp('Точку безубыточности (критический объем продаж) определим:')
    add_formula('N_{кр} = \\frac{S_{усл.пост.}}{Ц_{Б\\ произв\\ план} - S_{Б\\ перем}}')
    p = dp('Замечу, что предоставленная формула не учитывает того, что ')
    add_formula('S_{Б\\ перем}}', p)
    p.add_run(' может меняться от объёма продукции (см п. 4.3), поэтому воспользуемся иной формулой:')
    add_formula('N_{кр} \\cdot {Ц_{Б\\ произв\\ план} = S_{сум}(N_{кр})')
    add_formula('Q_{кр} = N_{кр} \\cdot {Ц_{Б\\ произв\\ план}} = ' + f'{fn(chapter_10.N_kr, 0)} \\cdot {fn(chapter_7.P_proizv_plan)} = {fn(chapter_10.Q_kr)}')
    dp('Решение найдём итеративно, с помощью бинарного поиска:')
    add_formula('N_{кр} = ' + f'{fn(chapter_10.N_kr, 0)}')

    document.add_page_break()
    dp('Коэффициент покрытия:')
    add_formula('k_{покр} = \\frac{Ц_{Б\\ произв\\ план} - S_{Б\\ перем}}{Ц_{Б\\ произв\\ план}')
    p = dp('Здесь формула снова не учитывает того, что ')
    add_formula('S_{Б\\ перем}}', p)
    p.add_run(' может меняться от объёма продукции:')
    add_formula('k_{покр\\ N} = \\frac{Ц_{Б\\ произв\\ план} - S_{Б\\ перем}(N)}{Ц_{Б\\ произв\\ план}')

    table = add_table([['N'] + [fn(e, 0) for e in chapter_10.k_pokr.input_data], [''] + [fn(e, 3) for e in chapter_10.k_pokr.output_data]])
    add_formula('k_{покр}', table.cell(1, 0).paragraphs[0])

    dp()
    dp('Запас финансовой прочности:')
    add_formula('Q_{фин\\ пр.} = \\frac{Q - Q_{кр}}{Q} \\cdot 100%')
    add_formula('Q_{фин\\ пр.\\ план} = \\frac{ ' + f'{fn(chapter_8.Q_plan)} - {chapter_10.Q_kr} }}{{ {fn(chapter_8.Q_plan)} }} = {fn(chapter_10.Q_fin_pr_plan * 100)} \\%')
    add_formula('Q_{фин\\ пр.\\ факт} = \\frac{ ' + f'{fn(chapter_8.Q_fact)} - {chapter_10.Q_kr} }}{{ {fn(chapter_8.Q_fact)} }} = {fn(chapter_10.Q_fin_pr_fact * 100)} \\%')

    dp()
    dp('Найдём эффект производственного рычага:')
    add_formula('E_{пр.\\ рыч.} = \\frac{Маржинальная\\ прибыль}{Прибыль\\ от\\ продаж}')
    add_formula(
        'E_{пр.\\ рыч.\\ план} = \\frac{ ' +
        f'{fn(chapter_8.Q_plan)} - {chapter_4.S_sum.variable} }}{{ {fn(chapter_8.P_pr_plan)} }} = {fn(chapter_10.proizv_richag_plan)}')
    add_formula(
        'E_{пр.\\ рыч.\\ факт} = \\frac{ ' +
        f'{fn(chapter_8.Q_plan)} - {fn(chapter_4.S_b_poln.variable)} \\cdot {fn(chapter_8.N_fact)} }}{{ {fn(chapter_8.P_pr_fact)} }} = {fn(chapter_10.proizv_richag_fact)}')

    document.add_page_break()

    dp('График 10, рентабельность изделия Б', table_name_text)

    plt.figure(figsize=(8, 8))
    plt.subplot(1, 1, 1)

    ct1 = chapter_4.ct1

    S_const_costs = [e.const for e in ct1.output_data]
    S_variable_costs = [e.variable for e in ct1.output_data]
    S_total_costs = [e.total for e in ct1.output_data]

    N_pl_values = chapter_4.N_pl_values

    plt.title('S(N)')
    plt.xlabel('N, шт. / год')
    plt.ylabel('Выручка, затраты, тыс. руб./год')
    plt.xticks([0, chapter_10.N_kr, initial_data.N_pl], ['0', 'N кр\n{:,.0f}'.format(chapter_10.N_kr), 'N пл\n{:,.0f}'.format(initial_data.N_pl)], rotation=0)
    plt.yticks([0, S_const_costs[-1] / 1e3, chapter_10.N_kr * chapter_7.P_proizv_plan / 1e3, S_total_costs[-1] / 1e3, chapter_8.Q_plan / 1e3],
               ['0', 'S усл.пост.\n{:,.0f}'.format(S_const_costs[-1] / 1e3), 'Q кр.\n{:,.0f}'.format(chapter_10.Q_kr / 1e3),
                '{:,.0f}\nS сум.'.format(S_total_costs[-1] / 1e3), 'Q пл\n{:,.0f}'.format(chapter_8.Q_plan / 1e3)])
    plt.grid(True)
    # plt.plot([0, N_kr], [N_kr * TS_B_proizv_plan / 1e3, N_kr * TS_B_proizv_plan / 1e3], color='#444444', ls=':')
    plt.plot(N_pl_values, np.array(S_const_costs) / 1e3, label='S усл-пост.', ls=':')
    plt.plot(N_pl_values, np.array(S_variable_costs) / 1e3, label='S перем.', ls=':')
    plt.plot(N_pl_values, np.array(S_total_costs) / 1e3, label='S тек.сум.')
    plt.plot([0, initial_data.N_pl], [0, chapter_8.Q_plan / 1e3], label='Q пл.')
    plt.legend()
    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile)

    picP = document.add_paragraph()
    picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picP.add_run().add_picture(memfile, width=Cm(17))
    picP.add_run().add_break(WD_BREAK.PAGE)


def gen_1_11():
    dp('Выводы', title_text)
    document.add_page_break()


def gen_2_1():
    dp('РАЗДЕЛ II', title_text)
    dp()
    dp('1. Планирование объемов производства изделий А, Б и В', title_text)
    dp('По результатам работы предприятия в первом периоде часть объема выпуска изделия Б не была реализована даже при рыночной цене (пониженной по сравнению с '
       'установленной плановой), предполагаем, что часть потребителей не удовлетворена качеством выпускаемого изделия. '
       'Причем одна группа потребителей готова приобретать аналог А − более высокого качества по более высокой цене, другая группа считает, что некоторые '
       'свойства изделия Б излишни, и готова приобретать изделие-аналог В более низкого качества и за более умеренную цену. '
       'Предприятие готово пойти навстречу и организовать производство изделий А и В')
    dp('Маркетинговое исследование показало что 50% покупателей удовлетворены качеством изделия Б, а 30% хотят что-либо изменить в нём.')

    dp('Таблица 21.1, объём производства', table_name_text)
    add_table([
        ['Изделия', 'А', 'Б', 'В'],
        ['Объём производства, %', fn(chapter_2_1.A_percent * 100, 0), fn(chapter_2_1.B_percent * 100, 0), fn(chapter_2_1.C_percent * 100, 0)],
        ['Объём производства, шт/год', fn(chapter_2_1.N_pl_A, 0), fn(chapter_2_1.N_pl_B, 0), fn(chapter_2_1.N_pl_C, 0)]
    ], [Cm(8), Cm(3), Cm(3), Cm(3)], True)

    document.add_page_break()


def gen_2_2():
    dp('2. Дополнительная потребность в оборудовании', title_text)
    dp('Расчет дополнительных потребностей в')
    dp('оборудовании (при этом составить стоимостную структуру основных средств)')
    dp('персонале, рассчитав численность работающих, величины фонда оплаты труда и единого социального налога')
    dp('материалах и комплектующих изделиях, используя исходную информацию')
    dp()

    dp('Таблица 22.1, вид материалов', table_name_text)
    table = add_table([
        ['Вид материала', None, '1', '2', '3', '4'],
        ['Стоимость, руб./ед.измер.', 'A'] + [str(e['cost']) for e in initial_data.materials_A.rows],
        [None, 'B'] + [str(e['cost']) for e in initial_data.materials_C.rows],
        ['Норма расхода, ед.измер./шт.', 'A'] + [str(e['amount']) for e in initial_data.materials_A.rows],
        [None, 'B'] + [str(e['amount']) for e in initial_data.materials_C.rows],
        ['Итого материалов на изделие А, руб./шт.', None, None, None, fn(initial_data.materials_A.calculate_sum(lambda x: x['amount'] * x['cost'])), None],
        ['Итого материалов на изделие B, руб./шт.', None, None, None, fn(initial_data.materials_C.calculate_sum(lambda x: x['amount'] * x['cost'])), None]
    ], [Cm(7), Cm(2), Cm(2), Cm(2), Cm(2), Cm(2)], True, style=table_style)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).merge(table.cell(2, 0))
    table.cell(3, 0).merge(table.cell(4, 0))
    table.cell(5, 0).merge(table.cell(5, 3))
    table.cell(5, 4).merge(table.cell(5, 5))
    table.cell(6, 0).merge(table.cell(6, 3))
    table.cell(6, 4).merge(table.cell(6, 5))

    dp()
    dp('Таблица 22.2, вид комплектующих', table_name_text)
    table = add_table([
        ['Вид материала', None, '1', '2', '3'],
        ['Стоимость, руб./ед.измер.', 'A'] + [str(e['cost']) for e in initial_data.accessories_A.rows],
        [None, 'B'] + [str(e['cost']) for e in initial_data.accessories_C.rows],
        ['Норма расхода, ед.измер./шт.', 'A'] + [str(e['amount']) for e in initial_data.accessories_A.rows],
        [None, 'B'] + [str(e['amount']) for e in initial_data.accessories_C.rows],
        ['Итого комплектующих на изделие А, руб./шт.', None, None, fn(initial_data.accessories_A.calculate_sum(lambda x: x['amount'] * x['cost'])), None],
        ['Итого комплектующих на изделие B, руб./шт.', None, None, fn(initial_data.accessories_C.calculate_sum(lambda x: x['amount'] * x['cost'])), None]
    ], [Cm(7), Cm(2), Cm(2), Cm(2), Cm(2)], True)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).merge(table.cell(2, 0))
    table.cell(3, 0).merge(table.cell(4, 0))
    table.cell(5, 0).merge(table.cell(5, 2))
    table.cell(5, 3).merge(table.cell(5, 4))
    table.cell(6, 0).merge(table.cell(6, 2))
    table.cell(6, 3).merge(table.cell(6, 4))

    document.add_page_break()

    dp('Таблица 22.3, технологическая трудоёмкость изделий', table_name_text)

    table = add_table(
        [
            ['Номер операции', 'Используемое оборудование', 'Стоимость используемого оборудования, тыс.руб./ед.оборуд.', 'Технологическая трудоёмкость'],
            ['Изделие А', None, None, None],
        ] + [[str(i + 1), e['name'], fn(e['cost'] // 1000, 0), fn(e['time'])] for i, e in enumerate(initial_data.operations_A.rows)] +
        [['Изделие Б', None, None, None]] + [[str(i + 1), e['name'], fn(e['cost'] // 1000, 0), fn(e['time'])] for i, e in enumerate(initial_data.operations_B.rows)] +
        [['Изделие В', None, None, None]] + [[str(i + 1), e['name'], fn(e['cost'] // 1000, 0), fn(e['time'])] for i, e in enumerate(initial_data.operations_C.rows)],
        [Cm(3), Cm(3.6), Cm(6.4), Cm(3.5)], True
    )
    table.cell(1, 0).merge(table.cell(1, 3)).paragraphs[0].runs[0].bold = True
    table.cell(2 + len(initial_data.operations_A), 0).merge(
        table.cell(2 + len(initial_data.operations_A), 3)).paragraphs[0].runs[0].bold = True
    table.cell(3 + len(initial_data.operations_A) + len(initial_data.operations_B), 0).merge(
        table.cell(3 + len(initial_data.operations_A) + len(initial_data.operations_B), 3)).paragraphs[0].runs[0].bold = True

    dp()
    dp('Как видим, оборудование б, в, г подорожало за прошедший год. Учтём это.', no_indent=True)
    add_formula('\\beta_{норм} = 0.7', dp('Предполагается, что все наименования изделий обрабатываются в основном на одном и том же оборудовании, поэтому с учетом '
                                          'необходимости переналадки планируемый коэффициент загрузки '))
    dp('Для определения необходимого количества оборудования используем формулу:')
    add_formula('n_{об\\ k_{расч}} = \\frac{t_{Ak}N_A+t_{Бk}N_Б+t_{Вk}N_В}{\\beta_{норм} F_{об\\ эф}}')

    document.add_page_break()

    dp('Таблица 22.4, дополнительная потребность в оборудовании во втором периоде', table_name_text)
    table = add_table(
        [['Вид и стоимость оборудования, тыс. руб.',
          'Техн. оборуд. во втором периоде', None,
          'Имеющееся количесво оборудования',
          'Доп. количество оборудования',
          'Стоимость доп. оборудования, руб.',
          'Фактич. нагрузка на оборуд.'], [None, None, None, None, None, None, None]] +
        [[f'{e["name"]}, {fn(e["cost"] // 1000, 0)}',
          fn(e['need_rasch']),
          fn(e['need_fact'], 0),
          fn(e['stock'], 0),
          fn(e['need_new'], 0),
          fn(e['cost'] * e['need_new'], 0),
          fn(e['b_fact'])] for e in chapter_2_2.machines.rows] +
        [['Итого', None,
          fn(chapter_2_2.machines.calculate_sum(lambda x: x['need_fact']), 0),
          fn(chapter_2_2.machines.calculate_sum(lambda x: x['stock']), 0),
          fn(chapter_2_2.machines.calculate_sum(lambda x: x['need_new']), 0),
          fn(chapter_2_2.new_machines_cost, 0), None]],
        [Cm(3.2), Cm(1.75), Cm(1.75), Cm(2.75), Cm(2.75), Cm(2.75), Cm(2)], style=table_style_12)
    add_formula('n_{об.i_{расч.}}', table.cell(1, 1).paragraphs[0])
    add_formula('n_{об.i_{прин.}}', table.cell(1, 2).paragraphs[0])
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 2))
    for i in [3, 4, 5, 6]:
        table.cell(0, i).merge(table.cell(1, i))

    dp()
    dp('Таблица 22.5, стоимостная структура ОС на начало второго периода', table_name_text)
    table = add_table(
        [['№',
          'Группа (виды) основных средств (ОС)',
          'Стоимость ОС на начало I периода, руб.',
          'Амортизация ОС за I период, руб.',
          'Стоимость ОС на начало II периода, руб.',
          'Изменение ОС, руб.',
          'Стоимость ОС на начало II периода, с учетом прироста ОС, руб.',
          '%']] +
        [[e['n'], e['name'], fn(e['cost I'], 0), fn(e['amortisation I'], 0),
          fn(e['cost II begin'], 0), fn(e['delta'], 0), fn(e['cost II'], 0), fn(e['cost II'] / chapter_2_2.S_os * 100, 1)] for e in chapter_2_2.main_resources.rows] +
        [[None, 'Итого',
          fn(chapter_2_2.main_resources.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['cost I']), 0),
          fn(chapter_2_2.main_resources.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['amortisation I']), 0),
          fn(chapter_2_2.main_resources.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['cost II begin']), 0),
          fn(chapter_2_2.main_resources.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['delta']), 0),
          fn(chapter_2_2.main_resources.filter_table(lambda x: len(x['n']) > 0).calculate_sum(lambda x: x['cost II']), 0),
          '100']],
        [Cm(0.7), Cm(3.25), Cm(2.25), Cm(2), Cm(2.25), Cm(2), Cm(3.25), Cm(1.0)],
        first_bold=True, style=table_style_10
    )
    table.cell(len(table.rows) - 1, 0).merge(table.cell(len(table.rows) - 1, 1))


def gen_2_3():
    dp('3. Дополнительная потребность в промышленно-производственном персонале', title_text)
    dp('Рассчитаем численность ППП, ФОТ, страховые взносы аналогично первому периоду. Оклады остаются неизменными')

    dp('Чтобы найти численность ОПР найдём в начале эффективный фонд времени одного работающего:')
    add_formula_with_description(
        'F_{раб_{эф}}=(T_{пл}-B-O-H) \\cdot D = ' +
        f"({chapter_2_1.T_pl}-{chapter_2_1.B}-{chapter_2_1.O}-{chapter_2_1.H}) \\cdot {chapter_2_1.D} = {chapter_2_3.F_rab_ef}\\ [час/год\\ чел.]",
        [
            ['O', f'продолжительность отпуска = {chapter_2_1.H} раб. дн.'],
            ['H', f'число планируемых невыходов = {chapter_2_1.O} раб. дн.'],
        ],
        style=formula_style_12)
    dp('Примем величину планируемых невыходов за 20 (производство изделий на токарных станках негативно сказывается на здоровье за счёт мелкодисперсной '
       'металлической стружки, переменных магнитных полей электродвигателей и так далее), отпуск 20 рабочих дней (28 – 8 нерабочих дней за 4 недели).')

    dp('Подставим полученное значение в формулу численности ОПР:')

    add_formula_with_description(
        'R_{ОПР}=\\frac{N_{плА,Б,В} \\sum_{i}^{m}{t_{техн\\ i\\ А,Б,В}}}{F_{раб_{эф}} k_{вн}} =' + f'{chapter_2_3.R_opr}\\ [чел.]',
        [['t_{техн\\ i\\ А,Б,В}', 'трудоёмкость i-той операции изделия А, Б, В']]
    )

    dp()
    p = dp(f'Численность ВРП примем за {fn(chapter_2_3.R_vpr / chapter_2_3.R_opr)} R')
    p.add_run('ОПР').font.subscript = True
    p.add_run(', а численность служащих за 0.6R')
    p.add_run('ОПР').font.subscript = True
    p.add_run(':')

    add_formula('R_{ВПР} = ' + fn(chapter_2_3.R_vpr / chapter_2_3.R_opr) + 'R_{ОПР} = ' + f'{chapter_2_3.R_vpr}')
    add_formula('R_{СЛ} = ' + fn(chapter_2_3.R_sl / chapter_2_3.R_opr) + 'R_{ОПР} = ' + f'{chapter_2_3.R_sl}')

    dp('Численность ППП:')
    add_formula('R_{ППП} = R_{ОПР} + R_{ВПР} + R_{СЛ} = ' + f'{chapter_2_3.R_opr} + {chapter_2_3.R_vpr} + {chapter_2_3.R_sl} = {chapter_2_3.R_ppp}')

    dp('Средняя тарифная ставка:')
    add_formula('С_{тар.\\ ст.} = \\frac{12 \\cdot ' + f'{fn(chapter_2_3.opr_salary)} }}{{ {fn(chapter_2_3.F_rab_ef)} }} = {fn(chapter_2_3.C_opr_mean)}')
    chapter_2_3.FOT_opr

    document.add_page_break()

    dp('Таблица 23.1, состав, структура и заработная плата персонала', table_name_text)
    add_employee_structure_table(chapter_2_3)
    document.add_page_break()
    dp('Таблица 23.2, суммарные заработные платы персонала за год', table_name_text)
    add_employee_salary_table(chapter_2_3)
    dp()
    dp(f'Отметим, что ОПР получают сдельную зарплату за {fn(chapter_2_3.R_opr_raw)} чел., но на предприятии их трудится {chapter_2_3.R_opr}, поэтому каждый из них получит меньше, '
       f'но их суммарная з/п будет равна ФОТ ОПР.').add_run().add_break(WD_BREAK.PAGE)

    dp('Найдём ФОТ ОПР на год:')
    add_formula('ФОТ_{ОПР без надб.} = p_{ср} N_{пл} m = ' + f'{fn(chapter_2_3.FOT_opr)}\\ [руб./год]')

    dp()
    dp('Учтём надбавки ОПР:')
    add_formula(
        'ФОТ_{ОПР} = ФОТ_{ОПР} + R_{ОПР} \\cdot (ОПР_{надб.} \\cdot 12 + ОПР_{тариф.ст.}) = ' +
        f'{chapter_2_3.R_opr} \\cdot ({fn(chapter_2_3.opr_extra)} \\cdot 12 + {fn(chapter_2_3.opr_salary)} = {fn(chapter_2_3.FOT_opr + chapter_2_3.FOT_opr_extra)}\\ [руб./год]')

    dp()
    dp('ФОТ ВПР и служащих:')
    add_formula_with_description('ФОТ_{ВПР+сл} = \\sum_{i}^{n}{(ТС_i \\cdot N_i \\cdot 12 + ТС_i)} = ' + f'{fn(chapter_2_3.FOT_vpr + chapter_2_3.FOT_sl)}', [
        ['n', 'число ОПР и служащих'],
        ['ТС_i', 'тарифная ставка'],
        ['N_i', 'численнойсть']
    ])

    dp()
    dp('Общий ФОТ:')
    add_formula('ФОТ_{общ} = ФОТ_{ОПР} + ФОТ_{ВПР+сл} = ' + f'{fn(chapter_2_3.FOT.total)}')

    dp()
    dp('Таблица 2.3, страховые взносы', table_name_text)
    table = add_table([['Взнос', 'Величина, %', 'Сумма, руб./год']] + [[e.name, fn(e.percent * 100, 1), fn(e.amount)] for e in chapter_2_3.insurance_fee.rows])
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    r.cells[0].paragraphs[0].add_run('Итого').bold = True
    r.cells[2].paragraphs[0].add_run(fn(chapter_2_3.FOT_fee.total)).bold = True
    dp()
    dp('Выводы: ')
    document.add_page_break()


def gen_2_4():
    dp('4. Суммарные затраты на производство и реализацию продукции', title_text)
    dp('Для определения затрат на производство необходимо определить дополнительную потребность в материалах и комплектующих изделиях')
    dp()
    dp('4.1. Дополнительная потребность в материалах и комплектующих изделиях', subtitle_text)
    dp('См 2 пункт II раздела')
    dp('Стоимость материалов и комплектующих:')
    dp(f'А: {fn(chapter_2_4.S_materials_i_comp_A)} [руб./шт.]')
    dp(f'Б: {fn(chapter_3.S_mat_i_comp)} [руб./шт.]')
    dp(f'В: {fn(chapter_2_4.S_materials_i_comp_C)} [руб./шт.]')

    dp()
    dp('Смета затрат на производство:')
    dp('1. Стоимость основных материалов')
    add_formula('S_{ом} = ' + f'{fn(chapter_2_4.costs["material_main"].total)}  [руб./год]')

    dp(f'2. Стоимость вспомогательных материалов (принято за {fn(chapter_3.help_materials_percent * 100)}% от *)')
    add_formula('S_{вм} = S_{ом} \\cdot k_{вм} = ' + f'{fn(chapter_2_4.costs["helper"].total)}\\ [руб./год]')

    dp(f'3. Транспортно-заготовительные расходы (принято за {fn(chapter_3.moving_save_percent * 100)}% от *)')
    add_formula('S_{т-з} = S_{ом} \\cdot k_{т-з} = ' + f'{fn(chapter_2_4.costs["move save"].total)}\\ [руб./год]')
    dp(f'Из них {fn(chapter_3.move_save_const_percent * 100, 0)}% - постоянные затраты {fn(chapter_2_4.costs["move save"].const)}\\ [руб./год]')

    dp(f'4. Инструменты, инвентарь (принято за {fn(chapter_3.inventory_percent * 100)}% от *)')
    add_formula('S_{инстр} = S_{ом} \\cdot k_{интср} = ' + f'{fn(chapter_2_4.costs["inventory"].total)}\\ [руб./год]')

    dp(f'5. Топливо и энергия (принято за {fn(chapter_3.fuel_percent * 100)}% от *)')
    add_formula('S_{топл +эн} = S_{ом} \\cdot k_{топл+эн} = ' + f'{fn(chapter_2_4.costs["fuel total"].total)}\\ [руб./год]')

    dp(f'5.1 Технологическое ({fn(chapter_3.fuel_tech_percent * 100)}% от топлива и энергии)')
    add_formula('S_{тех. топл + эн} = S_{топл + эн} \\cdot k_{техн} = ' + f'{fn(chapter_2_4.costs["fuel tech"].total)}\\ [руб./год]')
    dp(f'5.2 Нетехнологическое ({fn(chapter_3.fuel_non_tech_percent * 100)}% от топлива и энергии)')
    add_formula('S_{тех. топл + эн} = S_{топл + эн} \\cdot (1 - k_{тех.}) = ' + f'{fn(chapter_2_4.costs["fuel non tech"].total)}\\ [руб./год]')

    dp('* - от стоимости основных материалов и комплектующих')

    dp(f'Норму амортизации основных средств примем за {fn(chapter_3.OS_amortisation_percent * 100)}%:')
    add_formula('A_{ос} = k_{ам.ос.} \\cdot (S_{осн.ср.} - S_{земл.}) = ' + f'{fn(chapter_2_4.OS_amortisation)}\\ [руб./год]')

    document.add_page_break()

    dp(f'НМА = {fn(chapter_2_4.NMA)} руб.')
    dp(f'Норма амортизации НМА = {fn(chapter_3.NMA_amortisation_percent * 100)}% (от изначальной суммы):')
    add_formula('A_{НМА} = k_{ам. НМА} \\cdot НМА = ' + f'{fn(chapter_2_4.NMA_amortisation)}\\ [руб./год]')

    dp(f'Планируемые расходы на ремонт основных средств примем за {fn(chapter_3.OS_fix_percent * 100)}% от стоимости ОС:')
    add_formula('S_{рем.ос} = k_{рем ос} \\cdot (S_{осн.ср.} - S_{земл.}) = ' + f'{fn(chapter_2_4.OS_fix)}\\ [руб./год]')

    dp('1. Материальные затраты:')
    add_formula('S_{мат.зат} = S_{ом} + S_{вм} + S_{т-з} S_{инстр} S_{топл.+эн.} = ' + f'{fn(chapter_2_4.costs["material"].total)}\\ [руб./год]')

    dp('2. Затраты на оплату труда:')
    add_formula('S_{ФОТ} = ' + f'{fn(chapter_2.FOT.total)}\\ [руб./год]')

    dp('3. Страховые взносы:')
    add_formula('S_{страх.вз} = ' + f'{fn(chapter_2.FOT_fee.total)}\\ [руб./год]')

    dp('4. Амортизация основных средств и нематериальных активов:')
    add_formula('A_{ОС+НМА} = A_{ос} + A_{НМА}' + f'{fn(chapter_2_4.costs["amortisation"].total)}\\ [руб./год]')

    dp(f'5. Прочие затраты (примем за {fn(chapter_3.extra_percent * 100)}% от первых 4 пунктов + планируемые расходы на ремонт ОС):')
    add_formula('S_{проч.зат.} = k_{проч.зат.} \\cdot (S_{мат.зат.} + S_{ФОТ} + S_{страх.вз} + A_{ОС+НМА}) + S_{рем.ос} = ' + f'{fn(chapter_2_4.costs["extra"].total)}\\ [руб./год]')

    costs = [
        chapter_2_4.costs['material'],
        chapter_2_4.costs['fot'],
        chapter_2_4.costs['fot fee'],
        chapter_2_4.costs['amortisation'],
        chapter_2_4.costs['extra'],
    ]

    dp()
    dp('Таблица 24.1.1, смета затрат', table_name_text)
    table = add_table(
        [['№', 'Элемент сметы', 'Сумма, руб/год', '%']] +
        [[str(i + 1), e._display_name, fn(e.total), fn(e.total / chapter_2_4.costs.total * 100)] for i, e in enumerate(costs)],
        [Cm(1), Cm(8), Cm(4), Cm(2)]
    )
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    r.cells[0].paragraphs[0].add_run('Итого: ').bold = True
    r.cells[0].paragraphs[0].add_run('затраты на производство в текущем периоде ')
    add_formula('S_{пр.тек._{пл.}} ', r.cells[0].paragraphs[0])
    r.cells[2].paragraphs[0].add_run(fn(chapter_2_4.costs.total)).bold = True
    r.cells[3].paragraphs[0].add_run('100').bold = True

    dp(f'Примем коммерческие затраты равными {fn(chapter_4.S_kom_percent * 100)}% от величины затрат на производство в текущем периоде:')
    add_formula('S_{ком} = k_{ком} \\cdot S_{пр.тек.пл.} = ' + f'{fn(chapter_2_4.S_kom.total)} [руб./год]')
    dp(f'Из них {fn(chapter_4.S_kom_const_percent * 100)}% составляют постоянные затраты, {fn(chapter_2_4.S_kom.const)}\\ [руб./год]')
    dp(f'Суммарные затраты на производство: {fn(chapter_2_4.S_sum.total)} [руб.]')

    document.add_page_break()


def gen_2_5():
    dp('5. Расчет прямых и косвенных затрат', title_text)
    dp('Деление затрат на прямые и косвенные произведем с целью их распределения между единицами калькулирования и исчисления фактической себестоимости каждой единицы.')
    dp()
    dp('5.1 Расчет прямых затрат по каждому изделию А, Б и В', subtitle_text)
    dp('Рассчитаем прямые затраты по каждому изделию А, Б и В включив в них:')
    dp('– стоимость материалов и комплектующих;')
    dp('– заработную плату основных производственных рабочих, участвующих в изготовлении изделия;')
    dp('– величину страховых взносов;')
    dp('− сумму начисленной амортизации по основным средствам, используемые при производстве продукции.')

    dp('Предполагается, что ОПР принимают участие в изготовлении всех трех наименований изделий. '
       'Поэтому заработная плата ОПР, а также страховые взносы, распределяются по изделиям '
       'пропорционально трудоемкости всей производственной программы')

    add_formula_with_description('L\'_{ОПР j} = L\'_{ОПР\\ сум} \\frac{t_j}{\\sum^m_j{tjN_{пл.j}}}', [
        ['L\'_{ОПР j}', 'заработная плата и страховые взносы ОПР, приходящиеся на одно j-ое изделие, руб./шт.'],
        ['L\'_{ОПР\\ сум}', 'суммарная заработная плата и страховые взносы ОПР, руб./год.шт.'],
        ['t_j', 'технологическая трудоёмкость j-го изделия, час./шт.'],
        ['\\sum^m_j{tjN_{пл.j}}', 'суммарная технологическая трудоемкость производственной программы, час./год']
    ])

    dp('Сумму начисленной амортизации по основным средствам, используемым при производстве продукции, '
       'распределим аналогичным образом (пропорционально трудоемкости всей производственной программы)')

    document.add_page_break()
    dp('Таблица 25.1.1, прямые затраты по изделиям А, Б и В', table_name_text)
    table = add_table([
        ['Изделие', 'А, руб/год', 'Б, руб/год', 'В, руб/год'],
        ['Материалы', fn(chapter_2_4.S_materials_i_comp_A * chapter_2_1.N_pl_A),
         fn(chapter_2_4.S_materials_i_comp_B * chapter_2_1.N_pl_B), fn(chapter_2_4.S_materials_i_comp_C * chapter_2_1.N_pl_C)],
        ['Заработная плата ОПР', fn(chapter_2_5.direct_A.FOT * chapter_2_1.N_pl_A),
         fn(chapter_2_5.direct_B.FOT * chapter_2_1.N_pl_B), fn(chapter_2_5.direct_C.FOT * chapter_2_1.N_pl_C)],
        ['Страховые взносы', fn(chapter_2_5.direct_A.FOT_fee * chapter_2_1.N_pl_A),
         fn(chapter_2_5.direct_B.FOT_fee * chapter_2_1.N_pl_B), fn(chapter_2_5.direct_C.FOT_fee * chapter_2_1.N_pl_C)],
        ['Амортизация по основным средствам, используемые при производстве продукции', fn(chapter_2_5.direct_A.amortisation * chapter_2_1.N_pl_A),
         fn(chapter_2_5.direct_B.amortisation * chapter_2_1.N_pl_B), fn(chapter_2_5.direct_C.amortisation * chapter_2_1.N_pl_C)],
        ['Суммарные прямые затраты по изделиям', fn(chapter_2_5.direct_A.direct * chapter_2_1.N_pl_A),
         fn(chapter_2_5.direct_B.direct * chapter_2_1.N_pl_B), fn(chapter_2_5.direct_C.direct * chapter_2_1.N_pl_C)],
        ['Суммарные прямые затраты', fn(chapter_2_5.direct_total), None, None]
    ], [Cm(7), Cm(3), Cm(3), Cm(3)], first_bold=True)
    table.cell(6, 1).merge(table.cell(6, 3))

    dp('Таблица 25.1.1, прямые удельные затраты по изделиям А, Б и В', table_name_text)
    table = add_table([
        ['Изделие', 'А, руб', 'Б, руб', 'В, руб'],
        ['Материалы', fn(chapter_2_4.S_materials_i_comp_A), fn(chapter_2_4.S_materials_i_comp_B), fn(chapter_2_4.S_materials_i_comp_C)],
        ['Заработная плата ОПР', fn(chapter_2_5.direct_A.FOT), fn(chapter_2_5.direct_B.FOT), fn(chapter_2_5.direct_C.FOT)],
        ['Страховые взносы', fn(chapter_2_5.direct_A.FOT_fee), fn(chapter_2_5.direct_B.FOT_fee), fn(chapter_2_5.direct_C.FOT_fee)],
        ['Амортизация по основным средствам, используемые при производстве продукции',
         fn(chapter_2_5.direct_A.amortisation), fn(chapter_2_5.direct_B.amortisation), fn(chapter_2_5.direct_C.amortisation)],
        ['Суммарные прямые затраты по изделиям', fn(chapter_2_5.direct_A.direct), fn(chapter_2_5.direct_B.direct), fn(chapter_2_5.direct_C.direct)],
        ['Суммарные прямые затраты', fn(chapter_2_5.direct_total), None, None]
    ], [Cm(7), Cm(3), Cm(3), Cm(3)], first_bold=True)
    table.cell(6, 1).merge(table.cell(6, 3))

    dp()
    dp('5.2. Расчет косвенных затрат', subtitle_text)
    dp('Для расчета косвенных затрат следует составить смету косвенных расходов, включив в нее:')
    dp('– стоимость вспомогательных материалов;')
    dp('– транспортно-заготовительные расходы;')
    dp('– стоимость инструментов, инвентаря и хозяйственных принадлежностей;')
    dp('– стоимость топлива и энергии;')
    dp('– заработную плату служащих и вспомогательных рабочих;')
    dp('– страховые взносы на указанную заработную плату;')
    dp('– амортизацию основных средств по основным средствам, не используемые при производстве продукции.')
    dp('− амортизацию нематериальных активов;')
    dp('– прочие затраты.')

    document.add_page_break()

    dp('Таблица 25.2.1, смета косвенных расходов', table_name_text)
    sm = chapter_2_5.indirect.calculate_sum(lambda x: x['cost'].total * (len(x['n']) > 0))
    add_table(
        [['Элемент сметы', 'Сумма, руб./год.', '%']] +
        [[e['name'], fn(e['cost'].total), fn(e['cost'].total / sm * 100)] for e in chapter_2_5.indirect.rows] +
        [['Итого', fn(sm), '100']],
        [Cm(9), Cm(4.5), Cm(2.5)], first_bold=True
    )

    dp()
    dp('5.3. Разделение косвенных затрат на связанные и не связанные с работой оборудования', subtitle_text)
    dp('Перед отнесением косвенных затрат по соответствующим видам продукции разделим общую сумму косвенных расходов на две '
       'составляющие косвенные затраты, связанные с работой оборудования и косвенные затраты, не связанные с работой оборудования.')
    dp('Косвенные затраты связанные с работой оборудования включают:')
    dp('– стоимость технологической энергии,')
    dp('– 60% стоимости вспомогательных материалов,')
    dp('– 80% стоимости инструментов, инвентаря, хозяйственных принадлежностей.')
    dp('– затраты на ремонт оборудования.')
    dp('Косвенные расходы не связанные с работой оборудования:')
    dp('– стоимость нетехнологической энергии,')
    dp(' – 40% стоимости вспомогательных материалов,')
    dp('– заработная плата служащих, вспомогательных производственных рабочих (ВПР) и страховые взносы,')
    dp('− 20% стоимости инструментов, инвентаря, хозяйственных принадлежностей.')
    dp('– амортизационные отчисления от стоимости основных средств (кроме стоимости оборудования) и нематериальных активов,')
    dp('– транспортно-заготовительные расходы,')
    dp('– прочие расходы (кроме затрат на ремонт оборудования).')

    dp()
    dp('Таблица 25.3.1, распределение косвенных затрат на связанные и не связанные с работой оборудования')
    table = add_table(
        [['Элемент сметы', 'Связанные с работой оборудования, руб./год', 'Не связанные с работой оборудования, руб./год']] +
        [[e['name'], fn(e['cost'].work), fn(e['cost'].other)] for e in chapter_2_5.indirect.filter(lambda x: len(x['n']) > 0)] +
        [
            ['Итого', fn(chapter_2_5.indirect.calculate_sum(lambda x: x['cost'].work * (len(x['n']) > 0))),
             fn(chapter_2_5.indirect.calculate_sum(lambda x: x['cost'].other * (len(x['n']) > 0)))],
            ['Суммарные косвенные расходы', fn(sm), None]
        ],
        [Cm(8), Cm(4.25), Cm(4.25)], first_bold=True
    )
    table.rows[-1].cells[1].merge(table.rows[-1].cells[2])
    document.add_page_break()

    dp('5.4 Распределение косвенных затрат по изделиям А, Б и В', subtitle_text)
    dp('При распределении косвенных затрат по изделиям будем придерживаться определенного порядка:')
    dp('Косвенные затраты, связанные с работой оборудования, распределим пропорционально машино-часам, затраченным на обработку годового объема '
       'выпуска каждого изделия, определив предварительно стоимость одного машино-часа.')

    add_formula_with_description(
        'S_{м-ч} = \\frac{S^{косв}_{св.\\ с\\ раб.об.}}{\\sum^k_{j=1}{N_{пл.}t\'_j}} = \\frac{' +
        f'{fn(chapter_2_5.indirect.filter_table(lambda x: len(x["n"]) > 0).calculate_sum(lambda x: x["cost"].work))} }}'
        f'{{ {fn(chapter_2_5.total_machine_time)} }} = {fn(chapter_2_5.S_m_ch)}\\ [руб./час]',
        [['t\'_j', 'технологическая трудоемкость без учета ручной операции']])

    dp()
    dp('Косвенные затраты j-го изделия связанные с работой оборудования:')
    add_formula('S^j_{св.\\ с\\ раб.об} = N_{пл.\\ j} \\cdot t\'_j \\cdot S_{м-ч}')

    dp()
    for j, n, v, g in [
        ('A', chapter_2_1.N_pl_A, chapter_2_5.machine_time_A, chapter_2_5.S_sv_s_rab_ob_A),
        ('Б', chapter_2_1.N_pl_B, chapter_2_5.machine_time_B, chapter_2_5.S_sv_s_rab_ob_B),
        ('B', chapter_2_1.N_pl_C, chapter_2_5.machine_time_C, chapter_2_5.S_sv_s_rab_ob_C)
    ]:
        add_formula(f'S^{j}_{{св.\\ с\\ раб.об}} = N_{{пл.\\ {j} }} \\cdot t\'_{j} \\cdot S_{{м-ч}} = {fn(n, 0)} \\cdot {fn(v)} \\cdot {fn(chapter_2_5.S_m_ch)} = {fn(g)}\\ [руб.]')

    dp()
    dp('Косвенные затраты, не связанные с работой оборудования, распределим пропорционально основной заработной плате основных производственных рабочих и расходам '
       'по содержанию и эксплуатации оборудования, определив предварительно коэффициент косвенных затрат.')

    add_formula('k_{косв} = \\frac{S^{косв}_{не\\ св.\\ с\\ раб.об.}}{S^{косв}_{св.\\ с\\ раб.об.} + L_{ОПР\\ сум}} = \\frac{' +
                f'{fn(chapter_2_5.indirect.filter_table(lambda x: len(x["n"]) > 0).calculate_sum(lambda x: x["cost"].other))}}}{{ '
                f'{fn(chapter_2_5.indirect.filter_table(lambda x: len(x["n"]) > 0).calculate_sum(lambda x: x["cost"].work))} + '
                f'{fn(chapter_2_3.FOT.variable)} }} = {fn(chapter_2_5.k_kosv)}')

    dp()
    for j, n, v, s, g in [
        ('A', chapter_2_1.N_pl_A, chapter_2_5.direct_A.FOT, chapter_2_5.S_sv_s_rab_ob_A, chapter_2_5.S_ne_sv_s_rab_ob_A),
        ('Б', chapter_2_1.N_pl_B, chapter_2_5.direct_B.FOT, chapter_2_5.S_sv_s_rab_ob_B, chapter_2_5.S_ne_sv_s_rab_ob_B),
        ('B', chapter_2_1.N_pl_C, chapter_2_5.direct_C.FOT, chapter_2_5.S_sv_s_rab_ob_C, chapter_2_5.S_ne_sv_s_rab_ob_C)
    ]:
        add_formula(f'S^{j}_{{не\\ св.\\ с\\ раб.об}} = (L^{j}_{{ОПР}} + S^{j}_{{св.\\ с\\ раб.об.}}) \\cdot k_{{косв}} = ({fn(n * v, 0)} +'
                    f' {fn(s)} ) \\cdot {fn(chapter_2_5.k_kosv)} = {fn(g)}\\ [руб.]', style=formula_style_12)

    document.add_page_break()

    dp('5.5 Расчет себестоимости единицы каждого вида продукции (А, Б и В)', subtitle_text)
    dp('Расчет себестоимости единицы каждого вида продукции (А, Б, В) при запланированном объеме производства выполним в таблице 24. При оформлении калькуляции для всей '
       'номенклатуры изделий укажем абсолютную величину (руб./шт.) и структуру затрат (%).')

    dp()
    dp('Таблица 25.5.1, калькуляция изделия А')
    add_production_calculation_table(chapter_2_5.direct_A, chapter_2_5.S_rab_ob_A, chapter_2_4.S_kom.total * chapter_2_1.A_percent / (max(1, chapter_2_1.N_pl_A)))

    dp()
    dp('Таблица 25.5.2, калькуляция изделия Б')
    add_production_calculation_table(chapter_2_5.direct_B, chapter_2_5.S_rab_ob_B, chapter_2_4.S_kom.total * chapter_2_1.B_percent / (max(1, chapter_2_1.N_pl_B)))

    document.add_page_break()

    dp('Таблица 25.5.3, калькуляция изделия В')
    add_production_calculation_table(chapter_2_5.direct_C, chapter_2_5.S_rab_ob_C, chapter_2_4.S_kom.total * chapter_2_1.C_percent / (max(1, chapter_2_1.N_pl_C)))

    document.add_page_break()


def gen_2_6():
    dp('6. Планирование потребности в оборотных средствах на второй период', title_text)
    dp('Расчёт аналогичен пункту 5 раздела I')

    dp('Таблица 26.1, норма запаса материалов для изделия А', table_name_text)
    add_table([['Материал', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.materials_A.filter(lambda x: x['name'] != '-')])
    dp('Таблица 26.2 норма запаса комплектующих для изделия А', table_name_text)
    add_table([['Комплектующие', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.accessories_A.filter(lambda x: x['name'] != '-')])

    dp()
    dp('Таблица 26.3, норма запаса материалов для изделия В', table_name_text)
    add_table([['Материал', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.materials_C.filter(lambda x: x['name'] != '-')])
    dp('Таблица 26.4 норма запаса комплектующих для изделия В', table_name_text)
    add_table([['Комплектующие', 'число дней']] + [[e['name'], str(e['t_zap'])] for e in initial_data.accessories_C.filter(lambda x: x['name'] != '-')])

    document.add_page_break()

    add_formula('K_{об.ср.\\ мат\\ и\\ комп.} = ' + f'{fn(chapter_2_6.K_ob_sr_mk)}\\ [руб.]')
    dp(f'Рассчитанную величину оборотных средств в запасах материалов и комплектующих изделий увеличим на {fn(chapter_2_6.k_ob_sr_percent * 100, 0)}%. '
       f'Общая сумма оборотных средств в производственных запасах составит:')
    add_formula('K_{об.ср._{пр.зап.}} = ' + f'{fn(1 + chapter_2_6.k_ob_sr_percent)} \\cdot K_{{ об.ср.\\ мат.и\\ комп}} = {fn(chapter_2_6.K_ob_sr_pr_zap)}\\ [руб.]')
    dp()

    dp('6.2. Оборотные средства в незавершенном производстве', subtitle_text)

    dp('Коэффициент считается для каждого изделия:')
    add_formula('k_{нз\\ А} = \\frac{S_{мат\\ и\\ комп\\ А} + S_{А_{произв}}}{2 S_{А_{произв}}} = \\frac{' +
                f'{fn(chapter_2_4.S_materials_i_comp_A)} + {fn(chapter_2_5.S_A_proizv)} }}{{ 2 \\cdot {fn(chapter_2_5.S_A_proizv)} }} = {fn(chapter_2_6.k_nz_A)}')
    add_formula('k_{нз\\ Б} = \\frac{S_{мат\\ и\\ комп\\ Б} + S_{А_{произв}}}{2 S_{Б_{произв}}} = \\frac{' +
                f'{fn(chapter_2_4.S_materials_i_comp_B)} + {fn(chapter_2_5.S_B_proizv)} }}{{ 2 \\cdot {fn(chapter_2_5.S_B_proizv)} }} = {fn(chapter_2_6.k_nz_B)}')
    add_formula('k_{нз\\ В} = \\frac{S_{мат\\ и\\ комп\\ В} + S_{А_{произв}}}{2 S_{В_{произв}}} = \\frac{' +
                f'{fn(chapter_2_4.S_materials_i_comp_C)} + {fn(chapter_2_5.S_C_proizv)} }}{{ 2 \\cdot {fn(chapter_2_5.S_C_proizv)} }} = {fn(chapter_2_6.k_nz_C)}')

    dp('Производственный цикл (кален. дни) – отрезок времени между началом и окончанием производственного процесса '
       'изготовления одного изделия, включающий время технологических операций; время подготовительно-заключительных операций; '
       'длительность естественных процессов и вспомогательных операций; время межоперационных и междусменных перерывов; '
       'время ожидания обработки при передаче изделий на рабочие места по партиям')

    add_formula_with_description('T_ц = \\frac{\\sum^{m}_{i=1}{t_{техн\\ i}}\\gamma_ц}{C D}\\frac{T_{пл}}{Т_{пл} - B}', [
        ['\\gamma_ц', f'соотношение между производственным циклом и суммарной технологической трудоемкостью изготовления изделия, принято за {fn(chapter_2_6.gamma_cycle, 0)}']
    ])
    add_formula(f'T_{{ц\\ A}} = {fn(chapter_2_6.T_cycle_A, 3)}\\ [дн.]')
    add_formula(f'T_{{ц\\ Б}} = {fn(chapter_2_6.T_cycle_B, 3)}\\ [дн.]')
    add_formula(f'T_{{ц\\ В}} = {fn(chapter_2_6.T_cycle_C, 3)}\\ [дн.]')

    dp('Оборотные средства, находящиеся в незавершенном производстве - сумма оборотных средств для каждого изделия:')
    add_formula('K_{об._{нез.пр.}} = ' + f'{fn(chapter_2_6.K_ob_nez_pr)}')

    document.add_page_break()

    dp('6.3. Оборотные средства в готовой продукции', subtitle_text)
    dp('Время нахождения на складе:')
    add_formula('t_{реал.} = ' + f'{chapter_2_6.t_real}\\ [дн.]')
    dp('Оборотные средства, находящиеся в готовой продукции:')
    add_formula('K_{об._{гот.прод.}} = \\frac{S_{А,Б,В_{произв}} N_{пл\\ А,Б,В}}{Т_{пл\\ А,Б,В}}t_{реал} = \\frac{' +
                f'{fn(chapter_2_5.S_A_proizv)} \\cdot {fn(chapter_2_1.N_pl_A)} }}{{ {chapter_2_1.T_pl} }} \\cdot {chapter_2_6.t_real} + \\frac{{'
                f'{fn(chapter_2_5.S_B_proizv)} \\cdot {fn(chapter_2_1.N_pl_B)} }}{{ {chapter_2_1.T_pl} }} \\cdot {chapter_2_6.t_real} + \\frac{{'
                f'{fn(chapter_2_5.S_C_proizv)} \\cdot {fn(chapter_2_1.N_pl_C)} }}{{ {chapter_2_1.T_pl} }} \\cdot {chapter_2_6.t_real} = {fn(chapter_2_6.K_ob_got_prod)}\\ [руб.]')

    dp()
    dp('6.4. Суммарная потребность в оборотных средствах', subtitle_text)
    dp('Оборотные средства включают в себе не только оборотные средства в производственных запасах, незавершенном производстве и готовой продукции, '
       'а также в расходах будущих периодов, дебиторской задолженности, краткосрочных финансовых вложениях, денежных средствах и т.п. (т.е. прочие оборотные средства).')
    dp('Для упрощения расчетов в курсовой работе зададим удельный вес стоимости производственных запасов, незавершенного производства и готовой продукции в '
       'общей сумме оборотных средств:')

    add_formula('\\gamma_{об} = ' + f'{chapter_2_6.gamma_ob}')
    dp('Суммарные оборотные средства:')
    add_formula('K_{об_{сум}} = \\frac{K_{об.ср._{пр.зап.}} + K_{об._{нез.пр.}} + K_{об._{гот.прод.}}}{\\gamma_{об}}')
    add_formula('K_{об_{сум}} = \\frac{' + f'{fn(chapter_2_6.K_ob_sr_pr_zap)} + {fn(chapter_2_6.K_ob_nez_pr)} + {fn(chapter_2_6.K_ob_got_prod)} }}'
                                           f'{{ {chapter_2_6.gamma_ob} }} = {fn(chapter_2_6.K_ob_sum)}\\ [руб.]')
    dp('Прочие оборотные средства:')
    add_formula('K_{об_{проч}} = (1 - \\gamma_{об}) \\cdot K_{об_{сум}} = ' + f'{fn(chapter_2_6.K_ob_extra)}\\ [руб.]')

    dp('Расходы будущих периодов:')
    add_formula('K_{об._{РБП}} = ' + f'{fn(chapter_2_6.K_ob_RBP)},\\ [руб.]')

    dp('Денежные средства:')
    add_formula(
        'K_{об_{ДС}} = K_{об_{сум}} - (K_{об.ср._{пр.зап.}}) + K_{об._{РБП}}) = ' +
        f'{fn(chapter_2_6.K_ob_sum)} - ({fn(chapter_2_6.K_ob_sr_pr_zap)} + {fn(chapter_2_6.K_ob_RBP)}) = {fn(chapter_2_6.K_ob_ds)}\\ [руб.]',
        style=formula_style_12)

    document.add_page_break()


def gen_2_7():
    dp('7. Плановый бухгалтерский баланс на начало второго периода', title_text)
    dp('Составим прогнозный бухгалтерский баланс на начало второго периода.')
    dp()
    dp('Рассчитаем прибыль от реализации остатков готовой продукции:')
    add_formula('Пр_{реал\\ ост} = (Ц_{реал} - S_{Б\\ произв}) \\cdot N_{ост\\ Б} = ' + f'{fn(chapter_2_7.TS_real_ost)} \\cdot {fn(chapter_8.N_ost, 0)} = {fn(chapter_2_7.P_real_ost)}')
    dp('Чистая прибыль от реализации остатков:')
    add_formula('Пр_{чист} = Пр_{реал\\ ост} - S_{налоги} = ' + f'{fn(chapter_2_7.P_real_ost)} - {fn(chapter_2_7.S_nalogi)} = {fn(chapter_2_7.P_chistaya_real_ost)}')

    dp()
    dp('Реализуем избыточное оборудование по остаточной цене:')
    dp('Таблица 27.1, реализация избыточного оборудования', table_name_text)
    add_table(
        [['Оборудование', 'Излишек, шт.', 'Остаточная стоимость, руб./шт.']] +
        [[e['name'], fn(e['extra'], 0), fn(e['cost'])] for e in chapter_2_7.to_sell.rows] +
        [['Итого', fn(chapter_2_7.to_sell.calculate_sum(lambda x: x['extra']), 0), fn(chapter_2_7.S_sell_OS)]]
    )
    dp('Примем, что избыточное оборудование будет реализовано за ' + f'{fn(chapter_2_7.S_sell_OS)} руб.')
    dp()
    dp('Потребность в дополнительном оборудовании: ' + f'{fn(chapter_2_2.new_machines_cost)} руб.')
    dp('Дополнительная потребность в ОС (помимо нового оборудования): ' +
       f'{fn(chapter_2_2.S_os - chapter_1.S_os + chapter_3.costs["amortisation OS"].total - chapter_2_2.new_machines_cost)} руб.')
    dp('Суммарная дополнительная потребность в ОС: ' + f'{fn(chapter_2_2.S_os - chapter_1.S_os + chapter_3.costs["amortisation OS"].total)} руб.')
    dp()

    document.add_page_break()
    dp('Таблица 27.2, вспомогательные расчеты для составления планового бухгалтерского баланса на начало второго периода', table_name_text)
    tbl = chapter_2_7.active_passive.to_table('Конец I периода', 'Покупка оборудования', 'Покупка ОПФ', 'Продажа готовой продукции', 'Продажа оборудования', 'Начало II периода')

    for i in range(len(tbl)):
        for j in range(len(tbl[i])):
            t = type(tbl[i][j])
            if t != str:
                if t == int or t == float:
                    tbl[i][j] = fn(tbl[i][j], 0)

    split_active = []
    split_passive = []
    separator = len(tbl[0]) // 2
    for i in tbl:
        split_active.append(i[:separator])
        split_passive.append(i[separator:])

    t1 = add_table(split_active, style=table_style_10)
    document.add_page_break()
    dp('Продолжение таблицы 27.2', table_name_text)
    t2 = add_table(split_passive, style=table_style_10)

    for i in range(len(tbl)):
        for j in range(separator):
            if type(tbl[i][j]) == str and tbl[i][j].startswith(' '):
                t1.cell(i, j).paragraphs[0].paragraph_format.left_indent = Cm(0.5)
            if type(tbl[i][j + separator]) == str and tbl[i][j + separator].startswith(' '):
                t2.cell(i, j).paragraphs[0].paragraph_format.left_indent = Cm(0.5)

    for e in [0, 1, 6, 8, 18, 20]:
        t1.cell(e, 0).paragraphs[0].runs[0].bold = True
        t1.cell(e, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for e in [0, 1, 6, 8, 10, 12, 18, 20]:
        t2.cell(e, 0).paragraphs[0].runs[0].bold = True
        t2.cell(e, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()


def gen_2_8():
    dp('8. Планирование цен изделий А, Б и В', title_text)
    dp('Плановые цены трех изделий будем устанавливать на основе следующих методов ценообразования: параметрического, методов полных и переменных затрат.')
    dp()
    dp('8.1. Расчет цен параметрическим методом', subtitle_text)
    dp('Считаем, что качество изделий оценивается двумя группами потребителей I (качество изделии А) и II (качество изделия В) по трем показателям – Х1, Х2 и Х3.')
    dp('Увеличение показателя Х2 улучшает качество изделия в целом, увеличение показателя Х3 ухудшает.')
    dp('Абсолютные значения Х1, Х2, Х3 и значения показателей важности по группам потребителей приведены в таблице:')

    dp('Таблица 28.1, исходные данные для расчета цены параметрическим методом', table_name_text)
    table = add_table(
        [['Показатели качества', 'Абсолютные значения параметров', None, None, 'Важность', None], [None, 'A', 'Б', 'B', 'I', 'II']] +
        [[e['name'], fn(e['A'], 0), fn(e['B'], 0), fn(e['C'], 0), fn(e['importance A']), fn(e['importance C'])] for e in chapter_2_8.parametric_data.rows])

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(0, 5))
    dp()
    dp('Цена определяется по формуле:')
    add_formula_with_description(
        'Ц_{А,В} = \\sum^3_{i=1}{(\\frac{X_{i\\ А,В}}{X_{i\\ Б}}) ^{type} \\cdot b_i}',
        [
            ['type', 'равен 1, если увеличение параметра улучшает качество изделия, -1 если ухудшает'],
            ['b_i', 'важность i-го параметра (суммарная важность параметров для каждого изделия равна 1)']
        ])
    add_formula(f'Ц_А = {fn(chapter_7.P_fact)} \\cdot {fn(chapter_2_8.k_param_price_A)} = {fn(chapter_2_8.param_price_A)}\\ [руб.]')
    add_formula(f'Ц_B = {fn(chapter_7.P_fact)} \\cdot {fn(chapter_2_8.k_param_price_C)} = {fn(chapter_2_8.param_price_C)}\\ [руб.]')

    document.add_page_break()

    dp('8.2 Расчет цен методом полных затрат', subtitle_text)
    dp('Расчет осуществляется для трех изделий, на основе формулы п. 7.1')
    dp('Планируемую прибыль от продаж определим на основе рентабельности продукции за первый период:')
    add_formula_with_description(
        'П_{продаж} = S_{сум\\ 2\\ пер} R_{прод\\ факт} = ' +
        f' {fn(chapter_2_4.S_sum.total)} \\cdot {fn(chapter_10.R_production_fact)} = {fn(chapter_2_8.P_prodaj)}\\ [руб.]',
        [['R_{прод\\ факт}', 'фактическая рентабельность продукции в первом периоде']])

    dp()
    dp(f'Коэффициент наценки одинаковый для всей продукции ({fn(chapter_2_8.k_nats)}):')
    for n, p, pr in [
        ['A', chapter_2_5.S_A_proizv, chapter_2_8.total_cost_price_A],
        ['Б', chapter_2_5.S_B_proizv, chapter_2_8.total_cost_price_B],
        ['В', chapter_2_5.S_C_proizv, chapter_2_8.total_cost_price_C]
    ]:
        add_formula(f'Ц_{{ {n}\\ пост\\ затр}} = {fn(p + chapter_2_4.S_kom.total / initial_data.N_pl)} \\cdot (1 + {fn(chapter_2_8.k_nats)}) = {fn(pr)}')
    dp()
    dp('Выводы:')

    document.add_page_break()

    dp('8.3 Расчет цен методом переменных затрат', subtitle_text)

    dp('Таблица 28.3.1, условно-постоянные и переменные затраты для изделия А', table_name_text)
    add_const_and_variable_costs_table(chapter_2_8.S_A_sum, chapter_2_8.S_A_sum['fot'], chapter_2_8.S_A_sum['fot fee'], style=table_style_12_dense)

    dp('Таблица 28.3.2, условно-постоянные и переменные затраты для изделия Б', table_name_text)
    add_const_and_variable_costs_table(chapter_2_8.S_B_sum, chapter_2_8.S_B_sum['fot'], chapter_2_8.S_B_sum['fot fee'], style=table_style_12_dense)

    document.add_page_break()

    dp('Таблица 28.3.3, условно-постоянные и переменные затраты для изделия В', table_name_text)
    add_const_and_variable_costs_table(chapter_2_8.S_C_sum, chapter_2_8.S_C_sum['fot'], chapter_2_8.S_C_sum['fot fee'], style=table_style_12_dense)

    dp()
    dp('Расчет осуществим для трех изделий, на основе формулы п. 7.2')
    dp(f'Коэффициент наценки одинаковый: {fn(chapter_2_8.k_nats_variable)}')
    for n, p, pr in [
        ['A', chapter_2_8.S_A_poln.variable, chapter_2_8.variable_cost_price_A],
        ['Б', chapter_2_8.S_B_poln.variable, chapter_2_8.variable_cost_price_B],
        ['В', chapter_2_8.S_C_poln.variable, chapter_2_8.variable_cost_price_C]
    ]:
        add_formula(f'Ц_{{ {n}\\ перем\\ затр}} = {fn(p)} \\cdot (1 + {fn(chapter_2_8.k_nats_variable)}) = {fn(pr)}')

    dp()
    dp()
    dp('8.4. Установленные цены изделий А, Б и В', subtitle_text)
    dp('Полученные результаты в пунктах 8.1, 8.2 и 8.3 отразим в таблице:')
    dp('Таблица 28.4.1, итоговые результаты ценообразования, полученные разными методами', table_name_text)
    table = add_table([
        ['Наименование изделия', 'Полная себестоимость, руб./шт', 'Цена, полученная методом полных затрат, руб./шт.', 'Цена, полученная методом, переменных затрат, руб./шт.',
         'Цена, полученная «ценностным» методом, руб./шт.', 'Установленная цена, руб./шт.'],
        ['A', fn(chapter_2_8.S_A_poln.total), fn(chapter_2_8.total_cost_price_A), fn(chapter_2_8.variable_cost_price_A), fn(chapter_2_8.param_price_A), fn(chapter_2_8.TS_A_plan)],
        ['Б', fn(chapter_2_8.S_B_poln.total), fn(chapter_2_8.total_cost_price_B), fn(chapter_2_8.variable_cost_price_B), fn(chapter_2_8.param_price_B), fn(chapter_2_8.TS_B_plan)],
        ['В', fn(chapter_2_8.S_C_poln.total), fn(chapter_2_8.total_cost_price_C), fn(chapter_2_8.variable_cost_price_C), fn(chapter_2_8.param_price_C), fn(chapter_2_8.TS_C_plan)],
    ], [Cm(1.5), Cm(3.2), Cm(3), Cm(3), Cm(3.2), Cm(3)], style=table_style_12)

    textDirection = OxmlElement('w:textDirection')
    textDirection.set(ns.qn('w:val'), 'btLr')  # btLr tbRl
    table.cell(0, 0)._tc.get_or_add_tcPr().append(textDirection)

    dp('В качестве установленной цены примем среднюю из полученных').add_run().add_break(WD_BREAK.PAGE)


def gen_2_9():
    dp('9. Отчет о финансовых результатах на конец второго периода', title_text)

    dp('Рассчитаем планируемую выручку от реализации:')
    add_formula('Q_{реал.\\ план} = \\sum_{i={{A,Б,В}} }{Ц_i \\cdot N_{пл\\ i}} = ' + f'{fn(chapter_2_9.Q_plan)}\\ [руб.]')

    add_formula(
        'S_{пр.гот.пр_{план}} = S_{пр.тек_{пл.}} - K_{об_{нез.пр.}} - K_{об._{гот.прод.}} = ' +
        f'{fn(chapter_2_4.S_pr_tek_pl)} - {fn(chapter_2_6.K_ob_nez_pr)} - {fn(chapter_2_6.K_ob_got_prod)} = {fn(chapter_2_9.S_pr_got_pr_plan)}', style=formula_style_12)

    dp('Прочие доходы и расходы примерно равны прочим доходам и расходам за прошлый период.')

    dp('Таблица 8, отчёт о финансовых результатах на конец первого периода', table_name_text)
    table = add_table([
        ['Наименование показателя', 'Сумма, руб/год'],
        [None, 'план'],
        ['Выручка', fn(chapter_2_9.Q_plan)],
        ['Себестоимость продаж (проданной готовой продукции)', fn(chapter_2_9.S_pr_got_pr_plan)],
        ['Валовая прибыль (убыток)', fn(chapter_2_9.S_valovaya_plan)],
        ['Коммерческие расходы', fn(chapter_2_9.S_kom_plan)],
        ['Прибыль (убыток) от продаж', fn(chapter_2_9.P_pr_plan)],
        ['Прочие доходы ', fn(0), fn(0)],
        ['Прочие расходы', fn(chapter_2_9.S_prochie_dohidy_i_rashody_plan)],
        ['Прибыль (убыток) до налогообложения', fn(chapter_2_9.P_pr_do_nalogov_plan)],
        ['Налог на прибыль', fn(chapter_2_9.nalog_na_pribil_plan)],
        ['Чистая прибыль (убыток)', fn(chapter_2_9.P_chistaya_plan)],
    ], [Cm(9), Cm(8)])
    table.cell(0, 0).merge(table.cell(1, 0))


def gen_2_10():
    dp('10. Плановый бухгалтерский баланс на конец второго периода', title_text)

    dp(f'Амортизация основных средств: {fn(chapter_2_4.costs["amortisation OS"].total)} [руб./год] (см п.3)')
    dp(f'Амортизация НМА: {fn(chapter_2_4.costs["amortisation NMA"].total)} [руб./год] (см п.3)')
    dp('Рассчитаем оборотные средства в незавершённом производстве и готовой продукции:')
    add_formula('K_{об_{нез.пр.\\ план}} = ' + f'{fn(chapter_2_6.K_ob_nez_pr)}\\ [руб.]\\ (см.\\ п.\\ 5.2)')
    add_formula('K_{об_{гот.пр.\\ план}} = ' + f'{fn(chapter_2_6.K_ob_got_prod)}\\ [руб.]\\ (см.\\ п.\\ 5.2)')

    dp()
    dp('Рассчитаем денежные средства:')
    add_formula('K_{ден.ср.\\ план} = K_{ден.ср.нач.период.} - A_{НМА} - А_{ОС} + П_{чист.\\ план} - (K_{об_{нез.пр.\\ план}} + K_{об_{гот.пр.\\ план}}) = ' +
                f'{fn(chapter_2_7.active_passive["begin II"].active.K_ob_ds)} - {fn(chapter_2_4.costs["amortisation NMA"].total)} - {fn(chapter_2_4.costs["amortisation OS"].total)} + '
                f'{fn(chapter_2_9.P_chistaya_plan)} - ({fn(chapter_2_6.K_ob_nez_pr)} + {fn(chapter_2_9.K_ob_got_prod_plan)}) = {fn(chapter_2_10.K_den_sr_plan)}', style=formula_style_12)

    dp()
    if chapter_2_7.active_passive['begin II'].passive.kratkosroch_zaem_sredstva > 0:
        if chapter_2_10.valid_to_cope_kz_plan == 'full':
            dp('Возможно полное погашение краткосрочных заёмных средств для плановой суммы денежных средств:')
            add_formula('K_{ден.ср.конец\\ план} = K_{ден.ср.\\ план} - S_{кр.заёмн.ср.} = ' +
                        f'{fn(chapter_2_10.K_den_sr_plan)} - {fn(chapter_2_7.active_passive["begin II"].passive.kratkosroch_zaem_sredstva)} = {fn(chapter_2_10.K_den_sr_konez_plan)}\\ [руб.]',
                        style=formula_style_12)
        elif chapter_2_10.valid_to_cope_kz_plan == 'part':
            dp('Возможно частичное погашение краткосрочных заёмных средств для плановой суммы '
               'денежных средств (дальнейшая генерация неверна):').runs[0].font.color.rgb = RGBColor(255, 0, 0)
            add_formula('K_{ден.ср.конец\\ план} = K_{ден.ср.\\ план} - S_{кр.заёмн.ср.} = ' + f'{fn(500_000)} \\ [руб.]')
            add_formula('S_{кр.заёмн.ср.\\ план} = ' + f'{fn(chapter_2_10.S_kratkosroch_zaem_sredstva_konez_plan)} \\ [руб.]')
        else:
            dp('Погашение краткосрочных заёмных средств невозможно для плановой суммы денежных средств, '
               'дальнейшая генерация неверна.').runs[0].font.color.rgb = RGBColor(255, 0, 0)

    document.add_page_break()

    dp('Таблица 210.1, плановый бухгалтерский баланс на конец периода', table_name_text)
    add_active_passive_table(chapter_2_10.active_passive_plan)

    document.add_page_break()


def gen_2_11():
    dp('11. Основные показатели хозяйственной деятельности предприятия на конец второго периода', title_text)
    dp('Рассчитаем основные показатели хозяйственной деятельности. Построим графики рентабельности для изделий А, Б и В, '
       'определим точку безубыточности, запас финансовой прочности, величину операционного рычага.')

    def gen_2_11_1():
        dp()
        dp('11.1 Основные показатели хозяйственной деятельности предприятия', subtitle_text)

        dp()
        dp('Сумма хозяйственных средств:')
        add_formula('K_{хс.\\ план} = ' + f'{fn(chapter_2_10.active_passive_plan.active)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Собственные оборотные средства:')
        add_formula('k_{соб.об.ср.} = Оборотные\\ активы - Краткосрочные\\ обязательства')
        add_formula('k_{соб.об.ср.\\ план} = ' +
                    f'{fn(chapter_2_10.active_passive_plan.r2)} - {fn(chapter_2_10.active_passive_plan.r5)} = {fn(chapter_2_11.k_sob_ob_sr_plan)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Коэффициент обеспеченности собственными средствами:')
        add_formula('k_{обеспеч.соб.ср.} = \\frac{Оборотные\\ активы - Краткосрочные\\ обязательства}{Оборотные\\ активы}')
        add_formula(
            'k_{обеспеч.соб.ср.\\ план} = \\frac{' +
            f'{fn(chapter_2_11.k_sob_ob_sr_plan)} }}{{ {fn(chapter_2_10.active_passive_plan.r2)} }} = {fn(chapter_2_11.k_obespech_sob_sr_plan)}', style=formula_style_12)

        dp('Коэффициент абсолютной ликвидности:')
        add_formula('k_{абс.ликв.} = \\frac{Абсолютно\\ ликвидныке\\ активы}{Краткосрочные\\ обязательства}')
        add_formula('k_{абс.ликв.\\ план} = \\frac{' +
                    f'{fn(chapter_2_10.active_passive_plan.K_ob_ds)} }}{{ {fn(chapter_2_10.active_passive_plan.r5)} }} = {fn(chapter_2_11.k_abs_likvid_plan)}', style=formula_style_12)

        dp()
        dp('Коэффициент текущей ликвидности (или коэффициент покрытия баланса):')
        add_formula('k_{тек.ликв.} = \\frac{Сумма\\ оборотных\\ активов}{Краткосрочные\\ обязательства}')
        add_formula('k_{тек.ликв.\\ план} = \\frac{' +
                    f'{fn(chapter_2_10.active_passive_plan.r2)} }}{{ {fn(chapter_2_10.active_passive_plan.r5)} }} = {fn(chapter_2_11.k_tek_likvid_plan)}', style=formula_style_12)

        document.add_page_break()

        dp('Выручка от продажи продукции:')
        add_formula('Q_{план} = ' + f'{fn(chapter_2_9.Q_plan)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Нераспределенная прибыль:')
        add_formula('П_{нерасп.\\ план} = ' + f'{fn(chapter_2_10.active_passive_plan.neraspred_pribil)}\\ [руб.]', style=formula_style_12)

        dp('Выработка продукции на одного работника:')
        add_formula('V = \\frac{Объём\\ продукции}{Среднесписочное\\ кол-во\\ ППП} = \\frac{' +
                    f'{fn(initial_data.N_pl, 0)} }}{{ {fn(chapter_2_3.R_ppp, 0)} }} = {fn(chapter_2_11.V)}\\ [шт./чел.год]', style=formula_style_12)

        dp()
        dp('Среднегодовая стоимость ОПФ:')
        add_formula(
            'S_{ср.год.ст.ОПФ} = S_{ОПФ\\ нач.пер.} - А_{ОПФ} \\cdot 0.5 = ' +
            f'{fn(chapter_2_2.S_os_amortisable)} - {fn(chapter_2_4.costs["amortisation OS"].total)} = {fn(chapter_2_11.OS_year_mean)}\\ [руб.]', style=formula_style_12)

        dp()
        dp('Коэффициент фондоотдачи:')
        add_formula('k_{ФО\\ план} = \\frac{Q_{план}}{Среднегодовая\\ стоимость\\ ОПФ} = \\frac{' +
                    f'{fn(chapter_2_9.Q_plan)} }}{{ {fn(chapter_2_11.OS_year_mean)} }} = {fn(chapter_2_11.k_FO_plan)}', style=formula_style_12)

        dp()
        dp('Коэффициент фондоемкости:')
        add_formula('k_{ФЕ\\ план} = k_{ФО\\ план}^{-1} = ' + f'{fn(chapter_2_11.k_FO_plan)} ^ {{-1}} = {fn(chapter_2_11.k_FE_plan)}', style=formula_style_12)

        dp()
        dp('Число оборотов оборотных средств:')
        add_formula('Ср.сумм.исп.об.ср._{план} = ' + f'{fn(chapter_2_11.Z_ob_sr_year_mean_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('Z_{об} = \\frac{Выручка\\ от\\ реализации}{Средняя\\ сумма\\ используемых\\ обороных\\ средств}')
        add_formula('Z_{об\\ план} = \\frac{' +
                    f'{fn(chapter_2_9.Q_plan)} }}{{ {fn(chapter_2_11.Z_ob_sr_year_mean_plan)} }} = {fn(chapter_2_11.Z_ob_sr_year_mean_plan)} [раз/год]', style=formula_style_12)

        dp('Оборачиваемость собственного капитала:')
        add_formula('Ср.год.собств.кап_{план} = ' + f'{fn(chapter_2_11.S_sobstv_cap_year_mean_plan)}\\ [руб.]', style=formula_style_12)
        add_formula('k_{об.собств.кап.} = \\frac{Выручка\\ от\\ реализации}{Ср.год.собств.кап}')
        add_formula('k_{об.собств.кап.\\ план} = \\frac{' +
                    f'{fn(chapter_8.Q_plan)} }}{{ {fn(chapter_2_11.S_sobstv_cap_year_mean_plan)} }} = {fn(chapter_2_11.k_oborach_sobstv_capital_plan)}', style=formula_style_12)

        document.add_page_break()

        dp('Рентабельность продукции:')
        add_formula('R_{продукции} = \\frac{Прибыль\\ от\\ продаж}{Себестоимость\\ продаж}')
        add_formula('R_{продукции\\ план} = \\frac{' +
                    f'{fn(chapter_2_9.P_pr_plan)} }}{{ {fn(chapter_2_4.S_sum.total)} }} = {fn(chapter_2_11.R_production_plan)}', style=formula_style_12)

        dp()
        dp('Рентабельность продаж:')
        add_formula('R_{продаж} = \\frac{Чистая\\ прибыль}{Выручка}')
        add_formula('R_{продаж\\ план} = \\frac{' +
                    f'{fn(chapter_2_9.P_chistaya_plan)} }}{{ {fn(chapter_2_9.Q_plan)} }} = {fn(chapter_2_11.R_sell_plan)}', style=formula_style_12)

        dp('Рентабельность активов:')
        add_formula('R_{активов} = \\frac{Чистая\\ прибыль}{Актив}')
        add_formula('R_{активов\\ план} = \\frac{' +
                    f'{fn(chapter_2_9.P_chistaya_plan)} }}{{ {fn(chapter_2_10.active_passive_plan.active)} }} = {fn(chapter_2_11.R_active_plan)}', style=formula_style_12)

        dp()
        dp('Рентабельность собственного капитала:')
        add_formula('R_{собств.кап.} = \\frac{Чистая\\ прибыль}{Актив}')
        add_formula('R_{собств.кап.\\ план} = \\frac{' +
                    f'{fn(chapter_2_9.P_chistaya_plan)} }}{{ {fn(chapter_2_11.S_sobstv_cap_year_mean_plan)} }} = {fn(chapter_2_11.R_sobstv_capital_plan)}', style=formula_style_12)

        document.add_page_break()

        dp('Таблица 211.1, плановые значения основных показателей хозяйственной деятельности предприятия во II периоде', table_name_text)
        add_table([
            ['Наименование показателя и его размерность', 'План'],
            ['Сумма хозяйственных средств, [руб.]', fn(chapter_2_10.active_passive_plan.active)],
            ['Собственные оборотные средства, [руб.]', fn(chapter_2_11.k_sob_ob_sr_plan)],
            ['Коэффициент обеспеченности собственными средствами', fn(chapter_2_11.k_obespech_sob_sr_plan)],
            ['Коэффициент абсолютной ликвидности', fn(chapter_2_11.k_abs_likvid_plan)],
            ['Коэффициент текущей ликвидности', fn(chapter_2_11.k_tek_likvid_plan)],
            ['Выручка от продажи продукции, [руб.]', fn(chapter_2_9.Q_plan)],
            ['Нераспределенная прибыль, [руб.]', fn(chapter_2_10.active_passive_plan.neraspred_pribil)],
            ['Выработка продукции на одного работника [шт./чел.год]', fn(chapter_2_11.V)],
            ['Среднегодовая стоимость ОПФ, [руб.]', fn(chapter_2_11.OS_year_mean)],
            ['Коэффициент фондоотдачи [1/руб.]', fn(chapter_2_11.k_FO_plan), ],
            ['Коэффициент фондоемкости [руб.]', fn(chapter_2_11.k_FE_plan)],
            ['Число оборотов оборотных средств, [раз/год]', fn(chapter_2_11.Z_ob_plan)],
            ['Оборачиваемость собственного капитала', fn(chapter_2_11.k_oborach_sobstv_capital_plan)],
            ['Рентабельность продукции', fn(chapter_2_11.R_production_plan)],
            ['Рентабельность продаж', fn(chapter_2_11.R_sell_plan)],
            ['Рентабельность активов [1/год]', fn(chapter_2_11.R_active_plan)],
            ['Рентабельность собственного капитала [1/год]', fn(chapter_2_11.R_sobstv_capital_plan)],
        ], [Cm(9.25), Cm(7.5)], True, style=table_style_12)

        document.add_page_break()

    def gen_2_11_2():
        dp('11.2 Графики рентабельности изделий', subtitle_text)
        dp('В данном случае примем что мы не изменяем численность ОПР, необходимую для производства N изделий, а следовательно их премии и стимулирующие выплаты '
           'останутся неизменными, и мы сможем воспользоваться предложенными формулами.')
        dp('Точка безубыточности:')
        add_formula('N_{кр} = \\frac{S_{усл.пост.}}{Ц_{произв\\ план} - S_{перем}}')
        dp('Коэффициент покрытия:')
        add_formula('k_{покр} = \\frac{Ц_{произв\\ план} - S_{перем}}{Ц_{произв\\ план}')

        document.add_page_break()

        def gen_2_11_2_A():
            dp('11.2.1 для изделия А', subtitle_2_text)
            dp('Точка безубыточности:')
            add_formula('N_{кр\\ A} = \\frac{' +
                        f'{fn(chapter_2_8.S_A_sum.const)} }}{{ {fn(chapter_2_8.TS_A_plan)} - {fn(chapter_2_8.S_A_poln.variable)} }} = {fn(chapter_2_11.N_kr_A, 0)}\\ [шт./год]')
            add_formula('Q_{кр\\ A} = N_{кр\\ A} \\cdot Ц_{A\\ произв\\ план} = ' + f'{fn(chapter_2_11.Q_kr_A)}\\ [руб.]')
            dp('Коэффициент покрытия:')
            add_formula('k_{покр\\ A} = \\frac{' + f'{fn(chapter_2_8.TS_A_plan)} - {fn(chapter_2_8.S_A_poln.variable)} }}{{ {fn(chapter_2_8.TS_A_plan)} }} = {fn(chapter_2_11.k_pokr_A)}')
            dp('Запас финансовой прочности:')
            add_formula('Q_{фин\\ пр.\\ A} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_A)} - {fn(chapter_2_11.Q_kr_A)} }}{{ {fn(chapter_2_9.Q_plan_A)} }} = {fn(chapter_2_11.Q_fin_pr_A * 100)}\\%')

            dp('Эффект производственного рычага:')
            add_formula('E_{пр.\\ рыч.\\ план\\ А} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_A)} - {fn(chapter_2_8.S_A_sum.variable)} }}{{ {fn(chapter_2_9.P_pr_plan)} }} = {fn(chapter_2_11.proizv_richag_A)}')

            dp('График 211.2.1, рентабельность изделия А', table_name_text)

            plt.figure(figsize=(8, 8))
            plt.subplot(1, 1, 1)

            plt.title('S(N)')
            plt.xlabel('N, шт. / год')
            plt.ylabel('Выручка, затраты, тыс. руб./год')
            plt.xticks([0, chapter_2_11.N_kr_A, chapter_2_1.N_pl_A], ['0', 'N кр\n{:,.0f}'.format(chapter_2_11.N_kr_A), 'N пл\n{:,.0f}'.format(chapter_2_1.N_pl_A)], rotation=0)
            plt.yticks([0, chapter_2_8.S_A_sum.const / 1e3, chapter_2_11.Q_kr_A / 1e3, chapter_2_8.S_A_sum.total / 1e3, chapter_2_9.Q_plan_A / 1e3],
                       ['0', 'S усл.пост.\n{:,.0f}'.format(chapter_2_8.S_A_sum.total / 1e3), 'Q кр.\n{:,.0f}'.format(chapter_10.Q_kr / 1e3),
                        '{:,.0f}\nS сум.'.format(chapter_2_8.S_A_sum.total / 1e3), 'Q пл\n{:,.0f}'.format(chapter_2_9.Q_plan_A / 1e3)])
            plt.grid(True)
            plt.plot([0, chapter_2_1.N_pl_A], np.array([chapter_2_8.S_A_sum.const, chapter_2_8.S_A_sum.const]) / 1e3, label='S усл-пост.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_A], np.array([0, chapter_2_8.S_A_sum.variable]) / 1e3, label='S перем.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_A], np.array([chapter_2_8.S_A_sum.const, chapter_2_8.S_A_sum.total]) / 1e3, label='S тек.сум.')
            plt.plot([0, chapter_2_1.N_pl_A], [0, chapter_2_9.Q_plan_A / 1e3], label='Q пл.')
            plt.legend()
            plt.tight_layout()

            memfile = BytesIO()
            plt.savefig(memfile)

            picP = document.add_paragraph()
            picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            picP.add_run().add_picture(memfile, width=Cm(14))
            picP.add_run().add_break(WD_BREAK.PAGE)

        def gen_2_11_2_B():
            dp('11.2.1 для изделия Б', subtitle_2_text)
            dp('Точка безубыточности:')
            add_formula('N_{кр\\ Б} = \\frac{' +
                        f'{fn(chapter_2_8.S_B_sum.const)} }}{{ {fn(chapter_2_8.TS_B_plan)} - {fn(chapter_2_8.S_B_poln.variable)} }} = {fn(chapter_2_11.N_kr_B, 0)}\\ [шт./год]')
            add_formula('Q_{кр\\ Б} = N_{кр\\ Б} \\cdot Ц_{Б\\ произв\\ план} = ' + f'{fn(chapter_2_11.Q_kr_B)}\\ [руб.]')
            dp('Коэффициент покрытия:')
            add_formula('k_{покр\\ Б} = \\frac{' + f'{fn(chapter_2_8.TS_B_plan)} - {fn(chapter_2_8.S_B_poln.variable)} }}{{ {fn(chapter_2_8.TS_B_plan)} }} = {fn(chapter_2_11.k_pokr_B)}')
            dp('Запас финансовой прочности:')
            add_formula('Q_{фин\\ пр.\\ Б} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_B)} - {fn(chapter_2_11.Q_kr_B)} }}{{ {fn(chapter_2_9.Q_plan_B)} }} = {fn(chapter_2_11.Q_fin_pr_B * 100)}\\%')

            dp('Эффект производственного рычага:')
            add_formula('E_{пр.\\ рыч.\\ план\\ Б} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_B)} - {fn(chapter_2_8.S_B_sum.variable)} }}{{ {fn(chapter_2_9.P_pr_plan)} }} = {fn(chapter_2_11.proizv_richag_B)}')

            dp('График 211.2.2, рентабельность изделия Б', table_name_text)

            plt.figure(figsize=(8, 8))
            plt.subplot(1, 1, 1)

            plt.title('S(N)')
            plt.xlabel('N, шт. / год')
            plt.ylabel('Выручка, затраты, тыс. руб./год')
            plt.xticks([0, chapter_2_11.N_kr_B, chapter_2_1.N_pl_B], ['0', 'N кр\n{:,.0f}'.format(chapter_2_11.N_kr_B), 'N пл\n{:,.0f}'.format(chapter_2_1.N_pl_B)], rotation=0)
            plt.yticks([0, chapter_2_8.S_B_sum.const / 1e3, chapter_2_11.Q_kr_B / 1e3, chapter_2_8.S_B_sum.total / 1e3, chapter_2_9.Q_plan_B / 1e3],
                       ['0', 'S усл.пост.\n{:,.0f}'.format(chapter_2_8.S_B_sum.total / 1e3), 'Q кр.\n{:,.0f}'.format(chapter_10.Q_kr / 1e3),
                        '{:,.0f}\nS сум.'.format(chapter_2_8.S_B_sum.total / 1e3), 'Q пл\n{:,.0f}'.format(chapter_2_9.Q_plan_B / 1e3)])
            plt.grid(True)
            plt.plot([0, chapter_2_1.N_pl_B], np.array([chapter_2_8.S_B_sum.const, chapter_2_8.S_B_sum.const]) / 1e3, label='S усл-пост.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_B], np.array([0, chapter_2_8.S_B_sum.variable]) / 1e3, label='S перем.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_B], np.array([chapter_2_8.S_B_sum.const, chapter_2_8.S_B_sum.total]) / 1e3, label='S тек.сум.')
            plt.plot([0, chapter_2_1.N_pl_B], [0, chapter_2_9.Q_plan_B / 1e3], label='Q пл.')
            plt.legend()
            plt.tight_layout()

            memfile = BytesIO()
            plt.savefig(memfile)

            picP = document.add_paragraph()
            picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            picP.add_run().add_picture(memfile, width=Cm(14))
            picP.add_run().add_break(WD_BREAK.PAGE)

        def gen_2_11_2_C():
            dp('11.2.1 для изделия В', subtitle_2_text)
            dp('Точка безубыточности:')
            add_formula('N_{кр\\ В} = \\frac{' +
                        f'{fn(chapter_2_8.S_C_sum.const)} }}{{ {fn(chapter_2_8.TS_C_plan)} - {fn(chapter_2_8.S_C_poln.variable)} }} = {fn(chapter_2_11.N_kr_C, 0)}\\ [шт./год]')
            add_formula('Q_{кр\\ В} = N_{кр\\ В} \\cdot Ц_{В\\ произв\\ план} = ' + f'{fn(chapter_2_11.Q_kr_C)}\\ [руб.]')
            dp('Коэффициент покрытия:')
            add_formula('k_{покр\\ В} = \\frac{' + f'{fn(chapter_2_8.TS_C_plan)} - {fn(chapter_2_8.S_C_poln.variable)} }}{{ {fn(chapter_2_8.TS_C_plan)} }} = {fn(chapter_2_11.k_pokr_C)}')
            dp('Запас финансовой прочности:')
            add_formula('Q_{фин\\ пр.\\ В} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_C)} - {fn(chapter_2_11.Q_kr_C)} }}{{ {fn(chapter_2_9.Q_plan_C)} }} = {fn(chapter_2_11.Q_fin_pr_C * 100)}\\%')

            dp('Эффект производственного рычага:')
            add_formula('E_{пр.\\ рыч.\\ план\\ В} = \\frac{' +
                        f'{fn(chapter_2_9.Q_plan_C)} - {fn(chapter_2_8.S_C_sum.variable)} }}{{ {fn(chapter_2_9.P_pr_plan)} }} = {fn(chapter_2_11.proizv_richag_C)}')

            dp('График 211.2.3, рентабельность изделия В', table_name_text)

            plt.figure(figsize=(8, 8))
            plt.subplot(1, 1, 1)

            plt.title('S(N)')
            plt.xlabel('N, шт. / год')
            plt.ylabel('Выручка, затраты, тыс. руб./год')
            plt.xticks([0, chapter_2_11.N_kr_C, chapter_2_1.N_pl_C], ['0', 'N кр\n{:,.0f}'.format(chapter_2_11.N_kr_C), 'N пл\n{:,.0f}'.format(chapter_2_1.N_pl_C)], rotation=0)
            plt.yticks([0, chapter_2_8.S_C_sum.const / 1e3, chapter_2_11.Q_kr_C / 1e3, chapter_2_8.S_C_sum.total / 1e3, chapter_2_9.Q_plan_C / 1e3],
                       ['0', 'S усл.пост.\n{:,.0f}'.format(chapter_2_8.S_C_sum.total / 1e3), 'Q кр.\n{:,.0f}'.format(chapter_10.Q_kr / 1e3),
                        '{:,.0f}\nS сум.'.format(chapter_2_8.S_C_sum.total / 1e3), 'Q пл\n{:,.0f}'.format(chapter_2_9.Q_plan_C / 1e3)])
            plt.grid(True)
            plt.plot([0, chapter_2_1.N_pl_C], np.array([chapter_2_8.S_C_sum.const, chapter_2_8.S_C_sum.const]) / 1e3, label='S усл-пост.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_C], np.array([0, chapter_2_8.S_C_sum.variable]) / 1e3, label='S перем.', ls=':')
            plt.plot([0, chapter_2_1.N_pl_C], np.array([chapter_2_8.S_C_sum.const, chapter_2_8.S_C_sum.total]) / 1e3, label='S тек.сум.')
            plt.plot([0, chapter_2_1.N_pl_C], [0, chapter_2_9.Q_plan_C / 1e3], label='Q пл.')
            plt.legend()
            plt.tight_layout()

            memfile = BytesIO()
            plt.savefig(memfile)

            picP = document.add_paragraph()
            picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            picP.add_run().add_picture(memfile, width=Cm(14))
            picP.add_run().add_break(WD_BREAK.PAGE)

        gen_2_11_2_A()
        gen_2_11_2_B()
        gen_2_11_2_C()

    gen_2_11_1()
    gen_2_11_2()

    dp('Таблица 211.3, плановые и фактические значения основных показателей хозяйственной деятельности предприятия за I и II период', table_name_text)
    table = add_table([
        ['Наименование показателя и его размерность', 'Значение', None, None],
        [None, 'I период', None, 'II период'],
        [None, 'план', 'факт', 'план'],
        ['Сумма хозяйственных средств, [руб.]', fn(chapter_9.active_passive_plan.active), fn(chapter_9.active_passive_fact.active), fn(chapter_2_10.active_passive_plan.active)],
        ['Собственные оборотные средства, [руб.]', fn(chapter_10.k_sob_ob_sr_plan), fn(chapter_10.k_sob_ob_sr_fact), fn(chapter_2_11.k_sob_ob_sr_plan)],
        ['Коэффициент обеспеченности собственными средствами', fn(chapter_10.k_obespech_sob_sr_plan), fn(chapter_10.k_obespech_sob_sr_fact), fn(chapter_2_11.k_obespech_sob_sr_plan)],
        ['Коэффициент абсолютной ликвидности', fn(chapter_10.k_abs_likvid_plan), fn(chapter_10.k_abs_likvid_fact), fn(chapter_2_11.k_abs_likvid_plan)],
        ['Коэффициент текущей ликвидности', fn(chapter_10.k_tek_likvid_plan), fn(chapter_10.k_tek_likvid_fact), fn(chapter_2_11.k_tek_likvid_plan)],
        ['Выручка от продажи продукции, [руб.]', fn(chapter_8.Q_plan), fn(chapter_8.Q_fact), fn(chapter_2_9.Q_plan)],
        ['Прибыль до налогообложения, [руб.]', fn(chapter_8.P_pr_do_nalogov_plan), fn(chapter_8.P_pr_do_nalogov_fact), fn(chapter_2_9.P_pr_do_nalogov_plan)],
        ['Чистая прибыль, [руб.]', fn(chapter_8.P_chistaya_plan), fn(chapter_8.P_chistaya_fact), fn(chapter_2_9.P_chistaya_plan)],
        ['Нераспределенная прибыль, [руб.]', fn(chapter_9.active_passive_plan.neraspred_pribil), fn(chapter_9.active_passive_fact.neraspred_pribil), fn(chapter_2_10.active_passive_plan.neraspred_pribil)],
        ['Выработка продукции на одного работника [шт./чел.год]', fn(chapter_10.V), fn(chapter_10.V), fn(chapter_2_11.V)],
        ['Среднегодовая стоимость ОПФ, [руб.]', fn(chapter_10.OS_year_mean), fn(chapter_10.OS_year_mean), fn(chapter_2_11.OS_year_mean)],
        ['Коэффициент фондоотдачи [1/руб.]', fn(chapter_10.k_FO_plan), fn(chapter_10.k_FO_fact), fn(chapter_2_11.k_FO_plan), ],
        ['Коэффициент фондоемкости [руб.]', fn(chapter_10.k_FE_plan), fn(chapter_10.k_FE_fact), fn(chapter_2_11.k_FE_plan)],
        ['Число оборотов оборотных средств, [раз/год]', fn(chapter_10.Z_ob_plan), fn(chapter_10.Z_ob_fact), fn(chapter_2_11.Z_ob_plan)],
        ['Оборачиваемость собственного капитала', fn(chapter_10.k_oborach_sobstv_capital_plan), fn(chapter_10.k_oborach_sobstv_capital_fact), fn(chapter_2_11.k_oborach_sobstv_capital_plan)],
        ['Рентабельность продукции', fn(chapter_10.R_production_plan), fn(chapter_10.R_production_fact), fn(chapter_2_11.R_production_plan)],
        ['Рентабельность продаж', fn(chapter_10.R_sell_plan), fn(chapter_10.R_sell_fact), fn(chapter_2_11.R_sell_plan)],
        ['Рентабельность активов [1/год]', fn(chapter_10.R_active_plan), fn(chapter_10.R_active_fact), fn(chapter_2_11.R_active_plan)],
        ['Рентабельность собственного капитала [1/год]', fn(chapter_10.R_sobstv_capital_plan), fn(chapter_10.R_sobstv_capital_fact), fn(chapter_2_11.R_sobstv_capital_plan)],
    ], [Cm(7.5), Cm(3.3), Cm(3.3), Cm(3.3)], True, style=table_style_12)
    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(1, 1).merge(table.cell(1, 2))

    document.add_page_break()


def gen_final():
    dp('Заключение', title_text)


def main():
    init_styles()
    p = gen_first_list()
    gen_introduction()
    gen_initial_data()
    gen_1_1()
    gen_1_2()
    gen_1_3()
    gen_1_4()
    gen_1_5()
    gen_1_6()
    gen_1_7()
    gen_1_8()
    gen_1_9()
    gen_1_10()
    gen_1_11()

    gen_2_1()
    gen_2_2()
    gen_2_3()
    gen_2_4()
    gen_2_5()
    gen_2_6()
    gen_2_7()
    gen_2_8()
    gen_2_9()
    gen_2_10()
    gen_2_11()
    gen_final()

    return p


def add_page_numbers(paragraph):
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(fldChar4)


def add_table_of_content(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(ns.qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(ns.qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p


if __name__ == '__main__':
    import time

    t = time.time()
    paragraph = main()

    if paragraph is not None:
        document.sections[1].footer.is_linked_to_previous = False
        add_page_numbers(document.sections[1].footer.paragraphs[0])
        add_table_of_content(paragraph)

    sections = document.sections
    for section in sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    document.save('out.docx')
    print('time: {:,.2f}'.format(time.time() - t))
