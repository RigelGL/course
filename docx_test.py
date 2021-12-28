import latex2mathml.converter
import numpy as np
from lxml import etree
from matplotlib import pyplot as plt
from math import sin, cos, sqrt
from io import BytesIO
from docx import *
from docx.shared import *
from docx.enum.text import *
from docx.enum.section import *
from docx.enum.table import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def latex_to_word(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def main():
    d = Document()
    text = 'Lorem ipsum dolor sit amet, И еще немного рыбного текста'\
           'безо всяеого смвсла просто для запослнения текстом нескольких строк'
    p1 = d.add_paragraph()
    pf1 = p1.paragraph_format
    pf1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pf1.left_indent = Pt(30)
    pf1.space_before = Pt(20)
    pf1.line_spacing = 2.
    r1 = p1.add_run(text)
    r1.font.name = 'Arial'
    r1.font.size = Pt(24)
    p2 = d.add_paragraph()
    p2.paragraph_format.page_break_before = False
    r2 = p2.add_run('Lorem?')
    r2.font.color.rgb = RGBColor(0x92, 0x56, 0x78)

    d.add_paragraph().add_run()._element.append(latex_to_word('S_x=\\sqrt{\\frac{1}{n-1}\\sum_{i=1}^{n}{X_i^2-nX^2}}'))
    d.add_paragraph().add_run()._element.append(latex_to_word(
        'y=\\frac{\\sqrt[3]{(a+b)}}{\\int_{1}^{2}{\\sin{(x)}},dx}+\\sum_{i=1}^{10}{\\frac{x_i^2}{(x_i-1)^3}}'))

    s2 = d.add_section()
    s2.orientation = WD_ORIENTATION.LANDSCAPE
    s2.page_width, s2.page_height = s2.page_height, s2.page_width
    d.add_paragraph('Horizontal?')
    s2.header.is_linked_to_previous = False
    s2hp = s2.header.paragraphs[0]
    s2hr = s2hp.add_run('Horizontal')
    s2hr.font.name = 'Garamond'
    s2hr.font.size = Pt(18)
    s2hr.font.color.rgb = RGBColor(0xA8, 0xA8, 0xDF)
    s2hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    memfile = BytesIO()
    plt.plot([i for i in range(10)], [sin(i) for i in range(10)])
    plt.savefig(memfile)

    picP = d.add_paragraph()
    picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picP.add_run().add_picture(memfile, width=Mm(100))

    memfile.seek(0)
    labels = ['Доходы', 'Расходы', 'Сальдо']
    dohodi = [20000, 18000, 19600]
    rashodi = -np.array([19000, 16000, 18900])
    saldo = [dohodi[i] + rashodi[i] for i in range(len(dohodi))]

    x = np.arange(len(labels))
    width = (1.0 / len(labels)) * 0.8

    fig, ax = plt.subplots()
    ax.bar(x - width, [e[0] for e in [dohodi, rashodi, saldo]], width, label='1 Месяц')
    ax.bar(x, [e[1] for e in [dohodi, rashodi, saldo]], width, label='2 Месяц')
    ax.bar(x + width, [e[2] for e in [dohodi, rashodi, saldo]], width, label='3 Месяц')

    ax.set_ylabel('Денежные средства, руб.')
    ax.set_title('Расчёт сальдо за три месяца')
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.grid(axis='y')
    ax.legend()

    fig.tight_layout()
    plt.savefig(memfile)

    picP = d.add_paragraph()
    picP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picP.add_run().add_picture(memfile, width=Mm(100))

    s3 = d.add_section()
    s3.orientation = WD_ORIENTATION.LANDSCAPE
    s2.page_width, s2.page_height = s2.page_height, s2.page_width
    d.add_paragraph('Таблицы')
    s3.header.is_linked_to_previous = False
    s3hp = s3.header.paragraphs[0]
    s3hr = s3hp.add_run('Таблицы')
    s3hr.font.name = 'Times New Roman'
    s3hr.font.size = Pt(18)
    s3hr.font.color.rgb = RGBColor(0x00, 0xA8, 0xDF)
    s3hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table1 = d.add_table(rows=5, cols=4)
    table1.alignment = WD_TAB_ALIGNMENT.CENTER
    table1.style = 'Light Grid Accent 5'
    table1_data = [
        ['', '2018 г, тыс. руб', '2019 г, тыс.руб', 'Сумма'],
        ['Комплекты', 5650, 7402, ''],
        ['Фурнитура', 3630, 4510, ''],
        ['Итог', '', '', ''],
    ]

    for i in table1_data[1:-1]:
        i[3] = i[1] + i[2]
    for i in range(1, 3):
        table1_data[3][i] = table1_data[1][i] + table1_data[2][i]
    table1_data[3][3] = table1_data[1][3] + table1_data[2][3]

    table1.cell(4, 0).merge(table1.cell(4, 3))
    title_r = table1.cell(4, 0).paragraphs[0].add_run('Title')
    title_r.font.size = Pt(12)
    title_r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    title_r.font.bold = False
    set_cell_border(table1.cell(4, 0),
        bottom={"color": "#FFFFFF"},
        start={"color": "#FFFFFF"},
        end={"color": "#FFFFFF"},
    )

    for i in range(0, 4):
        for j in range(0, 4):
            p = table1.cell(i, j).paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Pt(6)
            p.paragraph_format.right_indent = Pt(6)
            r = p.add_run(str(table1_data[i][j]))
            r.font.size = Pt(18)
            table1.cell(i, j).width = Mm(35)
            if type(table1_data[i][j]) is int:
                if table1_data[i][j] > 10000:
                    tblCellProperties = table1.cell(i, j)._tc.get_or_add_tcPr()
                    clShading = OxmlElement('w:shd')
                    clShading.set(qn('w:fill'), "CC8888")
                    tblCellProperties.append(clShading)


    d.save('demo.docx')
    memfile.close()


if __name__ == '__main__':
    main()
